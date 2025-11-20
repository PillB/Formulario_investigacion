#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Nov 1 11:04:14 2025

@author: pabloillescasbuendia
"""
"""
app.py
==================

Esta aplicación implementa una versión de escritorio de la herramienta de
gestión de casos de fraude que previamente se desarrolló como sitio web.

Utiliza la biblioteca estándar ``tkinter`` para construir una interfaz de
usuario gráfica sin necesidad de arrancar un servidor web local. La
aplicación permite crear casos con sus clientes, colaboradores, productos,
riesgos identificados, normas transgredidas y narrativas. Soporta la
adición dinámica de registros, importación desde ficheros CSV, validación
de reglas de negocio, auto‑guardado del estado en un archivo JSON y
exportación de los datos normalizados en múltiples ficheros CSV. Al
almacenar cada entidad en tablas separadas y vincularlas mediante claves
primarias y foráneas, los datos pueden cargarse posteriormente en otros
sistemas o retomarse en otra estación de trabajo.

Para ejecutar la aplicación basta con tener Python 3 instalado; no se
requieren bibliotecas externas. La interfaz se abre en una ventana
independiente. Los datos se validan a medida que el usuario escribe y
se registran en un log interno para fines de análisis y mejora.

Funciones principales:
  * Gestión de los datos del caso (tipo de informe, categorías y
    modalidades, canal y proceso).
  * Gestión dinámica de clientes, colaboradores y productos con listas
    anidadas para teléfonos, correos, direcciones y asignaciones de
    montos.
  * Importación de clientes, colaboradores, productos y registros
    combinados desde ficheros CSV.
  * Auto‑poblado de datos de colaboradores a partir de un fichero
    ``team_details.csv`` si está disponible.
  * Validación de reglas de negocio (unicidad de combinaciones,
    coherencia de fechas y montos, requisitos de reclamos y analíticas,
    patrones de identificadores, etc.).
  * Auto‑guardado del estado del formulario en un archivo JSON cada vez
    que cambian los datos y restauración desde dicho archivo.
  * Exportación de los datos normalizados en varios CSV con la misma
    estructura que el archivo Excel de referencia.
  * Registro de eventos y errores en un log que se descarga junto con
    los demás ficheros.

Nota: el diseño de la interfaz y de las funciones de validación está
orientado a mantener la mayor parte de la funcionalidad del sitio web
original, pero por simplicidad visual se utilizan cuadros de texto
separados por punto y coma para campos con múltiples valores (teléfonos,
correos, direcciones y planes de acción).
"""

import csv
import json
import os
import shutil
import threading
import zipfile
from datetime import datetime
from decimal import Decimal
from pathlib import Path
from textwrap import dedent
from typing import Optional
from xml.sax.saxutils import escape

import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, ttk

try:  # python-docx es una dependencia opcional en tiempo de pruebas
    from docx import Document as DocxDocument
except ImportError:  # pragma: no cover - se usa la ruta de respaldo
    DocxDocument = None

from models import (build_detail_catalog_id_index, iter_massive_csv_rows,
                    load_detail_catalogs, normalize_detail_catalog_key,
                    parse_involvement_entries)
from settings import (AUTOSAVE_FILE, BASE_DIR, CANAL_LIST, CLIENT_ID_ALIASES,
                      CRITICIDAD_LIST, DETAIL_LOOKUP_ALIASES,
                      EXPORTS_DIR, EXTERNAL_LOGS_FILE, FLAG_CLIENTE_LIST,
                      FLAG_COLABORADOR_LIST, LOGS_FILE, STORE_LOGS_LOCALLY,
                      MASSIVE_SAMPLE_FILES, NORM_ID_ALIASES, PROCESO_LIST,
                      PRODUCT_ID_ALIASES, RISK_ID_ALIASES, TAXONOMIA,
                      TEAM_ID_ALIASES, TIPO_FALTA_LIST, TIPO_ID_LIST,
                      TIPO_INFORME_LIST, TIPO_MONEDA_LIST,
                      TIPO_PRODUCTO_LIST, TIPO_SANCION_LIST,
                      ensure_external_drive_dir)
from ui.frames import (ClientFrame, NormFrame, PRODUCT_MONEY_SPECS,
                       ProductFrame, RiskFrame, TeamMemberFrame)
from ui.tooltips import HoverTooltip
from validators import (drain_log_queue, FieldValidator, log_event,
                        normalize_without_accents, parse_decimal_amount,
                        resolve_catalog_product_type, should_autofill_field,
                        sum_investigation_components, TIPO_PRODUCTO_NORMALIZED,
                        validate_agency_code, validate_case_id,
                        validate_client_id, validate_codigo_analitica,
                        validate_date_text, validate_email_list,
                        validate_money_bounds, validate_multi_selection,
                        validate_norm_id, validate_phone_list,
                        validate_product_dates, validate_product_id,
                        validate_reclamo_id, validate_required_text,
                        validate_risk_id, validate_team_member_id)


_DOCX_STATIC_PARTS = {
    '[Content_Types].xml': dedent('''
        <?xml version="1.0" encoding="UTF-8"?>
        <Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
            <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
            <Default Extension="xml" ContentType="application/xml"/>
            <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
            <Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>
            <Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/>
            <Override PartName="/docProps/app.xml" ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/>
        </Types>
    ''').strip().encode('utf-8'),
    '_rels/.rels': dedent('''
        <?xml version="1.0" encoding="UTF-8"?>
        <Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
            <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
            <Relationship Id="rId2" Type="http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties" Target="docProps/core.xml"/>
            <Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties" Target="docProps/app.xml"/>
        </Relationships>
    ''').strip().encode('utf-8'),
    'docProps/core.xml': dedent('''
        <?xml version="1.0" encoding="UTF-8" standalone="yes"?>
        <cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
            xmlns:dc="http://purl.org/dc/elements/1.1/"
            xmlns:dcterms="http://purl.org/dc/terms/"
            xmlns:dcmitype="http://purl.org/dc/dcmitype/"
            xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">
            <dc:title>Informe de caso</dc:title>
            <cp:revision>1</cp:revision>
        </cp:coreProperties>
    ''').strip().encode('utf-8'),
    'docProps/app.xml': dedent('''
        <?xml version="1.0" encoding="UTF-8" standalone="yes"?>
        <Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"
            xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">
            <Application>Formulario Investigacion</Application>
        </Properties>
    ''').strip().encode('utf-8'),
    'word/styles.xml': dedent('''
        <?xml version="1.0" encoding="UTF-8" standalone="yes"?>
        <w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:style w:type="paragraph" w:default="1" w:styleId="Normal">
                <w:name w:val="Normal"/>
            </w:style>
        </w:styles>
    ''').strip().encode('utf-8'),
}


class _FallbackDocxCell:
    __slots__ = ('text',)

    def __init__(self):
        self.text = ''

    def to_xml(self):
        if not self.text:
            return '<w:tc><w:p/></w:tc>'
        safe = escape(self.text)
        return f'<w:tc><w:p><w:r><w:t xml:space="preserve">{safe}</w:t></w:r></w:p></w:tc>'


class _FallbackDocxRow:
    __slots__ = ('cells',)

    def __init__(self, cols):
        self.cells = [_FallbackDocxCell() for _ in range(cols)]

    def to_xml(self):
        return '<w:tr>' + ''.join(cell.to_xml() for cell in self.cells) + '</w:tr>'


class _FallbackDocxTable:
    __slots__ = ('rows', '_cols', 'style')

    def __init__(self, rows, cols):
        self._cols = cols
        self.rows = [_FallbackDocxRow(cols) for _ in range(rows)]
        self.style = None

    def add_row(self):
        row = _FallbackDocxRow(self._cols)
        self.rows.append(row)
        return row

    def to_xml(self):
        grid = ''.join('<w:gridCol w:w="2400"/>' for _ in range(self._cols))
        rows_xml = ''.join(row.to_xml() for row in self.rows)
        return f'<w:tbl><w:tblPr/><w:tblGrid>{grid}</w:tblGrid>{rows_xml}</w:tbl>'


class _FallbackDocxDocument:
    def __init__(self):
        self._blocks = []

    def add_paragraph(self, text):
        self._blocks.append(('paragraph', text or ''))
        return text

    def add_heading(self, text, level=1):  # noqa: ARG002 - nivel ignorado en el respaldo
        self._blocks.append(('paragraph', text or ''))
        return text

    def add_table(self, rows, cols):
        table = _FallbackDocxTable(rows, cols)
        self._blocks.append(('table', table))
        return table

    def save(self, path):
        document_xml = self._render_document_xml()
        _write_docx_package(path, document_xml)

    def _render_document_xml(self):
        pieces = [
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>',
            '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">',
            '<w:body>',
        ]
        for block_type, payload in self._blocks:
            if block_type == 'paragraph':
                pieces.append(_fallback_paragraph_xml(payload))
            elif block_type == 'table':
                pieces.append(payload.to_xml())
        pieces.extend([
            '<w:sectPr>',
            '<w:pgSz w:w="12240" w:h="15840"/>',
            '<w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440"/>',
            '</w:sectPr>',
            '</w:body>',
            '</w:document>',
        ])
        return '\n'.join(pieces).encode('utf-8')


def _fallback_paragraph_xml(text):
    safe = escape(text or '')
    if not safe:
        return '<w:p/>'
    return f'<w:p><w:r><w:t xml:space="preserve">{safe}</w:t></w:r></w:p>'


def _write_docx_package(path, document_xml_bytes):
    path = Path(path)
    with zipfile.ZipFile(path, 'w') as bundle:
        for name, payload in _DOCX_STATIC_PARTS.items():
            bundle.writestr(name, payload)
        bundle.writestr('word/document.xml', document_xml_bytes)


class FraudCaseApp:
    AUTOSAVE_DELAY_MS = 4000
    SUMMARY_REFRESH_DELAY_MS = 250
    LOG_FLUSH_INTERVAL_MS = 5000
    _external_drive_path: Optional[Path] = None
    _external_log_file_initialized: bool = False

    """Clase que encapsula la aplicación de gestión de casos de fraude."""

    def __init__(self, root):
        self.root = root
        self.root.title("Gestión de Casos de Fraude (App de escritorio)")
        self._suppress_messagebox = False
        # Lista para logs de navegación y validación
        self.logs = []
        self._hover_tooltips = []
        self.validators = []
        self._last_validated_risk_exposure_total = Decimal('0')
        self._log_flush_job_id: Optional[str] = None
        local_log_path = LOGS_FILE if STORE_LOGS_LOCALLY else None
        self._log_file_initialized = bool(local_log_path and os.path.exists(local_log_path))
        self._external_drive_path = self._prepare_external_drive()
        self._external_log_file_initialized = bool(
            EXTERNAL_LOGS_FILE and os.path.exists(EXTERNAL_LOGS_FILE)
        )
        self._export_base_path: Optional[Path] = None
        self.catalog_status_var = tk.StringVar(
            value="Catálogos de detalle pendientes. Usa 'Cargar catálogos' para habilitar el autopoblado."
        )
        self._catalog_loading = False
        self._catalog_ready = False
        self._catalog_prompted = False
        self._catalog_progress_visible = False
        self._catalog_loading_thread = None
        self._catalog_dependent_widgets = []
        self.catalog_load_button = None
        self.catalog_skip_button = None
        self.catalog_progress = None
        self.detail_catalogs = {}
        self.detail_lookup_by_id = {}
        self.team_lookup = {}
        self.client_lookup = {}
        self.product_lookup = {}
        self.claim_lookup = {}
        self.risk_lookup = {}
        self.norm_lookup = {}
        self.import_status_var = tk.StringVar(value="Listo para importar datos masivos.")
        self.import_progress = None
        self._import_progress_visible = False
        self._active_import_jobs = 0

        def register_tooltip(widget, text):
            if widget is None or not text:
                return None
            tip = HoverTooltip(widget, text)
            self._hover_tooltips.append(tip)
            return tip

        self.register_tooltip = register_tooltip
        self.root.protocol("WM_DELETE_WINDOW", self._handle_window_close)
        self._schedule_log_flush()
        # Datos en memoria: listas de frames
        self.client_frames = []
        self.team_frames = []
        self.product_frames = []
        self.risk_frames = []
        self.norm_frames = []
        self._client_frames_by_id = {}
        self._team_frames_by_id = {}
        self._product_frames_by_id = {}
        self.next_risk_number = 1
        self.summary_tables = {}
        self.summary_config = {}
        self.summary_context_menus = {}
        self.summary_tab = None
        self._summary_refresh_after_id = None
        self._summary_dirty_sections = set()
        self._summary_pending_dataset = None
        self._autosave_job_id = None
        self._autosave_dirty = False

        # Variables de caso
        self.id_caso_var = tk.StringVar()
        self.tipo_informe_var = tk.StringVar(value=TIPO_INFORME_LIST[0])
        self.cat_caso1_var = tk.StringVar(value=list(TAXONOMIA.keys())[0])
        subcats = list(TAXONOMIA[self.cat_caso1_var.get()].keys())
        self.cat_caso2_var = tk.StringVar(value=subcats[0])
        mods = TAXONOMIA[self.cat_caso1_var.get()][self.cat_caso2_var.get()]
        self.mod_caso_var = tk.StringVar(value=mods[0])
        self.canal_caso_var = tk.StringVar(value=CANAL_LIST[0])
        self.proceso_caso_var = tk.StringVar(value=PROCESO_LIST[0])

        # Referencias a cuadros de texto de análisis
        self.antecedentes_text = None
        self.modus_text = None
        self.hallazgos_text = None
        self.descargos_text = None
        self.conclusiones_text = None
        self.recomendaciones_text = None

        # Construir interfaz
        self.build_ui()
        # Cargar autosave si existe
        self.load_autosave()
        self.root.after(250, self._prompt_initial_catalog_loading)

    def _prepare_external_drive(self) -> Optional[Path]:
        try:
            return ensure_external_drive_dir()
        except OSError as exc:
            log_event("validacion", f"No se pudo preparar la carpeta externa: {exc}", self.logs)
            return None

    def _get_external_drive_path(self) -> Optional[Path]:
        path = getattr(self, "_external_drive_path", None)
        if path:
            return Path(path)
        self._external_drive_path = self._prepare_external_drive()
        if self._external_drive_path:
            return Path(self._external_drive_path)
        return None

    def _resolve_external_log_target(self) -> Optional[str]:
        if not EXTERNAL_LOGS_FILE:
            return None
        return EXTERNAL_LOGS_FILE if self._get_external_drive_path() else None

    def _has_log_targets(self) -> bool:
        if STORE_LOGS_LOCALLY and LOGS_FILE:
            return True
        return self._resolve_external_log_target() is not None

    def _get_text_content(self, widget: tk.Text) -> str:
        if widget is None:
            return ""
        return widget.get("1.0", "end-1c").strip()

    def _set_text_content(self, widget: tk.Text, value: str) -> None:
        if widget is None:
            return
        widget.delete("1.0", "end")
        if value:
            widget.insert("1.0", value)

    def _analysis_text_widgets(self):
        return {
            "antecedentes": getattr(self, "antecedentes_text", None),
            "modus_operandi": getattr(self, "modus_text", None),
            "hallazgos": getattr(self, "hallazgos_text", None),
            "descargos": getattr(self, "descargos_text", None),
            "conclusiones": getattr(self, "conclusiones_text", None),
            "recomendaciones": getattr(self, "recomendaciones_text", None),
        }

    def _get_exports_folder(self) -> Optional[Path]:
        base_path = getattr(self, "_export_base_path", None) or EXPORTS_DIR
        target = Path(base_path)
        try:
            target.mkdir(parents=True, exist_ok=True)
        except OSError as exc:
            message = f"No se pudo preparar la carpeta de exportación {target}: {exc}"
            log_event("validacion", message, self.logs)
            if not getattr(self, '_suppress_messagebox', False):
                messagebox.showerror("Error al guardar", message)
            return None
        return target

    def import_combined(self, filename=None):
        """Importa datos combinados de productos, clientes y colaboradores."""

        filename = filename or self._select_csv_file("combinado", "Seleccionar CSV combinado")
        if not filename:
            messagebox.showwarning("Sin archivo", "No hay CSV combinado disponible para importar.")
            return
        def worker():
            prepared_rows = []
            for row in iter_massive_csv_rows(filename):
                raw_row = dict(row)
                client_row, client_found = self._hydrate_row_from_details(raw_row, 'id_cliente', CLIENT_ID_ALIASES)
                team_row, team_found = self._hydrate_row_from_details(raw_row, 'id_colaborador', TEAM_ID_ALIASES)
                product_row, product_found = self._hydrate_row_from_details(raw_row, 'id_producto', PRODUCT_ID_ALIASES)
                collaborator_id = (team_row.get('id_colaborador') or '').strip()
                involvement_pairs = parse_involvement_entries(raw_row.get('involucramiento', ''))
                if not involvement_pairs and collaborator_id and raw_row.get('monto_asignado'):
                    involvement_pairs = [(collaborator_id, (raw_row.get('monto_asignado') or '').strip())]
                prepared_rows.append(
                    {
                        'raw_row': raw_row,
                        'client_row': client_row,
                        'client_found': client_found,
                        'team_row': team_row,
                        'team_found': team_found,
                        'product_row': product_row,
                        'product_found': product_found,
                        'involvement_pairs': involvement_pairs,
                    }
                )
            return prepared_rows

        self._start_background_import(
            "datos combinados",
            getattr(self, 'import_combined_button', None),
            worker,
            self._apply_combined_import_payload,
            "No se pudo importar el CSV combinado",
        )

    def import_risks(self, filename=None):
        """Importa riesgos desde un archivo CSV."""

        filename = filename or self._select_csv_file("riesgos", "Seleccionar CSV de riesgos")
        if not filename:
            messagebox.showwarning("Sin archivo", "No se encontró CSV de riesgos para importar.")
            return
        def worker():
            payload = []
            for row in iter_massive_csv_rows(filename):
                hydrated, _ = self._hydrate_row_from_details(row, 'id_riesgo', RISK_ID_ALIASES)
                payload.append(hydrated)
            return payload

        self._start_background_import(
            "riesgos",
            getattr(self, 'import_risks_button', None),
            worker,
            self._apply_risk_import_payload,
            "No se pudo importar riesgos",
        )

    def import_norms(self, filename=None):
        """Importa normas transgredidas desde un archivo CSV."""

        filename = filename or self._select_csv_file("normas", "Seleccionar CSV de normas")
        if not filename:
            messagebox.showwarning("Sin archivo", "No se encontró CSV de normas.")
            return
        def worker():
            payload = []
            for row in iter_massive_csv_rows(filename):
                hydrated, _ = self._hydrate_row_from_details(row, 'id_norma', NORM_ID_ALIASES)
                payload.append(hydrated)
            return payload

        self._start_background_import(
            "normas",
            getattr(self, 'import_norms_button', None),
            worker,
            self._apply_norm_import_payload,
            "No se pudo importar normas",
        )

    def import_claims(self, filename=None):
        """Importa reclamos desde un archivo CSV."""

        filename = filename or self._select_csv_file("reclamos", "Seleccionar CSV de reclamos")
        if not filename:
            messagebox.showwarning("Sin archivo", "No se encontró CSV de reclamos.")
            return
        def worker():
            payload = []
            for row in iter_massive_csv_rows(filename):
                hydrated, found = self._hydrate_row_from_details(row, 'id_producto', PRODUCT_ID_ALIASES)
                payload.append({'row': hydrated, 'found': found})
            return payload

        self._start_background_import(
            "reclamos",
            getattr(self, 'import_claims_button', None),
            worker,
            self._apply_claim_import_payload,
            "No se pudo importar reclamos",
        )

    # ---------------------------------------------------------------------
    # Construcción de la interfaz

    def build_ui(self):
        """Construye la interfaz del usuario en diferentes pestañas."""
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill="both", expand=True)
        self.notebook.bind("<<NotebookTabChanged>>", self._handle_notebook_tab_change)

        # --- Pestaña principal: caso y participantes ---
        self.main_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.main_tab, text="Caso y participantes")
        self.build_case_and_participants_tab(self.main_tab)

        # --- Pestaña Riesgos ---
        risk_tab = ttk.Frame(self.notebook)
        self.notebook.add(risk_tab, text="Riesgos")
        self.build_risk_tab(risk_tab)

        # --- Pestaña Normas ---
        norm_tab = ttk.Frame(self.notebook)
        self.notebook.add(norm_tab, text="Normas")
        self.build_norm_tab(norm_tab)

        # --- Pestaña Análisis ---
        analysis_tab = ttk.Frame(self.notebook)
        self.notebook.add(analysis_tab, text="Análisis y narrativas")
        self.build_analysis_tab(analysis_tab)

        # --- Pestaña Acciones ---
        actions_tab = ttk.Frame(self.notebook)
        self.notebook.add(actions_tab, text="Acciones")
        self.build_actions_tab(actions_tab)

        # --- Pestaña Resumen ---
        summary_tab = ttk.Frame(self.notebook)
        self.notebook.add(summary_tab, text="Resumen")
        self.build_summary_tab(summary_tab)

    def clipboard_get(self):
        """Proxy para ``Tk.clipboard_get`` que simplifica las pruebas unitarias."""

        return self.root.clipboard_get()

    def build_case_and_participants_tab(self, parent):
        """Agrupa en una sola vista los datos del caso, clientes, productos y equipo.

        Esta función encapsula la experiencia solicitada por los analistas:
        evita tener que cambiar de pestaña para capturar la información básica
        del expediente y coloca, en orden lógico, las secciones de caso,
        clientes, productos y colaboradores.  Para lograrlo, se crean
        ``LabelFrame`` consecutivos que actúan como contenedores y se reutiliza
        la lógica existente de ``build_case_tab``/``build_clients_tab``/etc.

        Args:
            parent (tk.Widget): Contenedor donde se inyectarán las secciones.

        Ejemplo::

            # Durante la construcción de la UI
            self.build_case_and_participants_tab(self.main_tab)
        """

        scroll_container = ttk.Frame(parent)
        scroll_container.pack(fill="both", expand=True)
        canvas = tk.Canvas(scroll_container, highlightthickness=0)
        scrollbar = ttk.Scrollbar(scroll_container, orient="vertical", command=canvas.yview)
        canvas.configure(yscrollcommand=scrollbar.set)
        scrollbar.pack(side="right", fill="y")
        canvas.pack(side="left", fill="both", expand=True)

        self._main_scrollable_frame = ttk.Frame(canvas)
        window_id = canvas.create_window((0, 0), window=self._main_scrollable_frame, anchor="nw")

        def _configure_canvas(event):
            canvas.configure(scrollregion=canvas.bbox("all"))

        def _resize_inner(event):
            canvas.itemconfigure(window_id, width=event.width)

        self._main_scrollable_frame.bind("<Configure>", _configure_canvas)
        canvas.bind("<Configure>", _resize_inner)

        case_section = ttk.LabelFrame(self._main_scrollable_frame, text="1. Datos generales del caso")
        case_section.pack(fill="x", expand=False, padx=5, pady=5)
        self.build_case_tab(case_section)

        clients_section = ttk.LabelFrame(self._main_scrollable_frame, text="2. Clientes implicados")
        clients_section.pack(fill="x", expand=True, padx=5, pady=5)
        self.build_clients_tab(clients_section)

        products_section = ttk.LabelFrame(self._main_scrollable_frame, text="3. Productos investigados")
        products_section.pack(fill="x", expand=True, padx=5, pady=5)
        self.build_products_tab(products_section)

        team_section = ttk.LabelFrame(self._main_scrollable_frame, text="4. Colaboradores involucrados")
        team_section.pack(fill="x", expand=True, padx=5, pady=5)
        self.build_team_tab(team_section)

    def _safe_update_idletasks(self):
        """Intenta refrescar la UI sin propagar errores cuando la ventana no existe."""

        try:
            self.root.update_idletasks()
        except tk.TclError:
            pass

    def _notify_taxonomy_warning(self, message):
        """Centraliza el aviso de inconsistencias en la taxonomía."""

        log_event('validacion', message, self.logs)
        try:
            messagebox.showwarning('Taxonomía inválida', message)
        except tk.TclError:
            pass

    def focus_main_tab(self):
        """Muestra la pestaña principal cuando una importación agrega registros.

        Al terminar de cargar clientes, productos o colaboradores desde la
        pestaña de acciones, los usuarios esperaban ver de inmediato los datos
        recién agregados.  Este helper selecciona la pestaña principal del
        ``Notebook`` (si existe) y registra el cambio en el log para que el
        auto-guardado capture el nuevo estado.
        """

        if getattr(self, "notebook", None) is None or getattr(self, "main_tab", None) is None:
            return
        try:
            self.notebook.select(self.main_tab)
            log_event("navegacion", "Se mostró la pestaña principal tras importar datos", self.logs)
        except tk.TclError:
            # Si el widget ya no existe (por ejemplo, al cerrar la app), se ignora el error.
            pass

    def sync_main_form_after_import(self, section_name, stay_on_summary=False):
        """Sincroniza la interfaz principal después de importar datos masivos.

        La función se diseñó como respuesta directa al feedback de que los
        registros cargados desde la pestaña de *Acciones* no se veían reflejados
        inmediatamente en las secciones de caso, clientes y productos.  Para
        garantizar esa visibilidad, se realizan cuatro pasos en orden:

            1. Se invocan ``update_client_options_global`` y
               ``update_team_options_global`` para notificar a todos los
               ``Combobox`` dependientes que hay nuevos IDs disponibles.
            2. Se llama a ``update_idletasks`` sobre la ventana raíz para que
               Tkinter repinte los marcos dinámicos y muestre los datos recién
               insertados sin esperar a que el usuario interactúe.
            3. Se selecciona la pestaña principal (método ``focus_main_tab``) de
               modo que el usuario visualice de inmediato los clientes,
               productos o colaboradores importados.
            4. Se registra un evento de navegación indicando qué tipo de datos
               disparó la sincronización, lo que permite auditar el proceso.

        Args:
            section_name (str): Nombre en español de la sección importada. Se
                utiliza únicamente para detallar el mensaje del log.
            stay_on_summary (bool): Cuando es ``True`` evita llamar a
                :meth:`focus_main_tab` para que la vista permanezca en la
                pestaña de Resumen tras pegar datos directamente en las tablas
                auxiliares.

        Ejemplo::

            self.sync_main_form_after_import("clientes")
        """

        self.update_client_options_global()
        self.update_team_options_global()
        try:
            self.root.update_idletasks()
        except tk.TclError:
            # Si la ventana ya no existe (por ejemplo al cerrar la app) no es
            # necesario continuar con la sincronización.
            return
        if not stay_on_summary:
            self.focus_main_tab()
        log_event(
            "navegacion",
            f"Sincronizó la pestaña principal tras importar {section_name}",
            self.logs,
        )

    def _notify_dataset_changed(self, summary_sections=None):
        """Marca el formulario como modificado y agenda el autosave diferido."""

        self.request_autosave()
        self._schedule_summary_refresh(sections=summary_sections)

    def build_case_tab(self, parent):
        """Construye la pestaña de detalles del caso."""
        frame = ttk.Frame(parent)
        frame.pack(fill="both", expand=True, padx=5, pady=5)
        # Case ID y tipo de informe
        row1 = ttk.Frame(frame)
        row1.pack(fill="x", pady=2)
        ttk.Label(row1, text="Número de caso (AAAA-NNNN):").pack(side="left")
        id_entry = ttk.Entry(row1, textvariable=self.id_caso_var, width=20)
        id_entry.pack(side="left", padx=5)
        id_entry.bind("<FocusOut>", lambda e: self._log_navigation_change("Modificó número de caso"))
        self.register_tooltip(id_entry, "Identificador del expediente con formato AAAA-NNNN.")
        ttk.Label(row1, text="Tipo de informe:").pack(side="left")
        tipo_cb = ttk.Combobox(row1, textvariable=self.tipo_informe_var, values=TIPO_INFORME_LIST, state="readonly", width=15)
        tipo_cb.pack(side="left", padx=5)
        self.register_tooltip(tipo_cb, "Selecciona si el reporte es interno o regulatorio.")

        # Categorías y modalidad
        row2 = ttk.Frame(frame)
        row2.pack(fill="x", pady=2)
        ttk.Label(row2, text="Categoría nivel 1:").pack(side="left")
        cat1_cb = ttk.Combobox(row2, textvariable=self.cat_caso1_var, values=list(TAXONOMIA.keys()), state="readonly", width=20)
        cat1_cb.pack(side="left", padx=5)
        cat1_cb.bind("<FocusOut>", lambda e: self.on_case_cat1_change())
        cat1_cb.bind("<<ComboboxSelected>>", lambda e: self.on_case_cat1_change())
        self.register_tooltip(cat1_cb, "Nivel superior de la taxonomía de fraude.")
        ttk.Label(row2, text="Categoría nivel 2:").pack(side="left")
        self.case_cat2_cb = ttk.Combobox(row2, textvariable=self.cat_caso2_var, values=list(TAXONOMIA[self.cat_caso1_var.get()].keys()), state="readonly", width=20)
        self.case_cat2_cb.pack(side="left", padx=5)
        self.case_cat2_cb.bind("<FocusOut>", lambda e: self.on_case_cat2_change())
        self.case_cat2_cb.bind("<<ComboboxSelected>>", lambda e: self.on_case_cat2_change())
        self.register_tooltip(self.case_cat2_cb, "Subcategoría que precisa el evento.")
        ttk.Label(row2, text="Modalidad:").pack(side="left")
        self.case_mod_cb = ttk.Combobox(row2, textvariable=self.mod_caso_var, values=TAXONOMIA[self.cat_caso1_var.get()][self.cat_caso2_var.get()], state="readonly", width=25)
        self.case_mod_cb.pack(side="left", padx=5)
        self.register_tooltip(self.case_mod_cb, "Modalidad específica dentro de la taxonomía.")

        # Canal y proceso
        row3 = ttk.Frame(frame)
        row3.pack(fill="x", pady=2)
        ttk.Label(row3, text="Canal:").pack(side="left")
        canal_cb = ttk.Combobox(row3, textvariable=self.canal_caso_var, values=CANAL_LIST, state="readonly", width=25)
        canal_cb.pack(side="left", padx=5)
        self.register_tooltip(canal_cb, "Canal donde se originó el evento.")
        ttk.Label(row3, text="Proceso impactado:").pack(side="left")
        proc_cb = ttk.Combobox(row3, textvariable=self.proceso_caso_var, values=PROCESO_LIST, state="readonly", width=25)
        proc_cb.pack(side="left", padx=5)
        self.register_tooltip(proc_cb, "Proceso que sufrió la desviación.")

        # Validaciones del caso
        self.validators.append(
            FieldValidator(
                id_entry,
                lambda: validate_case_id(self.id_caso_var.get()),
                self.logs,
                "Caso - ID",
                variables=[self.id_caso_var],
            )
        )
        self.validators.append(
            FieldValidator(
                tipo_cb,
                lambda: validate_required_text(self.tipo_informe_var.get(), "el tipo de informe"),
                self.logs,
                "Caso - Tipo de informe",
                variables=[self.tipo_informe_var],
            )
        )
        self.validators.append(
            FieldValidator(
                cat1_cb,
                lambda: validate_required_text(self.cat_caso1_var.get(), "la categoría nivel 1"),
                self.logs,
                "Caso - Categoría 1",
                variables=[self.cat_caso1_var],
            )
        )
        self.validators.append(
            FieldValidator(
                self.case_cat2_cb,
                lambda: validate_required_text(self.cat_caso2_var.get(), "la categoría nivel 2"),
                self.logs,
                "Caso - Categoría 2",
                variables=[self.cat_caso2_var],
            )
        )
        self.validators.append(
            FieldValidator(
                self.case_mod_cb,
                lambda: validate_required_text(self.mod_caso_var.get(), "la modalidad del caso"),
                self.logs,
                "Caso - Modalidad",
                variables=[self.mod_caso_var],
            )
        )
        self.validators.append(
            FieldValidator(
                canal_cb,
                lambda: validate_required_text(self.canal_caso_var.get(), "el canal del caso"),
                self.logs,
                "Caso - Canal",
                variables=[self.canal_caso_var],
            )
        )
        self.validators.append(
            FieldValidator(
                proc_cb,
                lambda: validate_required_text(self.proceso_caso_var.get(), "el proceso impactado"),
                self.logs,
                "Caso - Proceso",
                variables=[self.proceso_caso_var],
            )
        )

    def on_case_cat1_change(self):
        """Actualiza las opciones de categoría 2 y modalidad cuando cambia cat1 del caso."""
        cat1 = self.cat_caso1_var.get()
        subcats = list(TAXONOMIA.get(cat1, {}).keys())
        if not subcats:
            subcats = [""]
        self.case_cat2_cb['values'] = subcats
        self.cat_caso2_var.set('')
        self.case_cat2_cb.set('')
        self.case_mod_cb['values'] = []
        self.mod_caso_var.set('')
        self.case_mod_cb.set('')
        self._log_navigation_change("Modificó categoría 1 del caso")

    def on_case_cat2_change(self):
        cat1 = self.cat_caso1_var.get()
        cat2 = self.cat_caso2_var.get()
        mods = TAXONOMIA.get(cat1, {}).get(cat2, [])
        if not mods:
            mods = [""]
        self.case_mod_cb['values'] = mods
        self.mod_caso_var.set('')
        self.case_mod_cb.set('')
        self._log_navigation_change("Modificó categoría 2 del caso")
        if self.cat_caso2_var.get() == 'Fraude Externo':
            messagebox.showwarning(
                "Analítica de fraude externo",
                "Recuerda coordinar con el equipo de reclamos para registrar la analítica correcta en casos de Fraude Externo.",
            )

    def build_clients_tab(self, parent):
        """Construye la pestaña de clientes con lista dinámica."""
        frame = ttk.Frame(parent)
        frame.pack(fill="both", expand=True)
        # Contenedor para los clientes
        self.clients_container = ttk.Frame(frame)
        self.clients_container.pack(fill="x", pady=5)
        # Botón añadir cliente
        add_btn = ttk.Button(frame, text="Agregar cliente", command=self.add_client)
        add_btn.pack(side="left", padx=5)
        self.register_tooltip(add_btn, "Añade un nuevo cliente implicado en el caso.")
        # Inicialmente un cliente en blanco
        self.add_client()

    def add_client(self):
        """Crea y añade un nuevo marco de cliente a la interfaz.

        Se utiliza ``self.client_lookup`` para proporcionar datos de autopoblado
        al nuevo cliente, en caso de que exista un registro previo en
        ``client_details.csv``. Luego se actualizan las opciones de clientes
        disponibles para los productos.
        """
        idx = len(self.client_frames)
        client = ClientFrame(
            self.clients_container,
            idx,
            self.remove_client,
            self.update_client_options_global,
            self.logs,
            self.register_tooltip,
            client_lookup=self.client_lookup,
            summary_refresh_callback=self._schedule_summary_refresh,
            change_notifier=self._log_navigation_change,
            id_change_callback=self._handle_client_id_change,
        )
        self.client_frames.append(client)
        self.update_client_options_global()
        self._schedule_summary_refresh('clientes')

    def remove_client(self, client_frame):
        self._handle_client_id_change(client_frame, client_frame.id_var.get(), None)
        self.client_frames.remove(client_frame)
        # Renombrar las etiquetas
        for i, cl in enumerate(self.client_frames):
            cl.idx = i
            cl.frame.config(text=f"Cliente {i+1}")
        self.update_client_options_global()
        self._schedule_summary_refresh('clientes')

    def update_client_options_global(self):
        """Actualiza la lista de clientes en todos los productos y envolvimientos."""
        options = [c.id_var.get().strip() for c in self.client_frames if c.id_var.get().strip()]
        # Actualizar combobox de clientes en cada producto
        for prod in self.product_frames:
            prod.update_client_options()
        log_event("navegacion", "Actualizó opciones de cliente", self.logs)

    def build_team_tab(self, parent):
        frame = ttk.Frame(parent)
        frame.pack(fill="both", expand=True)
        self.team_container = ttk.Frame(frame)
        self.team_container.pack(fill="x", pady=5)
        add_btn = ttk.Button(frame, text="Agregar colaborador", command=self.add_team)
        add_btn.pack(side="left", padx=5)
        self.register_tooltip(add_btn, "Crea un registro para otro colaborador investigado.")
        self.add_team()

    def add_team(self):
        idx = len(self.team_frames)
        team = TeamMemberFrame(
            self.team_container,
            idx,
            self.remove_team,
            self.update_team_options_global,
            self.team_lookup,
            self.logs,
            self.register_tooltip,
            summary_refresh_callback=self._schedule_summary_refresh,
            change_notifier=self._log_navigation_change,
            id_change_callback=self._handle_team_id_change,
        )
        self.team_frames.append(team)
        self.update_team_options_global()
        self._schedule_summary_refresh('colaboradores')

    def remove_team(self, team_frame):
        self._handle_team_id_change(team_frame, team_frame.id_var.get(), None)
        self.team_frames.remove(team_frame)
        # Renombrar
        for i, tm in enumerate(self.team_frames):
            tm.idx = i
            tm.frame.config(text=f"Colaborador {i+1}")
        self.update_team_options_global()
        self._schedule_summary_refresh('colaboradores')

    def update_team_options_global(self):
        """Actualiza listas de colaboradores en productos e involucra."""
        options = [t.id_var.get().strip() for t in self.team_frames if t.id_var.get().strip()]
        for prod in self.product_frames:
            prod.update_team_options()
        log_event("navegacion", "Actualizó opciones de colaborador", self.logs)

    def build_products_tab(self, parent):
        frame = ttk.Frame(parent)
        frame.pack(fill="both", expand=True)
        self.product_container = ttk.Frame(frame)
        self.product_container.pack(fill="x", pady=5)
        add_btn = ttk.Button(frame, text="Agregar producto", command=self.add_product)
        add_btn.pack(side="left", padx=5)
        self.register_tooltip(add_btn, "Registra un nuevo producto investigado.")
        # No añadimos automáticamente un producto porque los productos están asociados a clientes

    def _apply_case_taxonomy_defaults(self, product_frame):
        """Configura un producto nuevo con la taxonomía seleccionada en el caso."""

        cat1 = self.cat_caso1_var.get().strip()
        cat2 = self.cat_caso2_var.get().strip()
        modalidad = self.mod_caso_var.get().strip()

        if not cat1 or cat1 not in TAXONOMIA:
            return

        previous_suppression = getattr(product_frame, '_suppress_change_notifications', False)
        product_frame._suppress_change_notifications = True
        try:
            product_frame.cat1_var.set(cat1)
            product_frame.on_cat1_change()
            if cat2 and cat2 in TAXONOMIA[cat1]:
                product_frame.cat2_var.set(cat2)
                if hasattr(product_frame, 'cat2_cb'):
                    product_frame.cat2_cb.set(cat2)
                product_frame.on_cat2_change()
                if modalidad and modalidad in TAXONOMIA[cat1][cat2]:
                    product_frame.mod_var.set(modalidad)
                    if hasattr(product_frame, 'mod_cb'):
                        product_frame.mod_cb.set(modalidad)
        finally:
            product_frame._suppress_change_notifications = previous_suppression

    def add_product(self, initialize_rows=True):
        idx = len(self.product_frames)
        prod = ProductFrame(
            self.product_container,
            idx,
            self.remove_product,
            self.get_client_ids,
            self.get_team_ids,
            self.logs,
            self.product_lookup,
            self.register_tooltip,
            claim_lookup=self.claim_lookup,
            summary_refresh_callback=self._schedule_summary_refresh,
            change_notifier=self._log_navigation_change,
            id_change_callback=self._handle_product_id_change,
            initialize_rows=initialize_rows,
        )
        self._apply_case_taxonomy_defaults(prod)
        self.product_frames.append(prod)
        # Renombrar
        for i, p in enumerate(self.product_frames):
            p.idx = i
            p.frame.config(text=f"Producto {i+1}")
        self._schedule_summary_refresh({'productos', 'reclamos'})

    def remove_product(self, prod_frame):
        self._handle_product_id_change(prod_frame, prod_frame.id_var.get(), None)
        self.product_frames.remove(prod_frame)
        for i, p in enumerate(self.product_frames):
            p.idx = i
            p.frame.config(text=f"Producto {i+1}")
        self._schedule_summary_refresh({'productos', 'reclamos'})

    def get_client_ids(self):
        return [c.id_var.get().strip() for c in self.client_frames if c.id_var.get().strip()]

    def get_team_ids(self):
        return [t.id_var.get().strip() for t in self.team_frames if t.id_var.get().strip()]

    def build_risk_tab(self, parent):
        frame = ttk.Frame(parent)
        frame.pack(fill="both", expand=True)
        self.risk_container = ttk.Frame(frame)
        self.risk_container.pack(fill="x", pady=5)
        add_btn = ttk.Button(frame, text="Agregar riesgo", command=self.add_risk)
        add_btn.pack(side="left", padx=5)
        self.register_tooltip(add_btn, "Registra un nuevo riesgo identificado.")
        self.add_risk()

    def add_risk(self):
        idx = len(self.risk_frames)
        default_risk_id = self._generate_next_risk_id()
        risk = RiskFrame(
            self.risk_container,
            idx,
            self.remove_risk,
            self.logs,
            self.register_tooltip,
            change_notifier=self._log_navigation_change,
            default_risk_id=default_risk_id,
        )
        self.risk_frames.append(risk)
        for i, r in enumerate(self.risk_frames):
            r.idx = i
            r.frame.config(text=f"Riesgo {i+1}")
        self._refresh_risk_auto_ids()
        self._schedule_summary_refresh('riesgos')

    def remove_risk(self, risk_frame):
        self.risk_frames.remove(risk_frame)
        for i, r in enumerate(self.risk_frames):
            r.idx = i
            r.frame.config(text=f"Riesgo {i+1}")
        self._refresh_risk_auto_ids()
        self._schedule_summary_refresh('riesgos')

    def _generate_next_risk_id(self, used_ids=None):
        """Obtiene el siguiente ID automático disponible para riesgos."""

        normalized_used = {rid for rid in (used_ids or set()) if rid}
        if used_ids is None:
            normalized_used.update(r.id_var.get().strip() for r in self.risk_frames if r.id_var.get().strip())
        while True:
            candidate = f"RSK-{self.next_risk_number:06d}"
            self.next_risk_number += 1
            if candidate not in normalized_used:
                return candidate
            # En caso de colisión continuar buscando
        # Not reachable

    def _refresh_risk_auto_ids(self):
        """Normaliza los IDs automáticos sin afectar los editados manualmente."""

        used_ids = set()
        frames_to_update = []
        for frame in self.risk_frames:
            rid = frame.id_var.get().strip()
            if rid and rid not in used_ids:
                used_ids.add(rid)
                continue
            frames_to_update.append(frame)
        for frame in frames_to_update:
            if frame.has_user_modified_id():
                continue
            new_id = self._generate_next_risk_id(used_ids)
            frame.assign_new_auto_id(new_id)
            used_ids.add(new_id)

    def build_norm_tab(self, parent):
        frame = ttk.Frame(parent)
        frame.pack(fill="both", expand=True)
        self.norm_container = ttk.Frame(frame)
        self.norm_container.pack(fill="x", pady=5)
        add_btn = ttk.Button(frame, text="Agregar norma", command=self.add_norm)
        add_btn.pack(side="left", padx=5)
        self.register_tooltip(add_btn, "Agrega otra norma transgredida.")

    def add_norm(self):
        idx = len(self.norm_frames)
        norm = NormFrame(
            self.norm_container,
            idx,
            self.remove_norm,
            self.logs,
            self.register_tooltip,
            change_notifier=self._log_navigation_change,
        )
        self.norm_frames.append(norm)
        for i, n in enumerate(self.norm_frames):
            n.idx = i
            n.frame.config(text=f"Norma {i+1}")
        self._schedule_summary_refresh('normas')

    def remove_norm(self, norm_frame):
        self.norm_frames.remove(norm_frame)
        for i, n in enumerate(self.norm_frames):
            n.idx = i
            n.frame.config(text=f"Norma {i+1}")
        self._schedule_summary_refresh('normas')

    def build_analysis_tab(self, parent):
        frame = ttk.Frame(parent)
        frame.pack(fill="both", expand=True)
        # Campos de análisis
        row1 = ttk.Frame(frame)
        row1.pack(fill="x", pady=1)
        ttk.Label(row1, text="Antecedentes:").pack(side="left")
        self.antecedentes_text = scrolledtext.ScrolledText(row1, width=80, height=4, wrap="word")
        self.antecedentes_text.pack(side="left", padx=5)
        self.antecedentes_text.bind("<FocusOut>", lambda e: self._log_navigation_change("Modificó antecedentes"))
        self.register_tooltip(self.antecedentes_text, "Resume los hechos previos y contexto del caso.")
        row2 = ttk.Frame(frame)
        row2.pack(fill="x", pady=1)
        ttk.Label(row2, text="Modus operandi:").pack(side="left")
        self.modus_text = scrolledtext.ScrolledText(row2, width=80, height=4, wrap="word")
        self.modus_text.pack(side="left", padx=5)
        self.modus_text.bind("<FocusOut>", lambda e: self._log_navigation_change("Modificó modus operandi"))
        self.register_tooltip(self.modus_text, "Describe la forma en que se ejecutó el fraude.")
        row3 = ttk.Frame(frame)
        row3.pack(fill="x", pady=1)
        ttk.Label(row3, text="Hallazgos principales:").pack(side="left")
        self.hallazgos_text = scrolledtext.ScrolledText(row3, width=80, height=4, wrap="word")
        self.hallazgos_text.pack(side="left", padx=5)
        self.hallazgos_text.bind("<FocusOut>", lambda e: self._log_navigation_change("Modificó hallazgos"))
        self.register_tooltip(self.hallazgos_text, "Menciona los hallazgos clave de la investigación.")
        row4 = ttk.Frame(frame)
        row4.pack(fill="x", pady=1)
        ttk.Label(row4, text="Descargos del colaborador:").pack(side="left")
        self.descargos_text = scrolledtext.ScrolledText(row4, width=80, height=4, wrap="word")
        self.descargos_text.pack(side="left", padx=5)
        self.descargos_text.bind("<FocusOut>", lambda e: self._log_navigation_change("Modificó descargos"))
        self.register_tooltip(self.descargos_text, "Registra los descargos formales del colaborador.")
        row5 = ttk.Frame(frame)
        row5.pack(fill="x", pady=1)
        ttk.Label(row5, text="Conclusiones:").pack(side="left")
        self.conclusiones_text = scrolledtext.ScrolledText(row5, width=80, height=4, wrap="word")
        self.conclusiones_text.pack(side="left", padx=5)
        self.conclusiones_text.bind("<FocusOut>", lambda e: self._log_navigation_change("Modificó conclusiones"))
        self.register_tooltip(self.conclusiones_text, "Escribe las conclusiones generales del informe.")
        row6 = ttk.Frame(frame)
        row6.pack(fill="x", pady=1)
        ttk.Label(row6, text="Recomendaciones y mejoras:").pack(side="left")
        self.recomendaciones_text = scrolledtext.ScrolledText(row6, width=80, height=4, wrap="word")
        self.recomendaciones_text.pack(side="left", padx=5)
        self.recomendaciones_text.bind("<FocusOut>", lambda e: self._log_navigation_change("Modificó recomendaciones"))
        self.register_tooltip(self.recomendaciones_text, "Propón acciones correctivas y preventivas.")

    def build_actions_tab(self, parent):
        frame = ttk.Frame(parent)
        frame.pack(fill="both", expand=True, padx=5, pady=5)
        catalog_group = ttk.LabelFrame(frame, text="Catálogos de detalle")
        catalog_group.pack(fill="x", expand=False, pady=(0, 10))
        ttk.Label(
            catalog_group,
            textvariable=self.catalog_status_var,
            wraplength=500,
            justify="left",
        ).pack(anchor="w", padx=5, pady=(5, 2))
        catalog_controls = ttk.Frame(catalog_group)
        catalog_controls.pack(fill="x", padx=5, pady=(0, 5))
        self.catalog_load_button = ttk.Button(
            catalog_controls,
            text="Cargar catálogos",
            command=self.request_catalog_loading,
        )
        self.catalog_load_button.pack(side="left", padx=2)
        self.catalog_skip_button = ttk.Button(
            catalog_controls,
            text="Iniciar sin catálogos",
            command=self._mark_catalogs_skipped,
        )
        self.catalog_skip_button.pack(side="left", padx=2)
        self.catalog_progress = ttk.Progressbar(catalog_controls, mode="indeterminate", length=160)
        self._catalog_progress_visible = False
        # Botones de importación
        ttk.Label(frame, text="Importar datos masivos (CSV)").pack(anchor="w")
        import_buttons = ttk.Frame(frame)
        import_buttons.pack(anchor="w", pady=2)
        btn_clientes = ttk.Button(import_buttons, text="Cargar clientes", command=self.import_clients)
        btn_clientes.pack(side="left", padx=2)
        self.register_tooltip(btn_clientes, "Importa clientes desde un CSV masivo.")
        btn_colabs = ttk.Button(import_buttons, text="Cargar colaboradores", command=self.import_team_members)
        btn_colabs.pack(side="left", padx=2)
        self.register_tooltip(btn_colabs, "Importa colaboradores y sus datos laborales.")
        btn_productos = ttk.Button(import_buttons, text="Cargar productos", command=self.import_products)
        btn_productos.pack(side="left", padx=2)
        self.register_tooltip(btn_productos, "Carga productos investigados desde un CSV.")
        btn_combo = ttk.Button(import_buttons, text="Cargar combinado", command=self.import_combined)
        btn_combo.pack(side="left", padx=2)
        self.register_tooltip(btn_combo, "Importa en un solo archivo clientes, productos y colaboradores.")
        # Nuevos botones para riesgos, normas y reclamos
        btn_riesgos = ttk.Button(import_buttons, text="Cargar riesgos", command=self.import_risks)
        btn_riesgos.pack(side="left", padx=2)
        self.register_tooltip(btn_riesgos, "Carga la matriz de riesgos desde CSV.")
        btn_normas = ttk.Button(import_buttons, text="Cargar normas", command=self.import_norms)
        btn_normas.pack(side="left", padx=2)
        self.register_tooltip(btn_normas, "Importa las normas vulneradas.")
        btn_reclamos = ttk.Button(import_buttons, text="Cargar reclamos", command=self.import_claims)
        btn_reclamos.pack(side="left", padx=2)
        self.register_tooltip(btn_reclamos, "Vincula reclamos con los productos.")
        self.import_clients_button = btn_clientes
        self.import_team_button = btn_colabs
        self.import_products_button = btn_productos
        self.import_combined_button = btn_combo
        self.import_risks_button = btn_riesgos
        self.import_norms_button = btn_normas
        self.import_claims_button = btn_reclamos
        for widget in (
            btn_clientes,
            btn_colabs,
            btn_productos,
            btn_combo,
            btn_riesgos,
            btn_normas,
            btn_reclamos,
        ):
            self._register_catalog_dependent_widget(widget)
        ttk.Label(
            frame,
            textvariable=self.import_status_var,
            wraplength=500,
            justify="left",
        ).pack(anchor="w", pady=(2, 0))
        self.import_progress = ttk.Progressbar(frame, mode="indeterminate", length=260)
        self._import_progress_visible = False

        # Botones de guardado y carga
        ttk.Label(frame, text="Guardar y cargar versiones").pack(anchor="w", pady=(10,2))
        version_buttons = ttk.Frame(frame)
        version_buttons.pack(anchor="w", pady=2)
        btn_save = ttk.Button(version_buttons, text="Guardar y enviar", command=self.save_and_send)
        btn_save.pack(side="left", padx=2)
        self.register_tooltip(btn_save, "Valida y exporta todos los archivos requeridos.")
        btn_load = ttk.Button(version_buttons, text="Cargar versión", command=self.load_version_dialog)
        btn_load.pack(side="left", padx=2)
        self.register_tooltip(btn_load, "Restaura una versión previa en formato JSON.")
        btn_clear = ttk.Button(
            version_buttons,
            text="Borrar todos los datos",
            command=lambda: self.clear_all(notify=True),
        )
        btn_clear.pack(side="left", padx=2)
        self.register_tooltip(btn_clear, "Limpia el formulario completo para iniciar desde cero.")

        # Información adicional
        ttk.Label(frame, text="El auto‑guardado se realiza automáticamente en un archivo JSON").pack(anchor="w", pady=(10,2))
        self._set_catalog_dependent_state(self._catalog_loading or self._active_import_jobs > 0)

    def _register_catalog_dependent_widget(self, widget):
        if widget is None:
            return
        self._catalog_dependent_widgets.append(widget)
        if self._catalog_loading:
            try:
                widget.state(['disabled'])
            except tk.TclError:
                pass

    def _set_catalog_dependent_state(self, disabled):
        for widget in self._catalog_dependent_widgets:
            if widget is None:
                continue
            try:
                if disabled:
                    widget.state(['disabled'])
                else:
                    widget.state(['!disabled'])
            except tk.TclError:
                continue

    def _show_catalog_progress(self):
        if not self.catalog_progress:
            return
        if not self._catalog_progress_visible:
            self.catalog_progress.pack(side="left", fill="x", expand=True, padx=5)
            self._catalog_progress_visible = True
        try:
            self.catalog_progress.start(10)
        except tk.TclError:
            pass

    def _hide_catalog_progress(self):
        if not self.catalog_progress:
            return
        try:
            self.catalog_progress.stop()
        except tk.TclError:
            pass
        if self._catalog_progress_visible:
            self.catalog_progress.pack_forget()
            self._catalog_progress_visible = False

    def _prompt_initial_catalog_loading(self):
        if self._catalog_prompted or self._catalog_loading or self._catalog_ready:
            return
        self._catalog_prompted = True
        try:
            should_load = messagebox.askyesno(
                "Catálogos de detalle",
                (
                    "El formulario puede autopoblar clientes, productos y colaboradores "
                    "si cargas los catálogos de detalle. ¿Deseas cargarlos ahora?"
                ),
                parent=self.root,
            )
        except tk.TclError:
            should_load = True
        if should_load:
            self.request_catalog_loading()
        else:
            self._mark_catalogs_skipped()

    def _mark_catalogs_skipped(self):
        if self._catalog_loading:
            return
        self._catalog_prompted = True
        self._catalog_ready = False
        self.catalog_status_var.set(
            "Trabajarás sin catálogos. Puedes cargarlos más adelante desde la pestaña Acciones."
        )
        self._set_catalog_dependent_state(False)

    def request_catalog_loading(self):
        if self._catalog_loading:
            return
        self._catalog_prompted = True
        self._catalog_loading = True
        self._catalog_ready = False
        self.catalog_status_var.set(
            "Cargando catálogos de detalle. Este proceso puede tardar algunos segundos..."
        )
        self._show_catalog_progress()
        self._set_catalog_dependent_state(True)
        for button in (self.catalog_load_button, self.catalog_skip_button):
            if button is None:
                continue
            try:
                button.state(['disabled'])
            except tk.TclError:
                pass
        self._catalog_loading_thread = threading.Thread(
            target=self._load_catalogs_in_background,
            daemon=True,
            name="catalog-loader",
        )
        self._catalog_loading_thread.start()

    def _load_catalogs_in_background(self):
        try:
            raw_catalogs = load_detail_catalogs()
            detail_catalogs, detail_lookup_by_id = self._normalize_detail_catalog_payload(raw_catalogs)
        except Exception as exc:  # pragma: no cover - errores inusuales de IO
            self._dispatch_to_ui(self._handle_catalog_load_failure, exc)
            return
        self._dispatch_to_ui(self._handle_catalog_load_success, detail_catalogs, detail_lookup_by_id)

    def _dispatch_to_ui(self, callback, *args, **kwargs):
        if getattr(self, 'root', None) is None:
            return
        try:
            self.root.after(0, lambda: callback(*args, **kwargs))
        except tk.TclError:
            pass

    def _handle_catalog_load_success(self, detail_catalogs, detail_lookup_by_id):
        self._catalog_ready = True
        self._apply_catalog_lookups(detail_catalogs, detail_lookup_by_id)
        if detail_lookup_by_id:
            self.catalog_status_var.set(
                "Catálogos de detalle cargados. El autopoblado e importación están habilitados."
            )
        else:
            self.catalog_status_var.set(
                "No se encontraron catálogos de detalle. Puedes seguir trabajando manualmente."
            )
        self._finalize_catalog_loading_state()

    def _handle_catalog_load_failure(self, error):
        self._catalog_ready = False
        self.catalog_status_var.set(
            "No se pudieron cargar los catálogos. Intenta nuevamente desde 'Cargar catálogos'."
        )
        self._finalize_catalog_loading_state(failed=True)
        try:
            messagebox.showerror(
                "Catálogos de detalle",
                f"No se pudieron cargar los catálogos: {error}",
            )
        except tk.TclError:
            pass

    def _finalize_catalog_loading_state(self, failed=False):
        self._ensure_import_runtime_state()
        self._catalog_loading = False
        self._hide_catalog_progress()
        self._set_catalog_dependent_state(self._active_import_jobs > 0)
        self._catalog_loading_thread = None
        target_text = "Recargar catálogos" if not failed else "Reintentar carga"
        if self.catalog_load_button is not None:
            try:
                self.catalog_load_button.config(text=target_text)
                self.catalog_load_button.state(['!disabled'])
            except tk.TclError:
                pass
        if self.catalog_skip_button is not None:
            try:
                self.catalog_skip_button.state(['!disabled'])
            except tk.TclError:
                pass

    def _ensure_import_runtime_state(self):
        if not hasattr(self, '_active_import_jobs'):
            self._active_import_jobs = 0
        if not hasattr(self, '_import_progress_visible'):
            self._import_progress_visible = False
        if not hasattr(self, 'import_progress'):
            self.import_progress = None
        if not hasattr(self, 'import_status_var'):
            self.import_status_var = None
        if not hasattr(self, '_catalog_loading'):
            self._catalog_loading = False
        if not hasattr(self, '_catalog_dependent_widgets'):
            self._catalog_dependent_widgets = []

    def _show_import_progress(self):
        self._ensure_import_runtime_state()
        if not self.import_progress:
            return
        if not self._import_progress_visible:
            self.import_progress.pack(anchor="w", padx=5, pady=(0, 5))
            self._import_progress_visible = True
        try:
            self.import_progress.start(10)
        except tk.TclError:
            pass

    def _hide_import_progress(self):
        self._ensure_import_runtime_state()
        if not self.import_progress:
            return
        try:
            self.import_progress.stop()
        except tk.TclError:
            pass
        if self._import_progress_visible:
            self.import_progress.pack_forget()
            self._import_progress_visible = False

    def _on_import_started(self, task_label):
        self._ensure_import_runtime_state()
        self._active_import_jobs += 1
        if self.import_status_var is not None:
            self.import_status_var.set(f"Importando {task_label}...")
        self._show_import_progress()
        if not self._catalog_loading:
            self._set_catalog_dependent_state(True)

    def _finalize_import_task(self, message, failed=False):
        self._ensure_import_runtime_state()
        self._active_import_jobs = max(0, self._active_import_jobs - 1)
        if self.import_status_var is not None and message:
            self.import_status_var.set(message)
        if self._active_import_jobs == 0:
            self._hide_import_progress()
            if not self._catalog_loading:
                self._set_catalog_dependent_state(False)

    def _start_background_import(self, task_label, button, worker, ui_callback, error_prefix, ui_error_prefix=None):
        self._ensure_import_runtime_state()
        if button is not None:
            try:
                button.state(['disabled'])
            except tk.TclError:
                pass
        self._on_import_started(task_label)
        run_async = getattr(self, 'root', None) is not None
        ui_error_prefix = ui_error_prefix or error_prefix

        def _execute_worker():
            try:
                payload = worker()
            except Exception as exc:  # pragma: no cover - errores inesperados
                self._handle_import_failure(task_label, button, exc, error_prefix)
                return
            self._handle_import_success(task_label, button, ui_callback, payload, ui_error_prefix)

        if not run_async:
            _execute_worker()
            return

        def _thread_target():
            try:
                payload = worker()
            except Exception as exc:  # pragma: no cover - errores inesperados
                self._dispatch_to_ui(self._handle_import_failure, task_label, button, exc, error_prefix)
                return
            self._dispatch_to_ui(self._handle_import_success, task_label, button, ui_callback, payload, ui_error_prefix)

        threading.Thread(
            target=_thread_target,
            name=f"import-{task_label}",
            daemon=True,
        ).start()

    def _handle_import_success(self, task_label, button, ui_callback, payload, error_prefix):
        message = f"Importación de {task_label} finalizada."
        failed = False
        captured_error = None
        try:
            ui_callback(payload)
        except Exception as exc:
            failed = True
            message = f"Importación de {task_label} con errores."
            captured_error = exc
            try:
                messagebox.showerror("Error", f"{error_prefix}: {exc}")
            except tk.TclError:
                pass
        finally:
            if button is not None and not self._catalog_loading:
                try:
                    button.state(['!disabled'])
                except tk.TclError:
                    pass
            self._finalize_import_task(message, failed=failed)
        if captured_error is not None and getattr(self, 'root', None) is None:
            raise captured_error

    def _handle_import_failure(self, task_label, button, error, error_prefix):
        if button is not None and not self._catalog_loading:
            try:
                button.state(['!disabled'])
            except tk.TclError:
                pass
        self._finalize_import_task(f"Error al importar {task_label}.", failed=True)
        try:
            messagebox.showerror("Error", f"{error_prefix}: {error}")
        except tk.TclError:
            pass
        if getattr(self, 'root', None) is None:
            raise error

    def _normalize_detail_catalog_payload(self, raw_catalogs):
        normalized = {
            normalize_detail_catalog_key(key): dict(value or {})
            for key, value in (raw_catalogs or {}).items()
        }
        lookup_by_id = build_detail_catalog_id_index(normalized)
        for canonical_key, aliases in (DETAIL_LOOKUP_ALIASES or {}).items():
            canonical = normalize_detail_catalog_key(canonical_key)
            lookup = normalized.get(canonical) or lookup_by_id.get(canonical)
            if not lookup:
                for alias in aliases or ():
                    alias_key = normalize_detail_catalog_key(alias)
                    alias_lookup = normalized.get(alias_key)
                    if alias_lookup:
                        lookup = alias_lookup
                        break
            if not lookup:
                continue
            normalized[canonical] = lookup
            lookup_by_id[canonical] = lookup
            for alias in aliases or ():
                alias_key = normalize_detail_catalog_key(alias)
                if not alias_key:
                    continue
                normalized[alias_key] = lookup
                lookup_by_id[alias_key] = lookup
        return normalized, lookup_by_id

    def _apply_catalog_lookups(self, detail_catalogs, detail_lookup_by_id):
        self.detail_catalogs = detail_catalogs or {}
        self.detail_lookup_by_id = detail_lookup_by_id or {}
        self.client_lookup = self._extract_lookup_or_empty("id_cliente")
        raw_team_lookup = self._extract_lookup_or_empty("id_colaborador")
        self.team_lookup = {
            normalized: value
            for key, value in raw_team_lookup.items()
            if (normalized := self._normalize_identifier(key))
        }
        self.product_lookup = self._extract_lookup_or_empty("id_producto")
        self.claim_lookup = self._extract_lookup_or_empty("id_reclamo")
        self.risk_lookup = self._extract_lookup_or_empty("id_riesgo")
        self.norm_lookup = self._extract_lookup_or_empty("id_norma")
        for frame in self.client_frames:
            if hasattr(frame, 'set_lookup'):
                frame.set_lookup(self.client_lookup)
                frame.on_id_change(preserve_existing=True, silent=True)
        for frame in self.team_frames:
            if hasattr(frame, 'set_lookup'):
                frame.set_lookup(self.team_lookup)
                frame.on_id_change(preserve_existing=True, silent=True)
        for frame in self.product_frames:
            if hasattr(frame, 'set_product_lookup'):
                frame.set_product_lookup(self.product_lookup)
                frame.on_id_change(preserve_existing=True, silent=True)
            if hasattr(frame, 'set_claim_lookup'):
                frame.set_claim_lookup(self.claim_lookup)
        for frame in self.risk_frames:
            if hasattr(frame, 'set_lookup'):
                frame.set_lookup(self.risk_lookup)
                frame.on_id_change(preserve_existing=True, silent=True)
        for frame in self.norm_frames:
            if hasattr(frame, 'set_lookup'):
                frame.set_lookup(self.norm_lookup)
                frame.on_id_change(preserve_existing=True, silent=True)

    def _extract_lookup_or_empty(self, canonical_key):
        normalized = normalize_detail_catalog_key(canonical_key)
        lookup = self.detail_lookup_by_id.get(normalized)
        if isinstance(lookup, dict):
            return lookup
        return {}

    def build_summary_tab(self, parent):
        """Construye la pestaña de resumen con tablas compactas."""

        self.summary_tab = parent
        container = ttk.Frame(parent)
        container.pack(fill="both", expand=True, padx=5, pady=5)
        ttk.Label(
            container,
            text="Resumen compacto de los datos capturados. Las tablas se actualizan tras cada guardado o importación.",
        ).pack(anchor="w", pady=(0, 5))

        config = [
            (
                "clientes",
                "Clientes registrados",
                [
                    ("id", "ID"),
                    ("tipo", "Tipo ID"),
                    ("flag", "Flag"),
                    ("telefonos", "Teléfonos"),
                    ("correos", "Correos"),
                    ("direcciones", "Direcciones"),
                    ("accionado", "Accionado"),
                ],
            ),
            (
                "colaboradores",
                "Colaboradores involucrados",
                [
                    ("id", "ID"),
                    ("division", "División"),
                    ("area", "Área"),
                    ("sancion", "Sanción"),
                ],
            ),
            (
                "involucramientos",
                "Asignaciones por colaborador",
                [
                    ("producto", "Producto"),
                    ("colaborador", "Colaborador"),
                    ("monto", "Monto asignado"),
                ],
            ),
            (
                "productos",
                "Productos investigados",
                [
                    ("id", "ID Producto"),
                    ("cliente", "Cliente"),
                    ("tipo", "Tipo"),
                    ("monto", "Monto investigado"),
                ],
            ),
            (
                "riesgos",
                "Riesgos registrados",
                [
                    ("id", "ID Riesgo"),
                    ("lider", "Líder"),
                    ("criticidad", "Criticidad"),
                    ("exposicion", "Exposición"),
                ],
            ),
            (
                "reclamos",
                "Reclamos asociados",
                [
                    ("id", "ID Reclamo"),
                    ("producto", "Producto"),
                    ("analitica", "Analítica"),
                    ("codigo", "Código analítica"),
                ],
            ),
            (
                "normas",
                "Normas transgredidas",
                [
                    ("id", "ID Norma"),
                    ("descripcion", "Descripción"),
                    ("vigencia", "Vigencia"),
                ],
            ),
        ]

        self.summary_tables.clear()
        self.summary_config = {key: columns for key, _, columns in config}
        for key, title, columns in config:
            section = ttk.LabelFrame(container, text=title)
            section.pack(fill="both", expand=True, pady=5)
            frame = ttk.Frame(section)
            frame.pack(fill="both", expand=True)
            tree = ttk.Treeview(frame, columns=[col for col, _ in columns], show="headings", height=5)
            scrollbar = ttk.Scrollbar(frame, orient="vertical")
            for col_id, heading in columns:
                tree.heading(col_id, text=heading)
                tree.column(col_id, width=150, stretch=True)
            tree.configure(yscrollcommand=scrollbar.set)
            scrollbar.configure(command=tree.yview)
            tree.pack(side="left", fill="both", expand=True)
            scrollbar.pack(side="right", fill="y")
            self.summary_tables[key] = tree
            self._register_summary_tree_bindings(tree, key)

        self._schedule_summary_refresh()

    def _handle_notebook_tab_change(self, event):
        if getattr(self, "notebook", None) is None:
            return
        if event.widget is not self.notebook:
            return
        if self._is_summary_tab_visible():
            self._flush_summary_refresh()

    def _register_summary_tree_bindings(self, tree, key):
        """Configura atajos de pegado y el menú contextual para una tabla."""

        tree.bind("<Control-v>", lambda event, target=key: self._handle_summary_paste(target))
        tree.bind("<Control-V>", lambda event, target=key: self._handle_summary_paste(target))
        menu = tk.Menu(tree, tearoff=False)
        menu.add_command(
            label="Pegar desde portapapeles",
            command=lambda target=key: self._handle_summary_paste(target),
        )
        tree.bind(
            "<Button-3>",
            lambda event, context_menu=menu: self._show_summary_context_menu(event, context_menu),
        )
        self.summary_context_menus[key] = menu

    def _show_summary_context_menu(self, event, menu):
        try:
            menu.tk_popup(event.x_root, event.y_root)
        finally:
            menu.grab_release()

    def _handle_summary_paste(self, key):
        """Lee el portapapeles, valida y carga filas en la tabla indicada."""

        tree = self.summary_tables.get(key)
        columns = self.summary_config.get(key, [])
        if not tree or not columns:
            return "break"
        try:
            clipboard_text = self.clipboard_get()
        except tk.TclError:
            messagebox.showerror("Portapapeles", "No se pudo leer el portapapeles desde el sistema.")
            return "break"
        try:
            parsed_rows = self._parse_clipboard_rows(clipboard_text, len(columns))
            sanitized_rows = self._transform_summary_clipboard_rows(key, parsed_rows)
        except ValueError as exc:
            messagebox.showerror("Pegado no válido", str(exc))
            return "break"
        ingestible_sections = {
            "clientes",
            "colaboradores",
            "involucramientos",
            "productos",
            "reclamos",
            "riesgos",
            "normas",
        }
        if key in ingestible_sections:
            try:
                self.ingest_summary_rows(key, sanitized_rows, stay_on_summary=True)
            except ValueError as exc:
                messagebox.showerror("Pegado no válido", str(exc))
            return "break"
        tree.delete(*tree.get_children())
        for row in sanitized_rows:
            tree.insert("", "end", values=row)
        return "break"

    def _parse_clipboard_rows(self, text, expected_columns):
        """Convierte el texto del portapapeles en una matriz con ``expected_columns`` celdas."""

        cleaned = (text or "").strip()
        if not cleaned:
            raise ValueError("El portapapeles está vacío.")
        lines = [line for line in cleaned.splitlines() if line.strip()]
        if not lines:
            raise ValueError("No se encontraron filas para pegar.")
        rows = []
        for idx, line in enumerate(lines, start=1):
            delimiter = "\t" if "\t" in line else ";"
            parts = [cell.strip() for cell in line.split(delimiter)]
            if len(parts) != expected_columns:
                raise ValueError(
                    f"La fila {idx} tiene {len(parts)} columnas y se esperaban {expected_columns}."
                )
            rows.append(parts)
        return rows

    def _transform_summary_clipboard_rows(self, key, rows):
        """Valida y normaliza filas de acuerdo al tipo de tabla del resumen."""

        handlers = {
            "clientes": self._transform_clipboard_clients,
            "colaboradores": self._transform_clipboard_colaboradores,
            "involucramientos": self._transform_clipboard_involucramientos,
            "productos": self._transform_clipboard_productos,
            "reclamos": self._transform_clipboard_reclamos,
            "riesgos": self._transform_clipboard_riesgos,
            "normas": self._transform_clipboard_normas,
        }
        handler = handlers.get(key)
        if not handler:
            raise ValueError("Esta tabla no admite pegado desde portapapeles.")
        return handler(rows)

    def _transform_clipboard_clients(self, rows):
        sanitized = []
        for idx, values in enumerate(rows, start=1):
            client_data = {
                "id_cliente": values[0].strip(),
                "tipo_id": values[1].strip(),
                "flag": values[2].strip(),
                "telefonos": values[3].strip(),
                "correos": values[4].strip(),
                "direcciones": values[5].strip(),
                "accionado": values[6].strip(),
            }
            tipo_id = client_data["tipo_id"]
            if tipo_id and tipo_id not in TIPO_ID_LIST:
                raise ValueError(
                    f"Cliente fila {idx}: el tipo de ID '{tipo_id}' no está en el catálogo CM."
                    " Corrige la hoja de Excel antes de volver a intentarlo."
                )
            flag_value = client_data["flag"]
            if flag_value and flag_value not in FLAG_CLIENTE_LIST:
                raise ValueError(
                    f"Cliente fila {idx}: el flag de cliente '{flag_value}' no está en el catálogo CM."
                    " Corrige la hoja de Excel antes de volver a intentarlo."
                )
            message = validate_client_id(client_data["tipo_id"], client_data["id_cliente"])
            if message:
                raise ValueError(f"Cliente fila {idx}: {message}")
            phone_required = validate_required_text(
                client_data["telefonos"], "los teléfonos del cliente"
            )
            if phone_required:
                raise ValueError(f"Cliente fila {idx}: {phone_required}")
            phone_message = validate_phone_list(client_data["telefonos"], "los teléfonos del cliente")
            if phone_message:
                raise ValueError(f"Cliente fila {idx}: {phone_message}")
            email_required = validate_required_text(
                client_data["correos"], "los correos del cliente"
            )
            if email_required:
                raise ValueError(f"Cliente fila {idx}: {email_required}")
            email_message = validate_email_list(client_data["correos"], "los correos del cliente")
            if email_message:
                raise ValueError(f"Cliente fila {idx}: {email_message}")
            sanitized.append(
                (
                    client_data["id_cliente"],
                    client_data["tipo_id"],
                    client_data["flag"],
                    client_data["telefonos"],
                    client_data["correos"],
                    client_data["direcciones"],
                    client_data["accionado"],
                )
            )
        return sanitized

    def _transform_clipboard_colaboradores(self, rows):
        sanitized = []
        for idx, values in enumerate(rows, start=1):
            collaborator = {
                "id_colaborador": values[0].strip(),
                "division": values[1].strip(),
                "area": values[2].strip(),
                "tipo_sancion": values[3].strip(),
            }
            if collaborator["tipo_sancion"] not in TIPO_SANCION_LIST:
                raise ValueError(
                    f"Colaborador fila {idx}: debe seleccionar un tipo de sanción válido."
                )
            message = validate_team_member_id(collaborator["id_colaborador"])
            if message:
                raise ValueError(f"Colaborador fila {idx}: {message}")
            sanitized.append(
                (
                    collaborator["id_colaborador"],
                    collaborator["division"],
                    collaborator["area"],
                    collaborator["tipo_sancion"],
                )
            )
        return sanitized

    def _resolve_product_type_for_involvement(self, product_id):
        """Busca el tipo de producto desde el resumen o el formulario si está disponible."""

        lookup = getattr(self, "product_lookup", None)
        if isinstance(lookup, dict):
            entry = lookup.get(product_id)
            if isinstance(entry, dict):
                tipo = (entry.get("tipo_producto") or "").strip()
                if tipo:
                    return tipo
        finder = getattr(self, "_find_product_frame", None)
        if callable(finder):
            frame = finder(product_id)
            tipo_var = getattr(frame, "tipo_prod_var", None) if frame else None
            if tipo_var and hasattr(tipo_var, "get"):
                return (tipo_var.get() or "").strip()
        return ""

    def _transform_clipboard_involucramientos(self, rows):
        sanitized = []
        for idx, values in enumerate(rows, start=1):
            product_id = (values[0] or "").strip()
            collaborator_id = (values[1] or "").strip()
            amount_text = (values[2] or "").strip()
            tipo_producto = self._resolve_product_type_for_involvement(product_id)
            if tipo_producto:
                product_message = validate_product_id(tipo_producto, product_id)
            else:
                product_message = validate_required_text(
                    product_id,
                    "el ID del producto involucrado",
                )
            if product_message:
                raise ValueError(f"Involucramiento fila {idx}: {product_message}")
            collaborator_message = validate_team_member_id(collaborator_id)
            if collaborator_message:
                raise ValueError(f"Involucramiento fila {idx}: {collaborator_message}")
            amount_message, _decimal_value, normalized_amount = validate_money_bounds(
                amount_text,
                "el monto asignado",
                allow_blank=False,
            )
            if amount_message:
                raise ValueError(f"Involucramiento fila {idx}: {amount_message}")
            sanitized.append((product_id, collaborator_id, normalized_amount))
        return sanitized

    def _transform_clipboard_productos(self, rows):
        sanitized = []
        for idx, values in enumerate(rows, start=1):
            product = {
                "id_producto": values[0].strip(),
                "id_cliente": values[1].strip(),
                "tipo_producto": values[2].strip(),
                "monto_investigado": values[3].strip(),
            }
            if not product["id_cliente"]:
                raise ValueError(f"Producto fila {idx}: el ID de cliente es obligatorio.")
            tipo_catalogo = resolve_catalog_product_type(product["tipo_producto"])
            if not tipo_catalogo:
                if product["tipo_producto"]:
                    raise ValueError(
                        f"Producto fila {idx}: el tipo de producto '{product['tipo_producto']}' no está en el catálogo CM."
                    )
                raise ValueError(f"Producto fila {idx}: debe ingresar el tipo de producto.")
            product["tipo_producto"] = tipo_catalogo
            message = validate_product_id(product["tipo_producto"], product["id_producto"])
            if message:
                raise ValueError(f"Producto fila {idx}: {message}")
            amount_message, decimal_value, _ = validate_money_bounds(
                product["monto_investigado"],
                "el monto investigado",
                allow_blank=False,
            )
            if amount_message:
                raise ValueError(f"Producto fila {idx}: {amount_message}")
            sanitized.append(
                (
                    product["id_producto"],
                    product["id_cliente"],
                    product["tipo_producto"],
                    f"{decimal_value:.2f}",
                )
            )
        return sanitized

    def _transform_clipboard_reclamos(self, rows):
        sanitized = []
        for idx, values in enumerate(rows, start=1):
            claim = {
                "id_reclamo": values[0].strip(),
                "id_producto": values[1].strip(),
                "nombre_analitica": values[2].strip(),
                "codigo_analitica": values[3].strip(),
            }
            message = validate_reclamo_id(claim["id_reclamo"])
            if message:
                raise ValueError(f"Reclamo fila {idx}: {message}")
            product_message = validate_required_text(claim["id_producto"], "el ID de producto")
            if product_message:
                raise ValueError(f"Reclamo fila {idx}: {product_message}")
            analytic_message = validate_required_text(
                claim["nombre_analitica"],
                "el nombre de la analítica",
            )
            if analytic_message:
                raise ValueError(f"Reclamo fila {idx}: {analytic_message}")
            code_message = validate_codigo_analitica(claim["codigo_analitica"])
            if code_message:
                raise ValueError(f"Reclamo fila {idx}: {code_message}")
            sanitized.append(
                (
                    claim["id_reclamo"],
                    claim["id_producto"],
                    claim["nombre_analitica"],
                    claim["codigo_analitica"],
                )
            )
        return sanitized

    def _transform_clipboard_riesgos(self, rows):
        sanitized = []
        valid_criticidades = set(CRITICIDAD_LIST)
        valid_criticidades_text = ", ".join(CRITICIDAD_LIST)
        for idx, values in enumerate(rows, start=1):
            risk = {
                "id_riesgo": values[0].strip().upper(),
                "lider": values[1].strip(),
                "criticidad": values[2].strip() or CRITICIDAD_LIST[0],
                "exposicion": values[3].strip(),
            }
            message = validate_risk_id(risk["id_riesgo"])
            if message:
                raise ValueError(f"Riesgo fila {idx}: {message}")
            if risk["criticidad"] not in valid_criticidades:
                raise ValueError(
                    f"Riesgo fila {idx}: la criticidad debe ser una de {valid_criticidades_text}."
                )
            exposure_message, exposure_decimal, _ = validate_money_bounds(
                risk["exposicion"],
                "la exposición residual",
                allow_blank=True,
            )
            if exposure_message:
                raise ValueError(f"Riesgo fila {idx}: {exposure_message}")
            exposure_text = f"{exposure_decimal:.2f}" if exposure_decimal is not None else ""
            sanitized.append(
                (
                    risk["id_riesgo"],
                    risk["lider"],
                    risk["criticidad"],
                    exposure_text,
                )
            )
        return sanitized

    def _transform_clipboard_normas(self, rows):
        sanitized = []
        for idx, values in enumerate(rows, start=1):
            norm = {
                "id_norma": values[0].strip(),
                "descripcion": values[1].strip(),
                "vigencia": values[2].strip(),
            }
            message = validate_norm_id(norm["id_norma"])
            if message:
                raise ValueError(f"Norma fila {idx}: {message}")
            desc_message = validate_required_text(norm["descripcion"], "la descripción de la norma")
            if desc_message:
                raise ValueError(f"Norma fila {idx}: {desc_message}")
            date_message = validate_date_text(norm["vigencia"], "la fecha de vigencia")
            if date_message:
                raise ValueError(f"Norma fila {idx}: {date_message}")
            sanitized.append(
                (
                    norm["id_norma"],
                    norm["descripcion"],
                    norm["vigencia"],
                )
            )
        return sanitized

    def ingest_summary_rows(self, section_key, rows, stay_on_summary=False):
        """Incorpora filas pegadas en las tablas de resumen al formulario principal."""

        if not rows:
            return 0
        section_key = (section_key or "").strip().lower()
        processed = 0
        missing_ids = []
        if section_key == "clientes":
            for idx, values in enumerate(rows, start=1):
                payload = {
                    "id_cliente": (values[0] or "").strip(),
                    "tipo_id": (values[1] or "").strip(),
                    "flag": (values[2] or "").strip(),
                    "telefonos": (values[3] or "").strip(),
                    "correos": (values[4] or "").strip(),
                    "direcciones": (values[5] or "").strip(),
                    "accionado": (values[6] or "").strip(),
                }
                hydrated, found = self._hydrate_row_from_details(payload, 'id_cliente', CLIENT_ID_ALIASES)
                client_id = (hydrated.get('id_cliente') or '').strip()
                if not client_id:
                    continue
                tipo_id = (hydrated.get('tipo_id') or '').strip()
                if tipo_id and tipo_id not in TIPO_ID_LIST:
                    raise ValueError(
                        f"Cliente fila {idx}: el tipo de ID '{tipo_id}' no está en el catálogo CM."
                        " Corrige la hoja de Excel antes de volver a intentarlo."
                    )
                flag_value = (hydrated.get('flag') or '').strip()
                if flag_value and flag_value not in FLAG_CLIENTE_LIST:
                    raise ValueError(
                        f"Cliente fila {idx}: el flag de cliente '{flag_value}' no está en el catálogo CM."
                        " Corrige la hoja de Excel antes de volver a intentarlo."
                    )
                frame = self._find_client_frame(client_id) or self._obtain_client_slot_for_import()
                merged = self._merge_client_payload_with_frame(frame, hydrated)
                self._populate_client_frame_from_row(frame, merged)
                self._trigger_import_id_refresh(
                    frame,
                    client_id,
                    notify_on_missing=True,
                    preserve_existing=True,
                )
                processed += 1
                if not found and 'id_cliente' in self.detail_catalogs:
                    missing_ids.append(client_id)
            if missing_ids:
                self._report_missing_detail_ids("clientes", missing_ids)
            if processed:
                self._notify_dataset_changed(summary_sections="clientes")
                self.sync_main_form_after_import("clientes", stay_on_summary=stay_on_summary)
            return processed
        if section_key == "colaboradores":
            for values in rows:
                payload = {
                    "id_colaborador": (values[0] or "").strip(),
                    "division": (values[1] or "").strip(),
                    "area": (values[2] or "").strip(),
                    "tipo_sancion": (values[3] or "").strip(),
                    "flag_colaborador": "No aplica",
                    "servicio": "",
                    "puesto": "",
                    "nombre_agencia": "",
                    "codigo_agencia": "",
                    "tipo_falta": "No aplica",
                }
                hydrated, found = self._hydrate_row_from_details(payload, 'id_colaborador', TEAM_ID_ALIASES)
                collaborator_id = (hydrated.get('id_colaborador') or '').strip()
                if not collaborator_id:
                    continue
                frame = self._find_team_frame(collaborator_id) or self._obtain_team_slot_for_import()
                merged = self._merge_team_payload_with_frame(frame, hydrated)
                self._populate_team_frame_from_row(frame, merged)
                self._trigger_import_id_refresh(
                    frame,
                    collaborator_id,
                    notify_on_missing=True,
                    preserve_existing=True,
                )
                processed += 1
                if not found and 'id_colaborador' in self.detail_catalogs:
                    missing_ids.append(collaborator_id)
            if missing_ids:
                self._report_missing_detail_ids("colaboradores", missing_ids)
            if processed:
                self._notify_dataset_changed(summary_sections="colaboradores")
                self.sync_main_form_after_import("colaboradores", stay_on_summary=stay_on_summary)
            return processed
        if section_key == "involucramientos":
            processed = 0

            def _fallback_frame(frames, identifier):
                normalized = (identifier or '').strip()
                if not normalized:
                    return None
                for frame in frames or []:
                    id_var = getattr(frame, 'id_var', None)
                    current = ''
                    if id_var and hasattr(id_var, 'get'):
                        current = (id_var.get() or '').strip()
                    if current == normalized:
                        return frame
                return None

            for idx, values in enumerate(rows, start=1):
                product_id = (values[0] or "").strip()
                collaborator_id = (values[1] or "").strip()
                amount_text = (values[2] or "").strip()
                if not product_id or not collaborator_id:
                    continue
                product_frame = self._find_product_frame(product_id) or _fallback_frame(getattr(self, 'product_frames', []), product_id)
                if not product_frame:
                    product_payload, product_found = self._hydrate_row_from_details(
                        {"id_producto": product_id},
                        'id_producto',
                        PRODUCT_ID_ALIASES,
                    )
                    if not product_found:
                        raise ValueError(
                            f"Involucramiento fila {idx}: el producto '{product_id}' no existe en el formulario ni en los catálogos de detalle."
                        )
                    client_for_product = (product_payload.get('id_cliente') or '').strip()
                    if client_for_product:
                        client_details, _ = self._hydrate_row_from_details({'id_cliente': client_for_product}, 'id_cliente', CLIENT_ID_ALIASES)
                        self._ensure_client_exists(client_for_product, client_details)
                    product_frame = self._obtain_product_slot_for_import()
                    merged = self._merge_product_payload_with_frame(product_frame, product_payload)
                    self._populate_product_frame_from_row(product_frame, merged)
                self._trigger_import_id_refresh(
                    product_frame,
                    product_id,
                    notify_on_missing=True,
                    preserve_existing=True,
                )
                team_frame = self._find_team_frame(collaborator_id) or _fallback_frame(getattr(self, 'team_frames', []), collaborator_id)
                if not team_frame:
                    collaborator_payload, collaborator_found = self._hydrate_row_from_details(
                        {"id_colaborador": collaborator_id},
                        'id_colaborador',
                        TEAM_ID_ALIASES,
                    )
                    if not collaborator_found:
                        raise ValueError(
                            f"Involucramiento fila {idx}: el colaborador '{collaborator_id}' no existe en el formulario ni en los catálogos de detalle."
                        )
                    team_frame, _created = self._ensure_team_member_exists(collaborator_id, collaborator_payload)
                self._trigger_import_id_refresh(
                    team_frame,
                    collaborator_id,
                    notify_on_missing=True,
                    preserve_existing=True,
                )
                amount_message, _amount_decimal, normalized_amount = validate_money_bounds(
                    amount_text,
                    "el monto asignado",
                    allow_blank=False,
                )
                if amount_message:
                    raise ValueError(f"Involucramiento fila {idx}: {amount_message}")
                existing_row = next(
                    (inv for inv in getattr(product_frame, 'involvements', []) if inv.team_var.get().strip() == collaborator_id),
                    None,
                )
                if not existing_row:
                    existing_row = self._obtain_involvement_slot(product_frame)
                existing_row.team_var.set(collaborator_id)
                team_widget = getattr(existing_row, 'team_cb', None)
                if team_widget is not None:
                    try:
                        team_widget.set(collaborator_id)
                    except tk.TclError:
                        pass
                existing_row.monto_var.set(normalized_amount)
                processed += 1
            if processed:
                self._notify_dataset_changed(summary_sections="involucramientos")
                self.save_auto()
                self.sync_main_form_after_import("involucramientos", stay_on_summary=stay_on_summary)
            return processed
        if section_key == "productos":
            for values in rows:
                payload = {
                    "id_producto": (values[0] or "").strip(),
                    "id_cliente": (values[1] or "").strip(),
                    "tipo_producto": (values[2] or "").strip(),
                    "monto_investigado": (values[3] or "").strip(),
                    "categoria1": "",
                    "categoria2": "",
                    "modalidad": "",
                    "canal": "",
                    "proceso": "",
                    "fecha_ocurrencia": "",
                    "fecha_descubrimiento": "",
                    "tipo_moneda": "",
                    "monto_perdida_fraude": "",
                    "monto_falla_procesos": "",
                    "monto_contingencia": "",
                    "monto_recuperado": "",
                    "monto_pago_deuda": "",
                    "id_reclamo": "",
                    "nombre_analitica": "",
                    "codigo_analitica": "",
                }
                hydrated, found = self._hydrate_row_from_details(payload, 'id_producto', PRODUCT_ID_ALIASES)
                product_id = (hydrated.get('id_producto') or '').strip()
                if not product_id:
                    continue
                frame = self._find_product_frame(product_id) or self._obtain_product_slot_for_import()
                client_id = (hydrated.get('id_cliente') or '').strip()
                if client_id:
                    client_details, _ = self._hydrate_row_from_details({'id_cliente': client_id}, 'id_cliente', CLIENT_ID_ALIASES)
                    self._ensure_client_exists(client_id, client_details)
                merged = self._merge_product_payload_with_frame(frame, hydrated)
                self._populate_product_frame_from_row(frame, merged)
                self._trigger_import_id_refresh(
                    frame,
                    product_id,
                    notify_on_missing=True,
                    preserve_existing=True,
                )
                processed += 1
                if not found and 'id_producto' in self.detail_catalogs:
                    missing_ids.append(product_id)
            if missing_ids:
                self._report_missing_detail_ids("productos", missing_ids)
            if processed:
                self._notify_dataset_changed(summary_sections="productos")
                self.sync_main_form_after_import("productos", stay_on_summary=stay_on_summary)
            return processed
        if section_key == "reclamos":
            missing_products = []
            unhydrated_products = []
            for values in rows:
                row_dict = {
                    'id_reclamo': (values[0] or "").strip(),
                    'id_producto': (values[1] or "").strip(),
                    'nombre_analitica': (values[2] or "").strip(),
                    'codigo_analitica': (values[3] or "").strip(),
                }
                hydrated, found = self._hydrate_row_from_details(row_dict, 'id_producto', PRODUCT_ID_ALIASES)
                product_id = (hydrated.get('id_producto') or '').strip()
                if not product_id:
                    continue
                if not found:
                    unhydrated_products.append(product_id)
                product_frame = self._find_product_frame(product_id)
                new_product = False
                if not product_frame:
                    product_frame = self._obtain_product_slot_for_import()
                    new_product = True
                client_id = (hydrated.get('id_cliente') or '').strip()
                if client_id:
                    client_details, _ = self._hydrate_row_from_details({'id_cliente': client_id}, 'id_cliente', CLIENT_ID_ALIASES)
                    self._ensure_client_exists(client_id, client_details)
                if new_product:
                    if found:
                        self._populate_product_frame_from_row(product_frame, hydrated)
                    else:
                        # Solo registrar el ID y mantener defaults hasta que el usuario complete los campos obligatorios.
                        product_frame.id_var.set(product_id)
                self._trigger_import_id_refresh(
                    product_frame,
                    product_id,
                    preserve_existing=False,
                )
                claim_payload = {
                    'id_reclamo': (hydrated.get('id_reclamo') or row_dict.get('id_reclamo') or '').strip(),
                    'nombre_analitica': (hydrated.get('nombre_analitica') or row_dict.get('nombre_analitica') or '').strip(),
                    'codigo_analitica': (hydrated.get('codigo_analitica') or row_dict.get('codigo_analitica') or '').strip(),
                }
                if not any(claim_payload.values()):
                    continue
                target = product_frame.find_claim_by_id(claim_payload['id_reclamo']) if claim_payload['id_reclamo'] else None
                if not target:
                    target = product_frame.obtain_claim_slot()
                target.set_data(claim_payload)
                self._sync_product_lookup_claim_fields(product_frame, product_id)
                product_frame.persist_lookup_snapshot()
                processed += 1
                if not found and 'id_producto' in self.detail_catalogs:
                    missing_products.append(product_id)
            if processed:
                self._notify_dataset_changed(summary_sections="reclamos")
                self.sync_main_form_after_import("reclamos", stay_on_summary=stay_on_summary)
                log_event("navegacion", f"Reclamos pegados desde resumen: {processed}", self.logs)
            if missing_products:
                self._report_missing_detail_ids("productos", missing_products)
            if unhydrated_products:
                self._notify_products_created_without_details(unhydrated_products)
            return processed
        if section_key == "riesgos":
            duplicate_ids = []
            for values in rows:
                risk_id = (values[0] or "").strip()
                if not risk_id:
                    continue
                if any(r.id_var.get().strip() == risk_id for r in self.risk_frames):
                    log_event("validacion", f"Riesgo duplicado {risk_id} en pegado", self.logs)
                    duplicate_ids.append(risk_id)
                    continue
                self.add_risk()
                frame = self.risk_frames[-1]
                frame.id_var.set(risk_id)
                frame.lider_var.set((values[1] or "").strip())
                criticidad = (values[2] or CRITICIDAD_LIST[0]).strip()
                if criticidad in CRITICIDAD_LIST:
                    frame.criticidad_var.set(criticidad)
                frame.exposicion_var.set((values[3] or "").strip())
                processed += 1
            if duplicate_ids:
                messagebox.showwarning(
                    "Riesgos duplicados",
                    "Se ignoraron los siguientes riesgos ya existentes:\n" + ", ".join(duplicate_ids),
                )
            if processed:
                self._notify_dataset_changed(summary_sections="riesgos")
                self.sync_main_form_after_import("riesgos", stay_on_summary=stay_on_summary)
            return processed
        if section_key == "normas":
            duplicate_ids = []
            for values in rows:
                norm_id = (values[0] or "").strip()
                if not norm_id:
                    continue
                if any(n.id_var.get().strip() == norm_id for n in self.norm_frames):
                    log_event("validacion", f"Norma duplicada {norm_id} en pegado", self.logs)
                    duplicate_ids.append(norm_id)
                    continue
                self.add_norm()
                frame = self.norm_frames[-1]
                frame.id_var.set(norm_id)
                frame.descripcion_var.set((values[1] or "").strip())
                frame.fecha_var.set((values[2] or "").strip())
                processed += 1
            if duplicate_ids:
                messagebox.showwarning(
                    "Normas duplicadas",
                    "Se ignoraron las siguientes normas ya existentes:\n" + ", ".join(duplicate_ids),
                )
            if processed:
                self._notify_dataset_changed(summary_sections="normas")
                self.sync_main_form_after_import("normas", stay_on_summary=stay_on_summary)
            return processed
        raise ValueError("Esta tabla no admite pegado directo al formulario principal.")

    def _schedule_summary_refresh(self, sections=None, data=None):
        """Marca secciones como sucias y actualiza el resumen cuando proceda."""

        if not self.summary_tables:
            return
        normalized = self._normalize_summary_sections(sections)
        if not normalized:
            return
        self._summary_dirty_sections.update(normalized)
        self._summary_pending_dataset = data
        summary_visible = self._is_summary_tab_visible()
        if not summary_visible:
            self._cancel_summary_refresh_job()
            return
        if self._summary_refresh_after_id:
            return
        try:
            self._summary_refresh_after_id = self.root.after(
                self.SUMMARY_REFRESH_DELAY_MS,
                self._run_scheduled_summary_refresh,
            )
        except tk.TclError:
            self._summary_refresh_after_id = None
            self._flush_summary_refresh(sections=normalized, data=data)

    def _run_scheduled_summary_refresh(self):
        self._summary_refresh_after_id = None
        if not self._is_summary_tab_visible():
            return
        self._flush_summary_refresh()

    def _flush_summary_refresh(self, sections=None, data=None):
        if sections is None:
            targets = set(self._summary_dirty_sections)
        else:
            targets = self._normalize_summary_sections(sections)
        if not targets:
            return
        dataset = data
        if dataset is None:
            dataset = self._summary_pending_dataset
        self.refresh_summary_tables(data=dataset, sections=targets)
        self._summary_pending_dataset = None

    def _normalize_summary_sections(self, sections):
        if not self.summary_tables:
            return set()
        if sections is None:
            return set(self.summary_tables.keys())
        if isinstance(sections, str):
            sections = {sections}
        return {section for section in sections if section in self.summary_tables}

    def _cancel_summary_refresh_job(self):
        if not self._summary_refresh_after_id:
            return
        try:
            self.root.after_cancel(self._summary_refresh_after_id)
        except tk.TclError:
            pass
        self._summary_refresh_after_id = None

    def _is_summary_tab_visible(self):
        if getattr(self, "notebook", None) is None or self.summary_tab is None:
            return False
        try:
            return self.notebook.select() == str(self.summary_tab)
        except tk.TclError:
            return False

    def refresh_summary_tables(self, data=None, sections=None):
        """Actualiza las tablas de resumen con la información actual."""

        if not self.summary_tables:
            return
        normalized = self._normalize_summary_sections(sections)
        if not normalized:
            return
        self._cancel_summary_refresh_job()
        self._summary_dirty_sections.difference_update(normalized)
        dataset = data or self.gather_data()
        ordered_sections = [key for key in self.summary_tables.keys() if key in normalized]
        for key in ordered_sections:
            rows = self._build_summary_rows(key, dataset)
            self._render_summary_rows(key, rows)

    def _build_summary_rows(self, section, dataset):
        if section == "clientes":
            return [
                (
                    client.get("id_cliente", ""),
                    client.get("tipo_id", ""),
                    client.get("flag", ""),
                    client.get("telefonos", ""),
                    client.get("correos", ""),
                    client.get("direcciones", ""),
                    client.get("accionado", ""),
                )
                for client in dataset.get("clientes", [])
            ]
        if section == "colaboradores":
            return [
                (
                    col.get("id_colaborador", ""),
                    col.get("division", ""),
                    col.get("area", ""),
                    col.get("tipo_sancion", ""),
                )
                for col in dataset.get("colaboradores", [])
            ]
        if section == "involucramientos":
            return [
                (
                    inv.get("id_producto", ""),
                    inv.get("id_colaborador", ""),
                    inv.get("monto_asignado", ""),
                )
                for inv in dataset.get("involucramientos", [])
            ]
        if section == "productos":
            return [
                (
                    prod.get("id_producto", ""),
                    prod.get("id_cliente", ""),
                    prod.get("tipo_producto", ""),
                    prod.get("monto_investigado", ""),
                )
                for prod in dataset.get("productos", [])
            ]
        if section == "riesgos":
            return [
                (
                    risk.get("id_riesgo", ""),
                    risk.get("lider", ""),
                    risk.get("criticidad", ""),
                    risk.get("exposicion_residual", ""),
                )
                for risk in dataset.get("riesgos", [])
            ]
        if section == "reclamos":
            return [
                (
                    rec.get("id_reclamo", ""),
                    rec.get("id_producto", ""),
                    rec.get("nombre_analitica", ""),
                    rec.get("codigo_analitica", ""),
                )
                for rec in dataset.get("reclamos", [])
            ]
        if section == "normas":
            return [
                (
                    norm.get("id_norma", ""),
                    norm.get("descripcion", ""),
                    norm.get("fecha_vigencia", ""),
                )
                for norm in dataset.get("normas", [])
            ]
        return []

    def _render_summary_rows(self, key, rows):
        tree = self.summary_tables.get(key)
        if not tree:
            return
        tree.delete(*tree.get_children())
        for row in rows:
            tree.insert("", "end", values=row)

    # ---------------------------------------------------------------------
    # Importación desde CSV

    def _select_csv_file(self, sample_key, dialog_title):
        """Obtiene un CSV desde diálogo o usa el archivo masivo de ejemplo."""

        filename = None
        try:
            filename = filedialog.askopenfilename(title=dialog_title, filetypes=[("CSV Files", "*.csv")])
        except tk.TclError:
            filename = None
        if not filename:
            sample_path = MASSIVE_SAMPLE_FILES.get(sample_key)
            if sample_path and os.path.exists(sample_path):
                filename = sample_path
                log_event(
                    "navegacion",
                    f"Se usó el archivo masivo de ejemplo {os.path.basename(sample_path)} para {sample_key}.",
                    self.logs,
                )
        return filename

    def _get_detail_lookup(self, id_column):
        """Obtiene el diccionario de detalles considerando alias configurados."""

        normalized = normalize_detail_catalog_key(id_column)
        lookup = self.detail_lookup_by_id.get(normalized)
        if isinstance(lookup, dict):
            return lookup
        candidate_keys = [normalized]
        candidate_keys.extend(
            normalize_detail_catalog_key(alias)
            for alias in DETAIL_LOOKUP_ALIASES.get(normalized, ())
        )
        seen = set()
        for key in candidate_keys:
            if not key or key in seen:
                continue
            seen.add(key)
            lookup = self.detail_catalogs.get(key)
            if isinstance(lookup, dict):
                return lookup
        return None

    def _hydrate_row_from_details(self, row, id_column, alias_headers):
        """Devuelve una copia de la fila complementada con catálogos de detalle."""

        hydrated = dict(row or {})
        alias_headers = alias_headers or ()
        canonical_id = ""
        for header in (id_column, *alias_headers):
            value = hydrated.get(header)
            if isinstance(value, str):
                value = value.strip()
            if value:
                canonical_id = str(value).strip()
                hydrated[id_column] = canonical_id
                break
        found = False
        lookup = self._get_detail_lookup(id_column)
        if canonical_id and lookup:
            details = lookup.get(canonical_id)
            if details:
                for key, value in details.items():
                    if hydrated.get(key):
                        continue
                    hydrated[key] = value
                found = True
        return hydrated, found

    def _has_meaningful_value(self, value):
        if value is None:
            return False
        if isinstance(value, str):
            return bool(value.strip())
        return True

    def _merge_payload_with_frame(self, payload, field_sources):
        merged = dict(payload or {})
        for field, getter in (field_sources or {}).items():
            incoming = merged.get(field)
            if self._has_meaningful_value(incoming):
                continue
            existing = getter() if callable(getter) else getter
            if self._has_meaningful_value(existing):
                merged[field] = existing.strip() if isinstance(existing, str) else existing
        return merged

    def _merge_client_payload_with_frame(self, frame, payload):
        return self._merge_payload_with_frame(
            payload,
            {
                'tipo_id': frame.tipo_id_var.get,
                'flag': frame.flag_var.get,
                'telefonos': frame.telefonos_var.get,
                'correos': frame.correos_var.get,
                'direcciones': frame.direcciones_var.get,
                'accionado': frame.accionado_var.get,
            },
        )

    def _merge_team_payload_with_frame(self, frame, payload):
        return self._merge_payload_with_frame(
            payload,
            {
                'flag_colaborador': frame.flag_var.get,
                'division': frame.division_var.get,
                'area': frame.area_var.get,
                'servicio': frame.servicio_var.get,
                'puesto': frame.puesto_var.get,
                'nombre_agencia': frame.nombre_agencia_var.get,
                'codigo_agencia': frame.codigo_agencia_var.get,
                'tipo_falta': frame.tipo_falta_var.get,
                'tipo_sancion': frame.tipo_sancion_var.get,
            },
        )

    def _merge_product_payload_with_frame(self, frame, payload):
        merged = self._merge_payload_with_frame(
            payload,
            {
                'id_cliente': frame.client_var.get,
                'categoria1': frame.cat1_var.get,
                'categoria2': frame.cat2_var.get,
                'modalidad': frame.mod_var.get,
                'canal': frame.canal_var.get,
                'proceso': frame.proceso_var.get,
                'tipo_producto': frame.tipo_prod_var.get,
                'fecha_ocurrencia': frame.fecha_oc_var.get,
                'fecha_descubrimiento': frame.fecha_desc_var.get,
                'tipo_moneda': frame.moneda_var.get,
                'monto_investigado': frame.monto_inv_var.get,
                'monto_perdida_fraude': frame.monto_perdida_var.get,
                'monto_falla_procesos': frame.monto_falla_var.get,
                'monto_contingencia': frame.monto_cont_var.get,
                'monto_recuperado': frame.monto_rec_var.get,
                'monto_pago_deuda': frame.monto_pago_var.get,
            },
        )
        claim_candidates = frame.extract_claims_from_payload(payload or {})
        if claim_candidates:
            merged['reclamos'] = claim_candidates
        elif 'reclamos' not in merged:
            merged['reclamos'] = [claim.get_data() for claim in frame.claims]
        return merged

    def _report_missing_detail_ids(self, entity_label, missing_ids):
        """Registra y muestra un único aviso de IDs sin detalle catalogado."""

        unique_ids = sorted({mid for mid in missing_ids if mid})
        if not unique_ids:
            return
        message = (
            f"No se encontraron detalles catalogados para los {entity_label} con ID: "
            f"{', '.join(unique_ids)}."
        )
        log_event("validacion", message, self.logs)
        try:
            messagebox.showwarning("Detalles faltantes", message)
        except tk.TclError:
            pass

    def _notify_products_created_without_details(self, product_ids):
        """Advierte que ciertos productos se generaron sin información de catálogo."""

        unique_ids = sorted({pid for pid in product_ids if pid})
        if not unique_ids:
            return
        message = (
            "Los siguientes productos se importaron desde reclamos sin detalle"
            " catalogado: "
            + ", ".join(unique_ids)
            + ".\nEl sistema solo puede crear el producto con su ID,"
            " por lo que debes completar manualmente la sección de Productos o"
            " actualizar el catálogo antes de continuar."
        )
        log_event("validacion", message, self.logs)
        try:
            messagebox.showwarning("Productos creados sin datos", message)
        except tk.TclError:
            pass

    @staticmethod
    def _normalize_identifier(identifier):
        return (identifier or '').strip().upper()

    @staticmethod
    def _frame_entity_label(frame):
        label = getattr(frame, 'ENTITY_LABEL', '')
        if isinstance(label, str) and label.strip():
            return label.strip()
        return 'registro'

    def _update_frame_id_index(self, index, frame, previous_id, new_id):
        if index is None:
            return
        previous = self._normalize_identifier(previous_id)
        new = self._normalize_identifier(new_id)
        existing = index.get(new)
        if new and existing is not None and existing is not frame:
            restore_value = previous_id if isinstance(previous_id, str) else (previous_id or '')
            if not isinstance(restore_value, str):
                restore_value = str(restore_value)
            if hasattr(frame, 'id_var') and hasattr(frame.id_var, 'set'):
                frame.id_var.set(restore_value)
            if hasattr(frame, '_last_tracked_id'):
                frame._last_tracked_id = previous
            entity_label = self._frame_entity_label(frame)
            message = (
                f"El ID '{new}' ya está asignado a otro {entity_label}. "
                f"Cada {entity_label} debe tener un ID único."
            )
            log_event('validacion', message, self.logs)
            if not getattr(self, '_suppress_messagebox', False):
                try:
                    messagebox.showerror('ID duplicado', message)
                except tk.TclError:
                    pass
            return
        for key, value in list(index.items()):
            if value is frame and key != new:
                index.pop(key, None)
        if hasattr(frame, '_last_tracked_id'):
            frame._last_tracked_id = new
        if not new:
            return
        index[new] = frame

    def _handle_client_id_change(self, frame, previous_id, new_id):
        self._ensure_frame_id_maps()
        self._update_frame_id_index(self._client_frames_by_id, frame, previous_id, new_id)

    def _handle_team_id_change(self, frame, previous_id, new_id):
        self._ensure_frame_id_maps()
        self._update_frame_id_index(self._team_frames_by_id, frame, previous_id, new_id)

    def _handle_product_id_change(self, frame, previous_id, new_id):
        self._ensure_frame_id_maps()
        self._update_frame_id_index(self._product_frames_by_id, frame, previous_id, new_id)

    def _rebuild_frame_id_indexes(self):
        self._ensure_frame_id_maps()
        def _rebuild(target, frames):
            target.clear()
            for frame in frames:
                identifier = self._normalize_identifier(frame.id_var.get() if hasattr(frame, 'id_var') else '')
                if hasattr(frame, '_last_tracked_id'):
                    frame._last_tracked_id = identifier
                if identifier:
                    target[identifier] = frame

        _rebuild(self._client_frames_by_id, self.client_frames)
        _rebuild(self._team_frames_by_id, self.team_frames)
        _rebuild(self._product_frames_by_id, self.product_frames)

    def _ensure_frame_id_maps(self):
        if not hasattr(self, '_client_frames_by_id'):
            self._client_frames_by_id = {}
        if not hasattr(self, '_team_frames_by_id'):
            self._team_frames_by_id = {}
        if not hasattr(self, '_product_frames_by_id'):
            self._product_frames_by_id = {}

    def _find_client_frame(self, client_id):
        client_id = self._normalize_identifier(client_id)
        if not client_id:
            return None
        return getattr(self, '_client_frames_by_id', {}).get(client_id)

    def _find_team_frame(self, collaborator_id):
        collaborator_id = self._normalize_identifier(collaborator_id)
        if not collaborator_id:
            return None
        return getattr(self, '_team_frames_by_id', {}).get(collaborator_id)

    def _find_product_frame(self, product_id):
        product_id = self._normalize_identifier(product_id)
        if not product_id:
            return None
        return getattr(self, '_product_frames_by_id', {}).get(product_id)

    def _obtain_client_slot_for_import(self):
        """Obtiene un ``ClientFrame`` vacío o crea uno nuevo para importación.

        Paso a paso:
            1. Recorre la lista existente de clientes y detecta el primero
               cuyo ``id_cliente`` esté en blanco.
            2. Si encuentra un espacio vacío, se reutiliza para mostrar los
               datos recién importados (esto evita que el usuario piense que no
               se cargó nada por tener clientes vacíos al inicio).
            3. Si no existe un espacio vacío, se invoca ``add_client`` para
               crear un nuevo marco y finalmente se devuelve.

        Returns:
            ClientFrame: Instancia lista para llenarse con datos externos.

        Ejemplo::

            slot = self._obtain_client_slot_for_import()
            slot.id_var.set("12345678")
        """

        for frame in self.client_frames:
            if not frame.id_var.get().strip():
                return frame
        self.add_client()
        return self.client_frames[-1]

    def _obtain_team_slot_for_import(self):
        for frame in self.team_frames:
            if not frame.id_var.get().strip():
                return frame
        self.add_team()
        return self.team_frames[-1]

    def _obtain_product_slot_for_import(self):
        for frame in self.product_frames:
            if not frame.id_var.get().strip():
                return frame
        self.add_product(initialize_rows=False)
        new_frame = self.product_frames[-1]
        return new_frame

    def _obtain_involvement_slot(self, product_frame):
        empty = next((inv for inv in product_frame.involvements if not inv.team_var.get().strip()), None)
        if empty:
            return empty
        product_frame.add_involvement()
        return product_frame.involvements[-1]

    def _trigger_import_id_refresh(self, frame, identifier, notify_on_missing=False, preserve_existing=False):
        """Ejecuta ``on_id_change`` tras importaciones evitando mensajes redundantes.

        Args:
            frame: Instancia del frame que contiene el campo ID.
            identifier (str): ID normalizado que se está cargando.
            notify_on_missing (bool): Si es ``True`` se propaga ``from_focus=True``
                para que los cuadros de diálogo de inexistencia aparezcan igual que
                cuando el usuario edita el campo manualmente.
            preserve_existing (bool): Si es ``True`` los campos con datos
                recientes no se sobreescriben durante el autopoblado.
        """

        identifier = (identifier or '').strip()
        if identifier and hasattr(frame, 'on_id_change'):
            frame.on_id_change(from_focus=notify_on_missing, preserve_existing=preserve_existing)

    def _sync_product_lookup_claim_fields(self, frame, product_id):
        """Actualiza ``product_lookup`` con la lista de reclamos visibles."""

        if not frame:
            return
        product_id = (product_id or '').strip()
        if not product_id:
            return
        claim_values = [claim.get_data() for claim in getattr(frame, 'claims', [])]
        lookups = []
        if isinstance(self.product_lookup, dict):
            lookups.append(self.product_lookup)
        frame_lookup = getattr(frame, 'product_lookup', None)
        if isinstance(frame_lookup, dict) and frame_lookup is not self.product_lookup:
            lookups.append(frame_lookup)
        for lookup in lookups:
            entry = lookup.setdefault(product_id, {})
            entry['reclamos'] = claim_values

    def _populate_client_frame_from_row(self, frame, row, preserve_existing=False):
        """Traslada los datos de una fila CSV a un ``ClientFrame``.

        Cada línea de este método tiene como fin dejar explícito qué campo se
        está poblando y por qué::

            1. Se normaliza el texto con ``strip`` para evitar espacios.
            2. Se actualiza el ``StringVar`` correspondiente en el frame.
            3. Se sincroniza el selector múltiple de ``accionado``.
            4. Se actualiza ``client_lookup`` para futuros autopoblados.

        Args:
            frame (ClientFrame): Cliente objetivo.
            row (dict): Fila leída por ``csv.DictReader``.
            preserve_existing (bool): Si es ``True`` solo completa campos vacíos.

        Ejemplo::

            fila = {"id_cliente": "123", "telefonos": "999"}
            self._populate_client_frame_from_row(cliente, fila)
        """

        id_cliente = (row.get('id_cliente') or row.get('IdCliente') or '').strip()
        frame.id_var.set(id_cliente)
        tipo_id = (row.get('tipo_id') or row.get('TipoID') or '').strip()
        if tipo_id and should_autofill_field(frame.tipo_id_var.get(), preserve_existing):
            frame.tipo_id_var.set(tipo_id)
        elif not tipo_id and not preserve_existing:
            frame.tipo_id_var.set('')
        flag_value = (row.get('flag') or row.get('Flag') or '').strip()
        if flag_value and should_autofill_field(frame.flag_var.get(), preserve_existing):
            frame.flag_var.set(flag_value)
        elif not flag_value and not preserve_existing:
            frame.flag_var.set('')
        telefonos = (row.get('telefonos') or row.get('Telefono') or '').strip()
        if telefonos and should_autofill_field(frame.telefonos_var.get(), preserve_existing):
            frame.telefonos_var.set(telefonos)
        elif not telefonos and not preserve_existing:
            frame.telefonos_var.set('')
        correos = (row.get('correos') or row.get('Correo') or '').strip()
        if correos and should_autofill_field(frame.correos_var.get(), preserve_existing):
            frame.correos_var.set(correos)
        elif not correos and not preserve_existing:
            frame.correos_var.set('')
        direcciones = (row.get('direcciones') or row.get('Direccion') or '').strip()
        if direcciones and should_autofill_field(frame.direcciones_var.get(), preserve_existing):
            frame.direcciones_var.set(direcciones)
        elif not direcciones and not preserve_existing:
            frame.direcciones_var.set('')
        accionado_val = (row.get('accionado') or row.get('Accionado') or '').strip()
        if accionado_val and should_autofill_field(frame.accionado_var.get(), preserve_existing):
            frame.set_accionado_from_text(accionado_val)
        elif not accionado_val and not preserve_existing:
            frame.set_accionado_from_text('')
        accionado_final = frame.accionado_var.get().strip()
        self.client_lookup[id_cliente] = {
            'tipo_id': frame.tipo_id_var.get(),
            'flag': frame.flag_var.get(),
            'telefonos': frame.telefonos_var.get(),
            'correos': frame.correos_var.get(),
            'direcciones': frame.direcciones_var.get(),
            'accionado': accionado_final,
        }

    def _populate_team_frame_from_row(self, frame, row):
        id_col = (
            row.get('id_colaborador')
            or row.get('IdColaborador')
            or row.get('IdTeamMember')
            or row.get('id_col')
            or ''
        ).strip()
        normalized_id = self._normalize_identifier(id_col)
        frame.id_var.set(normalized_id or id_col)
        flag_val = (
            row.get('flag_colaborador')
            or row.get('flag')
            or row.get('Flag')
            or 'No aplica'
        ).strip()
        frame.flag_var.set(flag_val or 'No aplica')
        frame.division_var.set((row.get('division') or '').strip())
        frame.area_var.set((row.get('area') or '').strip())
        frame.servicio_var.set((row.get('servicio') or '').strip())
        frame.puesto_var.set((row.get('puesto') or '').strip())
        frame.nombre_agencia_var.set((row.get('nombre_agencia') or '').strip())
        frame.codigo_agencia_var.set((row.get('codigo_agencia') or '').strip())
        frame.tipo_falta_var.set((row.get('tipo_falta') or '').strip() or 'No aplica')
        frame.tipo_sancion_var.set((row.get('tipo_sancion') or '').strip() or 'No aplica')
        lookup_key = normalized_id or id_col
        if lookup_key:
            self.team_lookup[lookup_key] = {
                'division': frame.division_var.get(),
                'area': frame.area_var.get(),
                'servicio': frame.servicio_var.get(),
                'puesto': frame.puesto_var.get(),
                'nombre_agencia': frame.nombre_agencia_var.get(),
                'codigo_agencia': frame.codigo_agencia_var.get(),
            }

    def _populate_product_frame_from_row(self, frame, row):
        product_id = (row.get('id_producto') or row.get('IdProducto') or '').strip()
        if not product_id:
            return
        frame.id_var.set(product_id)
        client_id = (row.get('id_cliente') or row.get('IdCliente') or '').strip()
        if client_id:
            values = list(frame.client_cb['values'])
            if client_id not in values:
                values.append(client_id)
                frame.client_cb['values'] = values
            frame.client_var.set(client_id)
            frame.client_cb.set(client_id)
        cat1 = (row.get('categoria1') or '').strip()
        cat2 = (row.get('categoria2') or '').strip()
        mod = (row.get('modalidad') or '').strip()
        if cat1:
            if cat1 in TAXONOMIA:
                if frame.cat1_var.get() != cat1:
                    frame.cat1_var.set(cat1)
                    frame.on_cat1_change()
            else:
                self._notify_taxonomy_warning(
                    f"Producto {product_id}: la categoría 1 '{cat1}' no está en la taxonomía."
                )
                frame.cat1_var.set(cat1)
        if cat2:
            if cat1 in TAXONOMIA and cat2 in TAXONOMIA[cat1]:
                frame.cat2_var.set(cat2)
                frame.cat2_cb.set(cat2)
                frame.on_cat2_change()
            else:
                self._notify_taxonomy_warning(
                    f"Producto {product_id}: la categoría 2 '{cat2}' no corresponde a {cat1}."
                )
                frame.cat2_var.set(cat2)
        if mod:
            valid_mods = TAXONOMIA.get(cat1, {}).get(cat2, [])
            if mod in valid_mods:
                frame.mod_var.set(mod)
                frame.mod_cb.set(mod)
            else:
                self._notify_taxonomy_warning(
                    f"Producto {product_id}: la modalidad '{mod}' no corresponde a {cat1}/{cat2}."
                )
                frame.mod_var.set(mod)
        canal = (row.get('canal') or '').strip()
        if canal in CANAL_LIST:
            frame.canal_var.set(canal)
        proceso = (row.get('proceso') or '').strip()
        if proceso in PROCESO_LIST:
            frame.proceso_var.set(proceso)
        tipo_prod = (row.get('tipo_producto') or row.get('tipoProducto') or '').strip()
        if tipo_prod:
            resolved_tipo = resolve_catalog_product_type(tipo_prod)
            if resolved_tipo:
                frame.tipo_prod_var.set(resolved_tipo)
            else:
                self._notify_taxonomy_warning(
                    f"Producto {product_id}: el tipo de producto '{tipo_prod}' no está en el catálogo CM."
                )
        fecha_occ = (row.get('fecha_ocurrencia') or '').strip()
        if fecha_occ:
            frame.fecha_oc_var.set(fecha_occ)
        fecha_desc = (row.get('fecha_descubrimiento') or '').strip()
        if fecha_desc:
            frame.fecha_desc_var.set(fecha_desc)
        frame.monto_inv_var.set((row.get('monto_investigado') or '').strip())
        moneda = (row.get('tipo_moneda') or '').strip()
        if moneda in TIPO_MONEDA_LIST:
            frame.moneda_var.set(moneda)
        frame.monto_perdida_var.set((row.get('monto_perdida_fraude') or '').strip())
        frame.monto_falla_var.set((row.get('monto_falla_procesos') or '').strip())
        frame.monto_cont_var.set((row.get('monto_contingencia') or '').strip())
        frame.monto_rec_var.set((row.get('monto_recuperado') or '').strip())
        frame.monto_pago_var.set((row.get('monto_pago_deuda') or '').strip())
        frame.set_claims_from_data(frame.extract_claims_from_payload(row))
        frame.persist_lookup_snapshot()

    def _ensure_client_exists(self, client_id, row_data=None):
        client_id = (client_id or '').strip()
        if not client_id:
            return None, False
        frame = self._find_client_frame(client_id)
        created = False
        if not frame:
            frame = self._obtain_client_slot_for_import()
            created = True
        if created and row_data:
            payload = dict(row_data)
            payload['id_cliente'] = client_id
            self._populate_client_frame_from_row(frame, payload)
        return frame, created

    def _ensure_team_member_exists(self, collaborator_id, row_data=None):
        collaborator_id = (collaborator_id or '').strip()
        if not collaborator_id:
            return None, False
        frame = self._find_team_frame(collaborator_id)
        created = False
        if not frame:
            frame = self._obtain_team_slot_for_import()
            created = True
        if created and row_data:
            payload = dict(row_data)
            payload['id_colaborador'] = collaborator_id
            self._populate_team_frame_from_row(frame, payload)
        return frame, created

    def _apply_client_import_payload(self, entries):
        imported = 0
        missing_ids = []
        for entry in entries or []:
            hydrated = entry.get('row', {})
            found = entry.get('found', False)
            id_cliente = (hydrated.get('id_cliente') or '').strip()
            if not id_cliente:
                continue
            frame = self._find_client_frame(id_cliente) or self._obtain_client_slot_for_import()
            self._populate_client_frame_from_row(frame, hydrated, preserve_existing=True)
            self._trigger_import_id_refresh(
                frame,
                id_cliente,
                notify_on_missing=True,
                preserve_existing=False,
            )
            imported += 1
            if not found and 'id_cliente' in self.detail_catalogs:
                missing_ids.append(id_cliente)
        self._notify_dataset_changed(summary_sections="clientes")
        log_event("navegacion", f"Clientes importados desde CSV: {imported}", self.logs)
        if imported:
            self.sync_main_form_after_import("clientes")
            messagebox.showinfo("Importación completa", f"Se cargaron {imported} clientes.")
        else:
            messagebox.showwarning("Sin cambios", "El archivo no aportó clientes nuevos.")
        self._report_missing_detail_ids("clientes", missing_ids)

    def _apply_team_import_payload(self, entries):
        imported = 0
        missing_ids = []
        for entry in entries or []:
            hydrated = entry.get('row', {})
            found = entry.get('found', False)
            collaborator_id = (hydrated.get('id_colaborador') or '').strip()
            if not collaborator_id:
                continue
            frame = self._find_team_frame(collaborator_id) or self._obtain_team_slot_for_import()
            self._populate_team_frame_from_row(frame, hydrated)
            self._trigger_import_id_refresh(
                frame,
                collaborator_id,
                notify_on_missing=True,
                preserve_existing=False,
            )
            imported += 1
            if not found and 'id_colaborador' in self.detail_catalogs:
                missing_ids.append(collaborator_id)
        self._notify_dataset_changed(summary_sections="colaboradores")
        log_event("navegacion", f"Colaboradores importados desde CSV: {imported}", self.logs)
        if imported:
            self.sync_main_form_after_import("colaboradores")
            messagebox.showinfo("Importación completa", "Colaboradores importados correctamente.")
        else:
            messagebox.showwarning("Sin cambios", "No se encontraron colaboradores nuevos en el archivo.")
        self._report_missing_detail_ids("colaboradores", missing_ids)

    def _apply_product_import_payload(self, entries):
        imported = 0
        missing_ids = []
        for entry in entries or []:
            hydrated = entry.get('row', {})
            found = entry.get('found', False)
            product_id = (hydrated.get('id_producto') or '').strip()
            if not product_id:
                continue
            frame = self._find_product_frame(product_id) or self._obtain_product_slot_for_import()
            client_id = (hydrated.get('id_cliente') or '').strip()
            if client_id:
                client_details, _ = self._hydrate_row_from_details({'id_cliente': client_id}, 'id_cliente', CLIENT_ID_ALIASES)
                self._ensure_client_exists(client_id, client_details)
            self._populate_product_frame_from_row(frame, hydrated)
            self._trigger_import_id_refresh(
                frame,
                product_id,
                notify_on_missing=True,
                preserve_existing=False,
            )
            imported += 1
            if not found and 'id_producto' in self.detail_catalogs:
                missing_ids.append(product_id)
        self._notify_dataset_changed(summary_sections="productos")
        log_event("navegacion", f"Productos importados desde CSV: {imported}", self.logs)
        if imported:
            self.sync_main_form_after_import("productos")
            messagebox.showinfo("Importación completa", "Productos importados correctamente.")
        else:
            messagebox.showwarning("Sin cambios", "No se detectaron productos nuevos en el archivo.")
        self._report_missing_detail_ids("productos", missing_ids)

    def _apply_combined_import_payload(self, entries):
        created_records = False
        missing_clients = []
        missing_team = []
        missing_products = []
        for entry in entries or []:
            client_row = dict(entry.get('client_row') or {})
            client_found = entry.get('client_found', False)
            client_id = (client_row.get('id_cliente') or '').strip()
            if client_id:
                if not client_row.get('flag') and client_row.get('flag_cliente'):
                    client_row['flag'] = client_row.get('flag_cliente')
                for key in ('telefonos', 'correos', 'direcciones', 'accionado', 'tipo_id'):
                    value = client_row.get(key)
                    if not value and entry.get('raw_row', {}).get(key):
                        client_row[key] = entry['raw_row'][key]
                client_frame, created_client = self._ensure_client_exists(client_id, client_row)
                if created_client:
                    self._trigger_import_id_refresh(
                        client_frame,
                        client_id,
                        notify_on_missing=False,
                        preserve_existing=False,
                    )
                created_records = created_records or created_client
                if not client_found and 'id_cliente' in self.detail_catalogs:
                    missing_clients.append(client_id)
            team_row = dict(entry.get('team_row') or {})
            team_found = entry.get('team_found', False)
            collaborator_id = (team_row.get('id_colaborador') or '').strip()
            if collaborator_id:
                team_frame, created_team = self._ensure_team_member_exists(collaborator_id, team_row)
                if created_team:
                    self._trigger_import_id_refresh(
                        team_frame,
                        collaborator_id,
                        notify_on_missing=False,
                        preserve_existing=False,
                    )
                created_records = created_records or created_team
                if not team_found and 'id_colaborador' in self.detail_catalogs:
                    missing_team.append(collaborator_id)
            product_row = dict(entry.get('product_row') or {})
            product_found = entry.get('product_found', False)
            product_id = (product_row.get('id_producto') or '').strip()
            product_frame = None
            if product_id:
                product_frame = self._find_product_frame(product_id)
                new_product = False
                if not product_frame:
                    product_frame = self._obtain_product_slot_for_import()
                    new_product = True
                client_for_product = (product_row.get('id_cliente') or '').strip()
                if client_for_product:
                    client_details, _ = self._hydrate_row_from_details({'id_cliente': client_for_product}, 'id_cliente', CLIENT_ID_ALIASES)
                    self._ensure_client_exists(client_for_product, client_details)
                self._populate_product_frame_from_row(product_frame, product_row)
                self._trigger_import_id_refresh(
                    product_frame,
                    product_id,
                    notify_on_missing=False,
                    preserve_existing=False,
                )
                created_records = created_records or new_product
                if not product_found and 'id_producto' in self.detail_catalogs:
                    missing_products.append(product_id)
            involvement_pairs = entry.get('involvement_pairs') or []
            if product_frame and involvement_pairs:
                for collab_id, amount in involvement_pairs:
                    collab_id = (collab_id or '').strip()
                    if not collab_id:
                        continue
                    collab_details, collab_found = self._hydrate_row_from_details({'id_colaborador': collab_id}, 'id_colaborador', TEAM_ID_ALIASES)
                    _, created_team = self._ensure_team_member_exists(collab_id, collab_details)
                    created_records = created_records or created_team
                    if not collab_found and 'id_colaborador' in self.detail_catalogs:
                        missing_team.append(collab_id)
                    inv_row = next((inv for inv in product_frame.involvements if inv.team_var.get().strip() == collab_id), None)
                    if not inv_row:
                        inv_row = self._obtain_involvement_slot(product_frame)
                    inv_row.team_var.set(collab_id)
                    amount_text = (amount or '').strip()
                    label = (
                        f"Monto asignado del colaborador {collab_id or 'sin ID'} "
                        f"en el producto {product_id or 'sin ID'}"
                    )
                    error, _amount, normalized_text = validate_money_bounds(
                        amount_text,
                        label,
                    )
                    if error:
                        raise ValueError(error)
                    inv_row.monto_var.set(normalized_text or amount_text)
                    created_records = True
        self._notify_dataset_changed()
        log_event("navegacion", "Datos combinados importados desde CSV", self.logs)
        if created_records:
            self.sync_main_form_after_import("datos combinados")
            messagebox.showinfo("Importación completa", "Datos combinados importados correctamente.")
        else:
            messagebox.showwarning("Sin cambios", "No se detectaron registros nuevos en el archivo.")
        self._report_missing_detail_ids("clientes", missing_clients)
        self._report_missing_detail_ids("colaboradores", missing_team)
        self._report_missing_detail_ids("productos", missing_products)

    def _apply_risk_import_payload(self, entries):
        imported = 0
        for hydrated in entries or []:
            rid = (hydrated.get('id_riesgo') or '').strip()
            if not rid:
                continue
            if any(r.id_var.get().strip() == rid for r in self.risk_frames):
                log_event("validacion", f"Riesgo duplicado {rid} en importación", self.logs)
                continue
            self.add_risk()
            rf = self.risk_frames[-1]
            rf.id_var.set(rid)
            rf.lider_var.set((hydrated.get('lider') or '').strip())
            rf.descripcion_var.set((hydrated.get('descripcion') or '').strip())
            crit = (hydrated.get('criticidad') or '').strip()
            if crit in CRITICIDAD_LIST:
                rf.criticidad_var.set(crit)
            rf.exposicion_var.set((hydrated.get('exposicion_residual') or '').strip())
            rf.planes_var.set((hydrated.get('planes_accion') or '').strip())
            imported += 1
        self._notify_dataset_changed(summary_sections="riesgos")
        log_event("navegacion", "Riesgos importados desde CSV", self.logs)
        if imported:
            messagebox.showinfo("Importación completa", "Riesgos importados correctamente.")
        else:
            messagebox.showwarning("Sin cambios", "No se añadieron riesgos nuevos.")

    def _apply_norm_import_payload(self, entries):
        imported = 0
        for hydrated in entries or []:
            nid = (hydrated.get('id_norma') or '').strip()
            if not nid:
                continue
            if any(n.id_var.get().strip() == nid for n in self.norm_frames):
                log_event("validacion", f"Norma duplicada {nid} en importación", self.logs)
                continue
            self.add_norm()
            nf = self.norm_frames[-1]
            nf.id_var.set(nid)
            nf.descripcion_var.set((hydrated.get('descripcion') or '').strip())
            nf.fecha_var.set((hydrated.get('fecha_vigencia') or '').strip())
            imported += 1
        self._notify_dataset_changed(summary_sections="normas")
        log_event("navegacion", "Normas importadas desde CSV", self.logs)
        if imported:
            messagebox.showinfo("Importación completa", "Normas importadas correctamente.")
        else:
            messagebox.showwarning("Sin cambios", "No se añadieron normas nuevas.")

    def _apply_claim_import_payload(self, entries):
        imported = 0
        missing_products = []
        for entry in entries or []:
            hydrated = entry.get('row', {})
            found = entry.get('found', False)
            product_id = (hydrated.get('id_producto') or '').strip()
            if not product_id:
                continue
            product_frame = self._find_product_frame(product_id)
            new_product = False
            if not product_frame:
                product_frame = self._obtain_product_slot_for_import()
                new_product = True
            client_id = (hydrated.get('id_cliente') or '').strip()
            if client_id:
                client_details, _ = self._hydrate_row_from_details({'id_cliente': client_id}, 'id_cliente', CLIENT_ID_ALIASES)
                self._ensure_client_exists(client_id, client_details)
            if new_product:
                self._populate_product_frame_from_row(product_frame, hydrated)
            if product_frame:
                self._trigger_import_id_refresh(
                    product_frame,
                    product_id,
                    preserve_existing=False,
                )
            claim_payload = {
                'id_reclamo': (hydrated.get('id_reclamo') or '').strip(),
                'nombre_analitica': (hydrated.get('nombre_analitica') or '').strip(),
                'codigo_analitica': (hydrated.get('codigo_analitica') or '').strip(),
            }
            if not any(claim_payload.values()):
                continue
            target = product_frame.find_claim_by_id(claim_payload['id_reclamo'])
            if not target:
                target = product_frame.obtain_claim_slot()
            target.set_data(claim_payload)
            self._sync_product_lookup_claim_fields(product_frame, product_id)
            product_frame.persist_lookup_snapshot()
            imported += 1
            if not found and 'id_producto' in self.detail_catalogs:
                missing_products.append(product_id)
        self._notify_dataset_changed(summary_sections="reclamos")
        log_event("navegacion", "Reclamos importados desde CSV", self.logs)
        if imported:
            self.sync_main_form_after_import("reclamos")
            messagebox.showinfo("Importación completa", "Reclamos importados correctamente.")
        else:
            messagebox.showwarning("Sin cambios", "Ningún reclamo se pudo vincular a productos existentes.")
        self._report_missing_detail_ids("productos", missing_products)

    def import_clients(self, filename=None):
        """Importa clientes desde un archivo CSV y los añade a la lista."""

        filename = filename or self._select_csv_file("clientes", "Seleccionar CSV de clientes")
        if not filename:
            messagebox.showwarning("Sin archivo", "No se seleccionó un CSV para clientes ni se encontró el ejemplo.")
            return
        def worker():
            payload = []
            for row in iter_massive_csv_rows(filename):
                hydrated, found = self._hydrate_row_from_details(row, 'id_cliente', CLIENT_ID_ALIASES)
                id_cliente = (hydrated.get('id_cliente') or '').strip()
                if not id_cliente:
                    continue
                payload.append({'row': hydrated, 'found': found})
            return payload

        self._start_background_import(
            "clientes",
            getattr(self, 'import_clients_button', None),
            worker,
            self._apply_client_import_payload,
            "No se pudo importar clientes",
        )

    def import_team_members(self, filename=None):
        """Importa colaboradores desde un archivo CSV y los añade a la lista."""

        filename = filename or self._select_csv_file("colaboradores", "Seleccionar CSV de colaboradores")
        if not filename:
            messagebox.showwarning("Sin archivo", "No hay CSV para colaboradores disponible.")
            return
        def worker():
            payload = []
            for row in iter_massive_csv_rows(filename):
                hydrated, found = self._hydrate_row_from_details(row, 'id_colaborador', TEAM_ID_ALIASES)
                collaborator_id = (hydrated.get('id_colaborador') or '').strip()
                if not collaborator_id:
                    continue
                payload.append({'row': hydrated, 'found': found})
            return payload

        self._start_background_import(
            "colaboradores",
            getattr(self, 'import_team_button', None),
            worker,
            self._apply_team_import_payload,
            "No se pudo importar colaboradores",
        )

    def import_products(self, filename=None):
        """Importa productos desde un archivo CSV y los añade a la lista."""

        filename = filename or self._select_csv_file("productos", "Seleccionar CSV de productos")
        if not filename:
            messagebox.showwarning("Sin archivo", "No se seleccionó CSV de productos ni se encontró el ejemplo.")
            return
        def worker():
            payload = []
            for row in iter_massive_csv_rows(filename):
                hydrated, found = self._hydrate_row_from_details(row, 'id_producto', PRODUCT_ID_ALIASES)
                product_id = (hydrated.get('id_producto') or '').strip()
                if not product_id:
                    continue
                payload.append({'row': hydrated, 'found': found})
            return payload

        self._start_background_import(
            "productos",
            getattr(self, 'import_products_button', None),
            worker,
            self._apply_product_import_payload,
            "No se pudo importar productos",
        )

    # ---------------------------------------------------------------------
    # Autoguardado y carga

    def save_auto(self, data=None):
        """Guarda automáticamente el estado actual en un archivo JSON."""

        dataset = data or self.gather_data()
        try:
            with open(AUTOSAVE_FILE, 'w', encoding="utf-8") as f:
                json.dump(dataset, f, ensure_ascii=False, indent=2)
        except Exception as ex:
            log_event("validacion", f"Error guardando autosave: {ex}", self.logs)
        self._schedule_summary_refresh(data=dataset)
        return dataset

    def load_autosave(self):
        """Carga el estado guardado automáticamente si el archivo existe."""
        if os.path.exists(AUTOSAVE_FILE):
            try:
                with open(AUTOSAVE_FILE, 'r', encoding="utf-8") as f:
                    data = json.load(f)
                self.populate_from_data(data)
                log_event("navegacion", "Se cargó el autosave", self.logs)
            except Exception as ex:
                log_event("validacion", f"Error cargando autosave: {ex}", self.logs)

    def _handle_window_close(self):
        self.flush_autosave()
        self.flush_logs_now(reschedule=False)
        self.root.destroy()

    def _schedule_log_flush(self) -> None:
        if not self._has_log_targets():
            return
        if self._log_flush_job_id is not None:
            return
        try:
            self._log_flush_job_id = self.root.after(
                self.LOG_FLUSH_INTERVAL_MS,
                self._on_log_flush_timer,
            )
        except tk.TclError:
            self._log_flush_job_id = None
            self._flush_log_queue_to_disk()

    def _on_log_flush_timer(self) -> None:
        self._log_flush_job_id = None
        self._flush_log_queue_to_disk()
        self._schedule_log_flush()

    def _cancel_log_flush_job(self) -> None:
        if self._log_flush_job_id is None:
            return
        try:
            self.root.after_cancel(self._log_flush_job_id)
        except tk.TclError:
            pass
        self._log_flush_job_id = None

    def flush_logs_now(self, reschedule: bool = True) -> None:
        self._cancel_log_flush_job()
        self._flush_log_queue_to_disk()
        if reschedule:
            self._schedule_log_flush()

    def _flush_log_queue_to_disk(self) -> None:
        if not self._has_log_targets():
            return
        rows = drain_log_queue()
        if not rows:
            return
        errors = []
        if STORE_LOGS_LOCALLY and LOGS_FILE:
            try:
                self._append_log_rows(
                    LOGS_FILE,
                    rows,
                    track_attr="_log_file_initialized",
                )
            except OSError as exc:
                errors.append((LOGS_FILE, exc))
        external_path = self._resolve_external_log_target()
        if external_path:
            try:
                self._append_log_rows(
                    external_path,
                    rows,
                    track_attr="_external_log_file_initialized",
                )
            except OSError as exc:
                errors.append((external_path, exc))
        if errors:
            messages = [f"No se pudo escribir el log en {path}: {exc}" for path, exc in errors]
            for message in messages:
                log_event("validacion", message, self.logs)
            if not getattr(self, '_suppress_messagebox', False):
                messagebox.showwarning("Registro no guardado", "\n".join(messages))

    def _append_log_rows(self, file_path: str, rows, *, track_attr: Optional[str] = None) -> None:
        target = Path(file_path)
        if target.parent:
            target.parent.mkdir(parents=True, exist_ok=True)
        file_exists = target.exists()
        if track_attr:
            file_exists = getattr(self, track_attr, False) or file_exists
        with target.open('a', newline='', encoding='utf-8') as file_handle:
            writer = csv.DictWriter(file_handle, fieldnames=['timestamp', 'tipo', 'mensaje'])
            if not file_exists:
                writer.writeheader()
            writer.writerows(rows)
        if track_attr:
            setattr(self, track_attr, True)

    def _log_navigation(self, message: str, autosave: bool = False) -> None:
        log_event("navegacion", message, self.logs)
        if autosave:
            self.request_autosave()

    def _log_navigation_change(self, message: str) -> None:
        self._log_navigation(message, autosave=True)

    def request_autosave(self) -> None:
        self._autosave_dirty = True
        if self._autosave_job_id is not None:
            return
        try:
            self._autosave_job_id = self.root.after(
                self.AUTOSAVE_DELAY_MS,
                self._perform_debounced_autosave,
            )
        except tk.TclError:
            self._autosave_job_id = None
            self.flush_autosave()

    def _perform_debounced_autosave(self) -> None:
        self._autosave_job_id = None
        if not self._autosave_dirty:
            return
        self._autosave_dirty = False
        dataset = self.save_auto()
        self.save_temp_version(dataset)

    def flush_autosave(self) -> None:
        if self._autosave_job_id is not None:
            try:
                self.root.after_cancel(self._autosave_job_id)
            except tk.TclError:
                pass
            self._autosave_job_id = None
        if self._autosave_dirty:
            self._autosave_dirty = False
            dataset = self.save_auto()
            self.save_temp_version(dataset)

    def load_version_dialog(self):
        """Abre un diálogo para cargar una versión previa del formulario.

        Esta función solicita al usuario que seleccione un archivo ``.json``
        previamente generado mediante la opción «Guardar y enviar» o un
        autosave temporal.  Tras seleccionar el archivo, se invoca
        ``load_version`` con la ruta escogida para restaurar el estado
        completo del formulario.  Si el usuario cancela la operación no
        se realizan cambios.

        """
        filename = filedialog.askopenfilename(title="Seleccionar versión JSON", filetypes=[("JSON Files", "*.json")])
        if not filename:
            return
        try:
            with open(filename, 'r', encoding="utf-8") as f:
                data = json.load(f)
            self.populate_from_data(data)
            log_event("navegacion", f"Se cargó versión desde {filename}", self.logs)
            messagebox.showinfo("Versión cargada", "La versión se cargó correctamente.")
        except Exception as ex:
            messagebox.showerror("Error", f"No se pudo cargar la versión: {ex}")

    def _clear_case_state(self, *, save_autosave: bool = True) -> None:
        """Elimina los datos cargados y restablece los frames dinámicos."""

        # Limpiar campos del caso
        self.id_caso_var.set("")
        self.tipo_informe_var.set(TIPO_INFORME_LIST[0])
        self.cat_caso1_var.set(list(TAXONOMIA.keys())[0])
        self.on_case_cat1_change()
        self.canal_caso_var.set(CANAL_LIST[0])
        self.proceso_caso_var.set(PROCESO_LIST[0])
        # Vaciar listas dinámicas
        for cf in self.client_frames:
            cf.frame.destroy()
        self.client_frames.clear()
        for tm in self.team_frames:
            tm.frame.destroy()
        self.team_frames.clear()
        for pr in self.product_frames:
            pr.frame.destroy()
        self.product_frames.clear()
        for rf in self.risk_frames:
            rf.frame.destroy()
        self.risk_frames.clear()
        for nf in self.norm_frames:
            nf.frame.destroy()
        self.norm_frames.clear()
        self.next_risk_number = 1
        self._rebuild_frame_id_indexes()
        # Reiniciar bitácora antes de poblar los frames por defecto
        self.logs.clear()
        drain_log_queue()
        # Volver a crear uno por cada sección donde corresponde
        self.add_client()
        self.add_team()
        self.add_risk()
        # Limpiar análisis
        for widget in self._analysis_text_widgets().values():
            self._set_text_content(widget, "")
        if save_autosave:
            self.save_auto()

    def _reset_form_state(self, confirm=True, save_autosave=True):
        """Restablece el formulario a su estado inicial.

        Args:
            confirm: Si es ``True`` solicita confirmación al usuario antes de
                continuar.
            save_autosave: Si es ``True`` persiste inmediatamente el estado
                vacío mediante ``save_auto``.

        Returns:
            ``True`` si el formulario se restableció, ``False`` si el usuario
            canceló la acción.
        """

        if confirm:
            confirmed = messagebox.askyesno(
                "Confirmar",
                "¿Desea borrar todos los datos? Esta acción no se puede deshacer.",
            )
            if not confirmed:
                return False
        self._clear_case_state(save_autosave=save_autosave)
        return True

    def clear_all(self, notify=True):
        """Elimina todos los datos actuales y restablece el formulario."""

        if not self._reset_form_state(confirm=True, save_autosave=True):
            return
        log_event("navegacion", "Se borraron todos los datos", self.logs)
        if notify:
            messagebox.showinfo("Datos borrados", "Todos los datos han sido borrados.")

    # ---------------------------------------------------------------------
    # Recolección y población de datos

    def gather_data(self):
        """Reúne todos los datos del formulario en una estructura de diccionarios."""
        data = {}
        data['caso'] = {
            "id_caso": self.id_caso_var.get().strip(),
            "tipo_informe": self.tipo_informe_var.get(),
            "categoria1": self.cat_caso1_var.get(),
            "categoria2": self.cat_caso2_var.get(),
            "modalidad": self.mod_caso_var.get(),
            "canal": self.canal_caso_var.get(),
            "proceso": self.proceso_caso_var.get(),
        }
        data['clientes'] = [c.get_data() for c in self.client_frames]
        data['colaboradores'] = [t.get_data() for t in self.team_frames]
        productos = []
        reclamos = []
        involucs = []
        for p in self.product_frames:
            prod_data = p.get_data()
            productos.append(prod_data['producto'])
            # Reclamos
            for claim in prod_data['reclamos']:
                if not any(claim.values()):
                    continue
                reclamos.append({
                    "id_reclamo": claim['id_reclamo'],
                    "id_caso": "",  # se añade al exportar
                    "id_producto": prod_data['producto']['id_producto'],
                    "nombre_analitica": claim['nombre_analitica'],
                    "codigo_analitica": claim['codigo_analitica'],
                })
            # Involucramientos
            for inv in prod_data['asignaciones']:
                involucs.append({
                    "id_producto": prod_data['producto']['id_producto'],
                    "id_caso": "",  # se completará al exportar
                    "id_colaborador": inv['id_colaborador'],
                    "monto_asignado": inv['monto_asignado'],
                })
        data['productos'] = productos
        data['reclamos'] = reclamos
        data['involucramientos'] = involucs
        data['riesgos'] = [r.get_data() for r in self.risk_frames]
        normas = []
        for n in self.norm_frames:
            norm_data = n.get_data()
            if not norm_data:
                continue
            normas.append(norm_data)
        data['normas'] = normas
        analysis_widgets = self._analysis_text_widgets()
        data['analisis'] = {
            "antecedentes": self._get_text_content(analysis_widgets["antecedentes"]),
            "modus_operandi": self._get_text_content(analysis_widgets["modus_operandi"]),
            "hallazgos": self._get_text_content(analysis_widgets["hallazgos"]),
            "descargos": self._get_text_content(analysis_widgets["descargos"]),
            "conclusiones": self._get_text_content(analysis_widgets["conclusiones"]),
            "recomendaciones": self._get_text_content(analysis_widgets["recomendaciones"]),
        }
        return data

    def _normalize_export_amount_strings(self, products):
        if not products:
            return
        for product in products:
            if not isinstance(product, dict):
                continue
            for field_name, _var_attr, _label, allow_blank, _ in PRODUCT_MONEY_SPECS:
                if not allow_blank:
                    continue
                raw_value = product.get(field_name)
                if isinstance(raw_value, str):
                    text = raw_value.strip()
                elif raw_value is None:
                    text = ""
                else:
                    text = str(raw_value).strip()
                if text:
                    continue
                product[field_name] = "0.00"

    def populate_from_data(self, data):
        """Puebla el formulario con datos previamente guardados."""
        # Limpiar primero sin confirmar ni sobrescribir el autosave
        self._clear_case_state(save_autosave=False)
        # Datos de caso
        def _set_dropdown_value(var, value, valid_values):
            """Establece el valor de un combobox solo si está en el catálogo."""

            normalized = value.strip() if isinstance(value, str) else value
            if normalized and normalized in valid_values:
                var.set(normalized)
            else:
                var.set('')

        caso = data.get('caso', {})
        self.id_caso_var.set(caso.get('id_caso', ''))
        if caso.get('tipo_informe') in TIPO_INFORME_LIST:
            self.tipo_informe_var.set(caso.get('tipo_informe'))
        if caso.get('categoria1') in TAXONOMIA:
            self.cat_caso1_var.set(caso.get('categoria1'))
            self.on_case_cat1_change()
            cat2_list = list(TAXONOMIA[self.cat_caso1_var.get()].keys())
            if caso.get('categoria2') in cat2_list:
                self.cat_caso2_var.set(caso.get('categoria2'))
                self.on_case_cat2_change()
                mod_list = TAXONOMIA[self.cat_caso1_var.get()][self.cat_caso2_var.get()]
                if caso.get('modalidad') in mod_list:
                    self.mod_caso_var.set(caso.get('modalidad'))
        _set_dropdown_value(self.canal_caso_var, caso.get('canal'), CANAL_LIST)
        _set_dropdown_value(self.proceso_caso_var, caso.get('proceso'), PROCESO_LIST)
        # Clientes
        for i, cliente in enumerate(data.get('clientes', [])):
            if i >= len(self.client_frames):
                self.add_client()
            cl = self.client_frames[i]
            cl.tipo_id_var.set(cliente.get('tipo_id', ''))
            cl.id_var.set(cliente.get('id_cliente', ''))
            cl.flag_var.set(cliente.get('flag', ''))
            cl.telefonos_var.set(cliente.get('telefonos', ''))
            cl.correos_var.set(cliente.get('correos', ''))
            cl.direcciones_var.set(cliente.get('direcciones', ''))
            cl.set_accionado_from_text(cliente.get('accionado', ''))
        # Colaboradores
        for i, col in enumerate(data.get('colaboradores', [])):
            if i >= len(self.team_frames):
                self.add_team()
            tm = self.team_frames[i]
            tm.id_var.set(col.get('id_colaborador', ''))
            tm.flag_var.set(col.get('flag', ''))
            tm.division_var.set(col.get('division', ''))
            tm.area_var.set(col.get('area', ''))
            tm.servicio_var.set(col.get('servicio', ''))
            tm.puesto_var.set(col.get('puesto', ''))
            tm.nombre_agencia_var.set(col.get('nombre_agencia', ''))
            tm.codigo_agencia_var.set(col.get('codigo_agencia', ''))
            tm.tipo_falta_var.set(col.get('tipo_falta', ''))
            tm.tipo_sancion_var.set(col.get('tipo_sancion', ''))
        # Productos y sus reclamos e involuc
        claims_map = {}
        for rec in data.get('reclamos', []):
            pid = (rec.get('id_producto') or '').strip()
            if not pid:
                continue
            claims_map.setdefault(pid, []).append(
                {
                    'id_reclamo': (rec.get('id_reclamo') or '').strip(),
                    'nombre_analitica': (rec.get('nombre_analitica') or '').strip(),
                    'codigo_analitica': (rec.get('codigo_analitica') or '').strip(),
                }
            )
        for i, prod in enumerate(data.get('productos', [])):
            if i >= len(self.product_frames):
                self.add_product(initialize_rows=False)
            pframe = self.product_frames[i]
            pframe.id_var.set(prod.get('id_producto', ''))
            pframe.client_var.set(prod.get('id_cliente', ''))
            cat1 = prod.get('categoria1')
            if cat1 in TAXONOMIA:
                pframe.cat1_var.set(cat1)
                pframe.on_cat1_change()
                cat2 = prod.get('categoria2')
                if cat2 in TAXONOMIA[cat1]:
                    pframe.cat2_var.set(cat2)
                    pframe.on_cat2_change()
                    mod = prod.get('modalidad')
                    if mod in TAXONOMIA[cat1][cat2]:
                        pframe.mod_var.set(mod)
            _set_dropdown_value(pframe.canal_var, prod.get('canal'), CANAL_LIST)
            _set_dropdown_value(pframe.proceso_var, prod.get('proceso'), PROCESO_LIST)
            pframe.fecha_oc_var.set(prod.get('fecha_ocurrencia', ''))
            pframe.fecha_desc_var.set(prod.get('fecha_descubrimiento', ''))
            pframe.monto_inv_var.set(prod.get('monto_investigado', ''))
            _set_dropdown_value(pframe.moneda_var, prod.get('tipo_moneda'), TIPO_MONEDA_LIST)
            pframe.monto_perdida_var.set(prod.get('monto_perdida_fraude', ''))
            pframe.monto_falla_var.set(prod.get('monto_falla_procesos', ''))
            pframe.monto_cont_var.set(prod.get('monto_contingencia', ''))
            pframe.monto_rec_var.set(prod.get('monto_recuperado', ''))
            pframe.monto_pago_var.set(prod.get('monto_pago_deuda', ''))
            tipo_producto = prod.get('tipo_producto')
            if tipo_producto in TIPO_PRODUCTO_LIST:
                pframe.tipo_prod_var.set(tipo_producto)
            pframe.set_claims_from_data(claims_map.get(pframe.id_var.get().strip(), []))
        # Involucramientos
        involvement_map = {}
        for inv in data.get('involucramientos', []):
            pid = (inv.get('id_producto') or '').strip()
            if not pid:
                continue
            involvement_map.setdefault(pid, []).append(inv)

        for pframe in self.product_frames:
            pid = pframe.id_var.get().strip()
            if pid not in involvement_map:
                continue
            pframe.clear_involvements()
            for inv in involvement_map[pid]:
                assign = pframe.add_involvement()
                assign.team_var.set(inv.get('id_colaborador', ''))
                assign.monto_var.set(inv.get('monto_asignado', ''))
        # Riesgos
        for i, risk in enumerate(data.get('riesgos', [])):
            if i >= len(self.risk_frames):
                self.add_risk()
            rf = self.risk_frames[i]
            rf.id_var.set(risk.get('id_riesgo', ''))
            rf.lider_var.set(risk.get('lider', ''))
            rf.descripcion_var.set(risk.get('descripcion', ''))
            rf.criticidad_var.set(risk.get('criticidad', CRITICIDAD_LIST[0]))
            rf.exposicion_var.set(risk.get('exposicion_residual', ''))
            rf.planes_var.set(risk.get('planes_accion', ''))
        # Normas
        for i, norm in enumerate(data.get('normas', [])):
            if i >= len(self.norm_frames):
                self.add_norm()
            nf = self.norm_frames[i]
            nf.id_var.set(norm.get('id_norma', ''))
            nf.descripcion_var.set(norm.get('descripcion', ''))
            nf.fecha_var.set(norm.get('fecha_vigencia', ''))
        # Analisis
        analisis = data.get('analisis', {})
        analysis_widgets = self._analysis_text_widgets()
        self._set_text_content(analysis_widgets['antecedentes'], analisis.get('antecedentes', ''))
        self._set_text_content(analysis_widgets['modus_operandi'], analisis.get('modus_operandi', ''))
        self._set_text_content(analysis_widgets['hallazgos'], analisis.get('hallazgos', ''))
        self._set_text_content(analysis_widgets['descargos'], analisis.get('descargos', ''))
        self._set_text_content(analysis_widgets['conclusiones'], analisis.get('conclusiones', ''))
        self._set_text_content(analysis_widgets['recomendaciones'], analisis.get('recomendaciones', ''))
        self._rebuild_frame_id_indexes()
        self._schedule_summary_refresh(data=data)

    # ---------------------------------------------------------------------
    # Validación de reglas de negocio

    def validate_data(self):
        """Valida los datos del formulario y retorna errores y advertencias."""
        errors = []
        warnings = []

        def _report_catalog_error(message):
            """Acumula los errores de catálogo para notificarlos en bloque."""

            errors.append(message)

        def _validate_product_catalog_field(value, label, catalog, catalog_label, product_label):
            text = (value or '').strip()
            message = validate_required_text(text, label)
            if message:
                return f"Producto {product_label}: {message}"
            if text not in catalog:
                catalog_message = (
                    f"Producto {product_label}: El {catalog_label} '{text}' no está en el catálogo CM."
                )
                _report_catalog_error(catalog_message)
                return None
            return None

        def _validate_product_taxonomy(producto, product_label):
            """Valida que categoría 1, 2 y modalidad del producto existan en TAXONOMIA."""

            messages = []
            cat1 = (producto.get('categoria1') or '').strip()
            cat2 = (producto.get('categoria2') or '').strip()
            modalidad = (producto.get('modalidad') or '').strip()

            cat1_message = validate_required_text(cat1, "la categoría 1 del producto")
            if cat1_message:
                messages.append(f"Producto {product_label}: {cat1_message}")
            elif cat1 not in TAXONOMIA:
                messages.append(
                    f"Producto {product_label}: La categoría 1 '{cat1}' no está en el catálogo CM."
                )
            cat1_valid = bool(cat1) and cat1 in TAXONOMIA

            cat2_message = validate_required_text(cat2, "la categoría 2 del producto")
            if cat2_message:
                messages.append(f"Producto {product_label}: {cat2_message}")
                cat2_valid = False
            else:
                if not cat1_valid:
                    messages.append(
                        f"Producto {product_label}: La categoría 2 '{cat2}' no puede validarse porque la categoría 1 es inválida."
                    )
                    cat2_valid = False
                elif cat2 not in TAXONOMIA[cat1]:
                    messages.append(
                        f"Producto {product_label}: La categoría 2 '{cat2}' no pertenece a la categoría 1 '{cat1}' del catálogo CM."
                    )
                    cat2_valid = False
                else:
                    cat2_valid = True

            mod_message = validate_required_text(modalidad, "la modalidad del producto")
            if mod_message:
                messages.append(f"Producto {product_label}: {mod_message}")
            else:
                if not (cat1_valid and cat2_valid):
                    messages.append(
                        f"Producto {product_label}: La modalidad '{modalidad}' no puede validarse porque las categorías registradas son inválidas."
                    )
                elif modalidad not in TAXONOMIA[cat1][cat2]:
                    messages.append(
                        f"Producto {product_label}: La modalidad '{modalidad}' no pertenece a la categoría 1 '{cat1}' y categoría 2 '{cat2}' del catálogo CM."
                    )
            return messages

        def _validate_team_catalog_value(value, label, catalog, collaborator_idx):
            text = (value or '').strip()
            if not text:
                errors.append(f"Colaborador {collaborator_idx}: Debe seleccionar {label}.")
            elif text not in catalog:
                errors.append(
                    f"Colaborador {collaborator_idx}: El {label} '{text}' no está en el catálogo CM."
                )
        # Validar número de caso
        id_caso = self.id_caso_var.get().strip()
        normalized_case_id = self._normalize_identifier(id_caso)
        case_message = validate_case_id(id_caso)
        if case_message:
            errors.append(case_message)
        # Validar campos obligatorios del caso antes de validar entidades hijas
        tipo_informe_value = (self.tipo_informe_var.get() or '').strip()
        tipo_message = validate_required_text(tipo_informe_value, "el tipo de informe")
        if tipo_message:
            errors.append(tipo_message)
        elif tipo_informe_value not in TIPO_INFORME_LIST:
            _report_catalog_error(
                f"El tipo de informe '{tipo_informe_value}' no está en el catálogo CM."
            )
        cat1_value = (self.cat_caso1_var.get() or '').strip()
        cat1_message = validate_required_text(cat1_value, "la categoría nivel 1")
        cat1_valid = False
        if cat1_message:
            errors.append(cat1_message)
        elif cat1_value not in TAXONOMIA:
            _report_catalog_error(
                f"La categoría nivel 1 '{cat1_value}' no está en el catálogo CM."
            )
        else:
            cat1_valid = True
        cat2_value = (self.cat_caso2_var.get() or '').strip()
        cat2_message = validate_required_text(cat2_value, "la categoría nivel 2")
        cat2_valid = False
        if cat2_message:
            errors.append(cat2_message)
        elif cat1_valid:
            if cat2_value not in TAXONOMIA[cat1_value]:
                _report_catalog_error(
                    f"La categoría nivel 2 '{cat2_value}' no está dentro de la categoría '{cat1_value}' del catálogo CM."
                )
            else:
                cat2_valid = True
        mod_value = (self.mod_caso_var.get() or '').strip()
        mod_message = validate_required_text(mod_value, "la modalidad del caso")
        if mod_message:
            errors.append(mod_message)
        elif cat1_valid and cat2_valid:
            modalidades = TAXONOMIA[cat1_value][cat2_value]
            if mod_value not in modalidades:
                _report_catalog_error(
                    f"La modalidad '{mod_value}' no existe dentro de la categoría '{cat1_value}'/'{cat2_value}' del catálogo CM."
                )
        canal_value = (self.canal_caso_var.get() or '').strip()
        canal_message = validate_required_text(canal_value, "el canal del caso")
        if canal_message:
            errors.append(canal_message)
        elif canal_value not in CANAL_LIST:
            _report_catalog_error(
                f"El canal del caso '{canal_value}' no está en el catálogo CM."
            )
        proceso_value = (self.proceso_caso_var.get() or '').strip()
        proceso_message = validate_required_text(proceso_value, "el proceso impactado")
        if proceso_message:
            errors.append(proceso_message)
        elif proceso_value not in PROCESO_LIST:
            _report_catalog_error(
                f"El proceso del caso '{proceso_value}' no está en el catálogo CM."
            )
        # Validar IDs de clientes
        client_id_occurrences = {}
        for idx, cframe in enumerate(self.client_frames, start=1):
            tipo_id_value = (cframe.tipo_id_var.get() or "").strip()
            tipo_message = validate_required_text(tipo_id_value, "el tipo de ID del cliente")
            if tipo_message:
                errors.append(f"Cliente {idx}: {tipo_message}")
            else:
                if tipo_id_value not in TIPO_ID_LIST:
                    errors.append(
                        f"Cliente {idx}: El tipo de ID '{tipo_id_value}' no está en el catálogo CM."
                    )
                else:
                    client_id_value = (cframe.id_var.get() or "").strip()
                    message = validate_client_id(tipo_id_value, client_id_value)
                    if message:
                        errors.append(f"Cliente {idx}: {message}")
                    normalized_client_id = self._normalize_identifier(client_id_value)
                    if normalized_client_id:
                        occurrences = client_id_occurrences.setdefault(normalized_client_id, [])
                        occurrences.append((idx, client_id_value or normalized_client_id))

            flag_var = getattr(cframe, "flag_var", None)
            if flag_var is not None:
                flag_value = (flag_var.get() or "").strip()
                flag_message = validate_required_text(flag_value, "el flag del cliente")
                if flag_message:
                    errors.append(f"Cliente {idx}: {flag_message}")
                elif flag_value not in FLAG_CLIENTE_LIST:
                    errors.append(
                        f"Cliente {idx}: El flag de cliente '{flag_value}' no está en el catálogo CM."
                    )
            if hasattr(cframe, 'telefonos_var'):
                phone_value = cframe.telefonos_var.get()
                phone_required = validate_required_text(
                    phone_value, "los teléfonos del cliente"
                )
                if phone_required:
                    errors.append(f"Cliente {idx}: {phone_required}")
                else:
                    phone_message = validate_phone_list(
                        phone_value, "los teléfonos del cliente"
                    )
                    if phone_message:
                        errors.append(f"Cliente {idx}: {phone_message}")
            if hasattr(cframe, 'correos_var'):
                email_value = cframe.correos_var.get()
                email_required = validate_required_text(
                    email_value, "los correos del cliente"
                )
                if email_required:
                    errors.append(f"Cliente {idx}: {email_required}")
                else:
                    email_message = validate_email_list(
                        email_value, "los correos del cliente"
                    )
                    if email_message:
                        errors.append(f"Cliente {idx}: {email_message}")
            if hasattr(cframe, 'accionado_var'):
                accionado_message = validate_multi_selection(
                    cframe.accionado_var.get(), "Accionado"
                )
                if accionado_message:
                    errors.append(f"Cliente {idx}: {accionado_message}")
        for normalized_client_id, occurrences in client_id_occurrences.items():
            if len(occurrences) > 1:
                formatted_positions = ", ".join(str(pos) for pos, _ in occurrences)
                display_value = occurrences[0][1] or normalized_client_id
                errors.append(
                    (
                        f"El ID de cliente {display_value} está duplicado en los clientes {formatted_positions}. "
                        "Cada cliente debe tener un ID único."
                    )
                )
        # Validar duplicidad del key técnico (caso, producto, cliente, colaborador, fecha ocurrencia, reclamo)
        key_set = set()
        product_client_map = {}
        total_investigado = Decimal('0')
        total_componentes = Decimal('0')
        normalized_amounts = []
        team_id_occurrences = {}
        for idx, tm in enumerate(self.team_frames, start=1):
            current_idx = idx
            team_id_value = (tm.id_var.get() or "").strip()
            tm_id_message = validate_team_member_id(team_id_value)
            if tm_id_message:
                errors.append(f"Colaborador {current_idx}: {tm_id_message}")
            normalized_team_id = self._normalize_identifier(team_id_value)
            if normalized_team_id:
                team_id_occurrences.setdefault(normalized_team_id, []).append(current_idx)
            agency_message = validate_agency_code(tm.codigo_agencia_var.get(), allow_blank=True)
            if agency_message:
                errors.append(f"Colaborador {current_idx}: {agency_message}")
            flag_value = (tm.flag_var.get() if hasattr(tm, 'flag_var') else '').strip()
            flag_message = validate_required_text(flag_value, "el flag del colaborador")
            if flag_message:
                errors.append(f"Colaborador {current_idx}: {flag_message}")
            elif flag_value not in FLAG_COLABORADOR_LIST:
                errors.append(
                    f"Colaborador {current_idx}: El flag del colaborador '{flag_value}' no está en el catálogo CM."
                )
            falta_value = tm.tipo_falta_var.get() if hasattr(tm, 'tipo_falta_var') else ''
            sancion_value = tm.tipo_sancion_var.get() if hasattr(tm, 'tipo_sancion_var') else ''
            _validate_team_catalog_value(
                falta_value,
                "el tipo de falta del colaborador",
                TIPO_FALTA_LIST,
                current_idx,
            )
            _validate_team_catalog_value(
                sancion_value,
                "el tipo de sanción del colaborador",
                TIPO_SANCION_LIST,
                current_idx,
            )
            division_value = (tm.division_var.get() if hasattr(tm, 'division_var') else '').strip()
            area_value = (tm.area_var.get() if hasattr(tm, 'area_var') else '').strip()
            division_norm = division_value.lower().replace('á', 'a').replace('é', 'e').replace('ó', 'o')
            area_norm = area_value.lower().replace('á', 'a').replace('é', 'e').replace('ó', 'o')
            needs_agency = (
                'dca' in division_norm or 'canales de atencion' in division_norm
            ) and ('area comercial' in area_norm)
            if needs_agency:
                nombre_agencia = (
                    tm.nombre_agencia_var.get().strip()
                    if hasattr(tm, 'nombre_agencia_var')
                    else ''
                )
                codigo_agencia = (
                    tm.codigo_agencia_var.get().strip()
                    if hasattr(tm, 'codigo_agencia_var')
                    else ''
                )
                if not nombre_agencia or not codigo_agencia:
                    errors.append(
                        "El colaborador {idx} debe registrar nombre y código de agencia por pertenecer a canales comerciales.".format(
                            idx=current_idx
                        )
                    )
        for collaborator_id, positions in team_id_occurrences.items():
            if len(positions) > 1:
                formatted_positions = ", ".join(str(pos) for pos in positions)
                errors.append(
                    (
                        f"El ID de colaborador {collaborator_id} está duplicado en los colaboradores {formatted_positions}. "
                        "Cada colaborador debe tener un ID único."
                    )
                )
        collaborator_ids = set()
        for tm in self.team_frames:
            normalized_team_id = self._normalize_identifier(tm.id_var.get())
            if normalized_team_id:
                collaborator_ids.add(normalized_team_id)
        for idx, p in enumerate(self.product_frames, start=1):
            prod_data = p.get_data()
            producto = prod_data['producto']
            pid = producto.get('id_producto', '')
            pid_norm = self._normalize_identifier(pid)
            pid_key = pid_norm or pid or ''
            producto_label = pid_norm or pid or f"Producto {idx}"
            pid_message = validate_product_id(p.tipo_prod_var.get(), p.id_var.get())
            if pid_message:
                errors.append(f"Producto {idx}: {pid_message}")
            cid = producto['id_cliente']
            cid_norm = self._normalize_identifier(cid)
            if not cid:
                errors.append(
                    f"Producto {idx}: el cliente vinculado fue eliminado. Selecciona un nuevo titular antes de exportar."
                )
            if pid_key in product_client_map:
                prev_entry = product_client_map[pid_key]
                if prev_entry['client_norm'] != cid_norm:
                    prev_client_display = prev_entry['client_display'] or prev_entry['client_norm'] or 'sin ID'
                    current_client_display = cid or cid_norm or 'sin ID'
                    errors.append(
                        f"El producto {producto_label} está asociado a dos clientes distintos ({prev_client_display} y {current_client_display})."
                    )
                else:
                    errors.append(f"El producto {producto_label} está duplicado en el formulario.")
            else:
                product_client_map[pid_key] = {
                    'client_norm': cid_norm,
                    'client_display': cid,
                }
            for taxonomy_error in _validate_product_taxonomy(producto, producto_label):
                errors.append(taxonomy_error)
            catalog_validations = [
                (
                    producto.get('canal'),
                    "el canal del producto",
                    CANAL_LIST,
                    "canal",
                ),
                (
                    producto.get('proceso'),
                    "el proceso del producto",
                    PROCESO_LIST,
                    "proceso",
                ),
                (
                    producto.get('tipo_moneda'),
                    "la moneda del producto",
                    TIPO_MONEDA_LIST,
                    "tipo de moneda",
                ),
            ]
            for value, label, catalog, catalog_label in catalog_validations:
                catalog_error = _validate_product_catalog_field(
                    value,
                    label,
                    catalog,
                    catalog_label,
                    producto_label,
                )
                if catalog_error:
                    errors.append(catalog_error)
            # For each involvement; if no assignments, use empty string for id_colaborador
            claim_rows = prod_data['reclamos'] or [{'id_reclamo': ''}]
            product_occurrence_date = prod_data['producto'].get('fecha_ocurrencia')
            if not prod_data['asignaciones']:
                for claim in claim_rows:
                    claim_id = (claim.get('id_reclamo') or '').strip()
                    claim_id_norm = self._normalize_identifier(claim_id)
                    key = (
                        normalized_case_id,
                        pid_norm,
                        cid_norm,
                        '',
                        product_occurrence_date,
                        claim_id_norm,
                    )
                    if key in key_set:
                        errors.append(f"Registro duplicado de clave técnica (producto {producto_label})")
                    key_set.add(key)
            for inv_idx, inv in enumerate(prod_data['asignaciones'], start=1):
                collaborator_id = (inv.get('id_colaborador') or '').strip()
                collaborator_norm = self._normalize_identifier(collaborator_id)
                amount_value = (inv.get('monto_asignado') or '').strip()
                amount_label = (
                    f"Monto asignado del colaborador {collaborator_id or f'sin ID ({inv_idx})'} "
                    f"en el producto {producto_label}"
                )
                amount_error, _amount, normalized_amount = validate_money_bounds(
                    amount_value,
                    amount_label,
                )
                if amount_error:
                    errors.append(amount_error)
                else:
                    inv['monto_asignado'] = normalized_amount or amount_value
                if amount_value and not collaborator_id:
                    errors.append(
                        f"Producto {pid}: la asignación {inv_idx} tiene un monto sin colaborador."
                    )
                if collaborator_id and not amount_value:
                    errors.append(
                        f"Producto {pid}: la asignación {inv_idx} tiene un colaborador sin monto."
                    )
                if collaborator_id and collaborator_norm not in collaborator_ids:
                    errors.append(
                        f"Producto {pid}: la asignación {inv_idx} referencia un colaborador eliminado (ID {collaborator_id})."
                    )
                for claim in claim_rows:
                    claim_id = (claim.get('id_reclamo') or '').strip()
                    key = (
                        normalized_case_id,
                        pid_norm,
                        cid_norm,
                        collaborator_norm,
                        product_occurrence_date,
                        self._normalize_identifier(claim_id),
                    )
                    if key in key_set:
                        collaborator_label = collaborator_norm or collaborator_id or 'sin ID'
                        errors.append(
                            f"Registro duplicado de clave técnica (producto {producto_label}, colaborador {collaborator_label})"
                        )
                    key_set.add(key)
        # Validar fechas y montos por producto
        for p in self.product_frames:
            data = p.get_data()
            producto = data['producto']
            tipo_producto = producto.get('tipo_producto', '').strip()
            if not tipo_producto:
                errors.append(f"Producto {producto['id_producto']}: Debe ingresar el tipo de producto.")
            else:
                normalized_tipo = normalize_without_accents(tipo_producto).lower()
                if normalized_tipo not in TIPO_PRODUCTO_NORMALIZED:
                    errors.append(
                        f"Producto {producto['id_producto']}: El tipo de producto '{tipo_producto}' no está en el catálogo."
                    )
            # Fechas
            date_message = validate_product_dates(
                producto.get('id_producto'),
                producto.get('fecha_ocurrencia'),
                producto.get('fecha_descubrimiento'),
            )
            if date_message:
                errors.append(date_message)
            # Montos
            money_values = {}
            money_error = False
            for field, _, label, allow_blank, _ in PRODUCT_MONEY_SPECS:
                message, decimal_value, normalized_text = validate_money_bounds(
                    producto.get(field, ''),
                    f"{label} del producto {producto['id_producto']}",
                    allow_blank=allow_blank,
                )
                if message:
                    errors.append(message)
                    money_error = True
                if not message and normalized_text is not None:
                    producto[field] = normalized_text if normalized_text else ''
                money_values[field] = decimal_value if decimal_value is not None else Decimal('0')
            if money_error:
                continue
            m_inv = money_values['monto_investigado']
            m_perd = money_values['monto_perdida_fraude']
            m_fall = money_values['monto_falla_procesos']
            m_cont = money_values['monto_contingencia']
            m_rec = money_values['monto_recuperado']
            m_pago = money_values['monto_pago_deuda']
            normalized_amounts.append({
                'perdida': m_perd,
                'falla': m_fall,
                'contingencia': m_cont,
            })
            componentes = sum_investigation_components(
                perdida=m_perd,
                falla=m_fall,
                contingencia=m_cont,
                recuperado=m_rec,
            )
            if componentes != m_inv:
                errors.append(
                    f"Las cuatro partidas (pérdida, falla, contingencia y recuperación) deben ser iguales al monto investigado en el producto {producto['id_producto']}"
                )
            if m_rec > m_inv:
                errors.append(
                    f"El monto recuperado no puede superar el monto investigado en el producto {producto['id_producto']}"
                )
            if m_pago > m_inv:
                errors.append(f"El monto pagado de deuda excede el monto investigado en el producto {producto['id_producto']}")
            total_investigado += m_inv
            total_componentes += componentes
            # Reclamo y analíticas
            requiere_reclamo = m_perd > 0 or m_fall > 0 or m_cont > 0
            complete_claim_found = False
            seen_claim_ids = set()
            for claim in data['reclamos'] or []:
                claim_id = (claim.get('id_reclamo') or '').strip()
                normalized_claim_id = self._normalize_identifier(claim_id)
                claim_name = (claim.get('nombre_analitica') or '').strip()
                claim_code = (claim.get('codigo_analitica') or '').strip()
                has_any_value = any([claim_id, claim_name, claim_code])
                if claim_id:
                    if normalized_claim_id in seen_claim_ids:
                        errors.append(
                            f"Producto {producto['id_producto']}: El ID de reclamo {claim_id} está duplicado."
                        )
                    seen_claim_ids.add(normalized_claim_id)
                    reclamo_message = validate_reclamo_id(claim_id)
                    if reclamo_message:
                        errors.append(f"Producto {producto['id_producto']}: {reclamo_message}")
                if claim_code:
                    codigo_message = validate_codigo_analitica(claim_code)
                    if codigo_message:
                        errors.append(f"Producto {producto['id_producto']}: {codigo_message}")
                if has_any_value:
                    if not (claim_id and claim_name and claim_code):
                        errors.append(
                            f"Producto {producto['id_producto']}: El reclamo {claim_id or '(sin ID)'} debe tener ID, nombre y código de analítica."
                        )
                    else:
                        complete_claim_found = True
            if requiere_reclamo and not complete_claim_found:
                errors.append(
                    f"Debe ingresar al menos un reclamo completo en el producto {producto['id_producto']} porque hay montos de pérdida, falla o contingencia"
                )
            # Longitud id_producto
            # Tipo producto vs contingencia
            tipo_prod = normalize_without_accents(producto['tipo_producto']).lower()
            if any(word in tipo_prod for word in ['credito', 'tarjeta']):
                if m_cont != m_inv:
                    errors.append(f"El monto de contingencia debe ser igual al monto investigado en el producto {producto['id_producto']} porque es un crédito o tarjeta")
            # Fraude externo
            if producto['categoria2'] == 'Fraude Externo':
                warnings.append(
                    f"Producto {producto['id_producto']} con categoría 2 'Fraude Externo': verifique la analítica registrada."
                )
        if self.product_frames and total_componentes != total_investigado:
            errors.append(
                "Las cuatro partidas (pérdida, falla, contingencia y recuperación) sumadas en el caso no coinciden con el total investigado."
            )
        # Validar que al menos un producto coincida con categorías del caso
        match_found = False
        for p in self.product_frames:
            pd = p.get_data()['producto']
            if pd['categoria1'] == self.cat_caso1_var.get() and pd['categoria2'] == self.cat_caso2_var.get() and pd['modalidad'] == self.mod_caso_var.get():
                match_found = True
                break
        if not match_found:
            errors.append("Ningún producto coincide con las categorías y modalidad seleccionadas para el caso.")
        # Validar reporte tipo Interno vs pérdidas y sanciones
        if self.tipo_informe_var.get() == 'Interno':
            any_loss = any(
                amounts['perdida'] > Decimal('0')
                or amounts['falla'] > Decimal('0')
                or amounts['contingencia'] > Decimal('0')
                for amounts in normalized_amounts
            )
            any_sanction = any(
                t.tipo_sancion_var.get() not in ('No aplica', '')
                for t in self.team_frames
            )
            if any_loss or any_sanction:
                errors.append("No se puede seleccionar tipo de informe 'Interno' si hay pérdidas, fallas, contingencias o sanciones registradas.")
        # Validar riesgos
        risk_exposure_total = Decimal('0')
        risk_ids = set()
        plan_ids = set()
        for idx, r in enumerate(self.risk_frames, start=1):
            rd = r.get_data()
            rid = rd['id_riesgo']
            risk_message = validate_risk_id(rid)
            if risk_message:
                errors.append(f"Riesgo {idx}: {risk_message}")
            elif rid in risk_ids:
                errors.append(f"ID de riesgo duplicado: {rid}")
            if rid:
                risk_ids.add(rid)
            criticidad_value = (rd.get('criticidad') or '').strip()
            if not criticidad_value:
                errors.append(
                    f"Riesgo {idx}: Debe seleccionar la criticidad del riesgo."
                )
            elif criticidad_value not in CRITICIDAD_LIST:
                errors.append(
                    f"Riesgo {idx}: La criticidad '{criticidad_value}' no está en el catálogo CM."
                )
            # Exposición
            message, exposure_decimal, normalized_text = validate_money_bounds(
                rd['exposicion_residual'],
                f"Exposición residual del riesgo {rid}",
                allow_blank=True,
            )
            if message:
                errors.append(message)
            elif exposure_decimal is not None:
                risk_exposure_total += exposure_decimal
            if (
                not message
                and normalized_text
                and normalized_text != (rd.get('exposicion_residual') or '').strip()
            ):
                if hasattr(r, 'exposicion_var'):
                    r.exposicion_var.set(normalized_text)
                rd['exposicion_residual'] = normalized_text
            # Planes de acción
            for plan in [p.strip() for p in rd['planes_accion'].split(';') if p.strip()]:
                if plan in plan_ids:
                    errors.append(f"Plan de acción {plan} duplicado entre riesgos")
                plan_ids.add(plan)
        self._last_validated_risk_exposure_total = risk_exposure_total
        # Validar normas
        norm_ids = set()
        for idx, n in enumerate(self.norm_frames, start=1):
            nd = n.get_data()
            if not nd:
                continue
            nid = nd['id_norma']
            descripcion = nd['descripcion']
            message = validate_norm_id(nid)
            if message:
                errors.append(f"Norma {idx}: {message}")
            elif nid in norm_ids:
                errors.append(f"ID de norma duplicado: {nid}")
            else:
                norm_ids.add(nid)
            if not descripcion:
                errors.append(f"Norma {idx}: Debe ingresar la descripción de la norma.")
            # Fecha vigencia
            fvig = (nd.get('fecha_vigencia') or '').strip()
            fvig_message = validate_date_text(fvig, "la fecha de vigencia", allow_blank=False)
            if fvig_message:
                errors.append(f"Norma {idx}: {fvig_message}")
            else:
                fv = datetime.strptime(fvig, "%Y-%m-%d")
                if fv > datetime.now():
                    errors.append(f"Fecha de vigencia futura en norma {nid or 'sin ID'}")
        return errors, warnings

    # ---------------------------------------------------------------------
    # Exportación de datos

    def _build_report_context(self, data):
        case = data['caso']
        analysis = data['analisis']
        clients = data['clientes']
        team = data['colaboradores']
        products = data['productos']
        riesgos = data['riesgos']
        normas = data['normas']
        total_inv = sum((parse_decimal_amount(p.get('monto_investigado')) or Decimal('0')) for p in products)
        destinatarios = sorted({
            " - ".join(filter(None, [col.get('division', '').strip(), col.get('area', '').strip(), col.get('servicio', '').strip()]))
            for col in team
            if any([col.get('division'), col.get('area'), col.get('servicio')])
        })
        destinatarios = [d for d in destinatarios if d]
        destinatarios_text = ", ".join(destinatarios) if destinatarios else "Sin divisiones registradas"
        reclamos_por_producto = {}
        for record in data['reclamos']:
            pid = record.get('id_producto')
            if not pid:
                continue
            reclamos_por_producto.setdefault(pid, []).append(record)
        client_rows = [
            [
                f"Cliente {idx}",
                client.get('tipo_id', ''),
                client.get('id_cliente', ''),
                client.get('flag', ''),
                client.get('telefonos', ''),
                client.get('correos', ''),
                client.get('direcciones', ''),
                client.get('accionado', ''),
            ]
            for idx, client in enumerate(clients, start=1)
        ]
        team_rows = [
            [
                f"Colaborador {idx}",
                col.get('id_colaborador', ''),
                col.get('flag', ''),
                col.get('division', ''),
                col.get('area', ''),
                col.get('servicio', ''),
                col.get('puesto', ''),
                col.get('nombre_agencia', ''),
                col.get('codigo_agencia', ''),
                col.get('tipo_falta', ''),
                col.get('tipo_sancion', ''),
            ]
            for idx, col in enumerate(team, start=1)
        ]
        product_rows = []
        for idx, prod in enumerate(products, start=1):
            claim_values = reclamos_por_producto.get(prod['id_producto'], [])
            claims_text = "; ".join(
                f"{rec.get('id_reclamo', '')} / {rec.get('codigo_analitica', '')}"
                for rec in claim_values
            )
            product_rows.append([
                f"Producto {idx}",
                prod.get('id_producto', ''),
                prod.get('id_cliente', ''),
                prod.get('tipo_producto', ''),
                prod.get('canal', ''),
                prod.get('proceso', ''),
                prod.get('categoria1', ''),
                prod.get('categoria2', ''),
                prod.get('modalidad', ''),
                (
                    f"INV:{prod.get('monto_investigado', '')} | PER:{prod.get('monto_perdida_fraude', '')} "
                    f"| FALLA:{prod.get('monto_falla_procesos', '')} | CONT:{prod.get('monto_contingencia', '')} "
                    f"| REC:{prod.get('monto_recuperado', '')} | PAGO:{prod.get('monto_pago_deuda', '')}"
                ),
                claims_text,
            ])
        risk_rows = [
            [
                risk.get('id_riesgo', ''),
                risk.get('lider', ''),
                risk.get('criticidad', ''),
                risk.get('exposicion_residual', ''),
                risk.get('planes_accion', ''),
            ]
            for risk in riesgos
        ]
        norm_rows = [
            [
                norm.get('id_norma', ''),
                norm.get('descripcion', ''),
                norm.get('fecha_vigencia', ''),
            ]
            for norm in normas
        ]
        summary_sentence = (
            f"Se documentaron {len(clients)} clientes, {len(team)} colaboradores y {len(products)} productos. "
            f"El caso está tipificado como {case['categoria1']} / {case['categoria2']} en modalidad {case['modalidad']}."
        )
        return {
            'case': case,
            'analysis': analysis,
            'total_investigado': total_inv,
            'destinatarios_text': destinatarios_text,
            'client_rows': client_rows,
            'team_rows': team_rows,
            'product_rows': product_rows,
            'risk_rows': risk_rows,
            'norm_rows': norm_rows,
            'summary_sentence': summary_sentence,
        }

    def build_markdown_report(self, data):
        """Genera el informe solicitado en formato Markdown."""

        context = self._build_report_context(data)
        case = context['case']
        analysis = context['analysis']

        def md_table(headers, rows):
            if not rows:
                return ["Sin registros."]
            safe = lambda cell: str(cell or '').replace('|', '\\|')
            lines = ["| " + " | ".join(headers) + " |", "| " + " | ".join(['---'] * len(headers)) + " |"]
            for row in rows:
                lines.append("| " + " | ".join(safe(col) for col in row) + " |")
            return lines

        lines = [
            "Banco de Crédito - BCP",
            "SEGURIDAD CORPORATIVA, INVESTIGACIONES & CRIMEN CIBERNÉTICO",
            "INVESTIGACIONES & CIBERCRIMINOLOGÍA",
            f"Informe {case['tipo_informe']} N.{case['id_caso']}",
            f"Dirigido a: {context['destinatarios_text']}",
            (
                "Referencia: "
                f"{len(context['team_rows'])} colaboradores investigados, {len(context['product_rows'])} productos afectados, "
                f"monto investigado total {context['total_investigado']:.2f} y modalidad {case['modalidad']}."
            ),
            "",
            "## 1. Antecedentes",
            analysis.get('antecedentes') or "Pendiente",
            "",
            "## 2. Tabla de clientes",
        ]
        lines.extend(md_table([
            "Cliente", "Tipo ID", "ID", "Flag", "Teléfonos", "Correos", "Direcciones", "Accionado"
        ], context['client_rows']))
        lines.extend([
            "",
            "## 3. Tabla de team members involucrados",
        ])
        lines.extend(md_table([
            "Colaborador", "ID", "Flag", "División", "Área", "Servicio", "Puesto", "Agencia", "Código", "Falta", "Sanción"
        ], context['team_rows']))
        lines.extend([
            "",
            "## 4. Tabla de productos combinado",
        ])
        lines.extend(md_table([
            "Registro", "ID", "Cliente", "Tipo", "Canal", "Proceso", "Cat.1", "Cat.2", "Modalidad", "Montos", "Reclamo/Analítica"
        ], context['product_rows']))
        lines.extend([
            "",
            "## 5. Resumen automatizado",
            context['summary_sentence'],
            "",
            "## 6. Modus Operandi",
            analysis.get('modus_operandi') or "Pendiente",
            "",
            "## 7. Hallazgos Principales",
            analysis.get('hallazgos') or "Pendiente",
            "",
            "## 8. Descargo de colaboradores",
            analysis.get('descargos') or "Pendiente",
            "",
            "## 9. Tabla de riesgos identificados",
        ])
        lines.extend(md_table([
            "ID Riesgo", "Líder", "Criticidad", "Exposición US$", "Planes"
        ], context['risk_rows']))
        lines.extend([
            "",
            "## 10. Tabla de normas transgredidas",
        ])
        lines.extend(md_table([
            "N° de norma", "Descripción", "Fecha de vigencia"
        ], context['norm_rows']))
        lines.extend([
            "",
            "## 11. Conclusiones",
            analysis.get('conclusiones') or "Pendiente",
            "",
            "## 12. Recomendaciones y mejoras de procesos",
            analysis.get('recomendaciones') or "Pendiente",
            "",
        ])
        return "\n".join(lines)

    def _create_word_document(self):
        if DocxDocument is not None:
            return DocxDocument()
        return _FallbackDocxDocument()

    def build_word_report(self, data):
        """Genera el informe solicitado en formato Word (.docx)."""

        context = self._build_report_context(data)
        case = context['case']
        analysis = context['analysis']
        document = self._create_word_document()

        header_lines = [
            "Banco de Crédito - BCP",
            "SEGURIDAD CORPORATIVA, INVESTIGACIONES & CRIMEN CIBERNÉTICO",
            "INVESTIGACIONES & CIBERCRIMINOLOGÍA",
            f"Informe {case['tipo_informe']} N.{case['id_caso']}",
            f"Dirigido a: {context['destinatarios_text']}",
            (
                "Referencia: "
                f"{len(context['team_rows'])} colaboradores investigados, {len(context['product_rows'])} productos afectados, "
                f"monto investigado total {context['total_investigado']:.2f} y modalidad {case['modalidad']}."
            ),
            "",
        ]
        for line in header_lines:
            document.add_paragraph(line)

        def add_heading_and_text(title, body_text):
            document.add_heading(title, level=2)
            document.add_paragraph(body_text or "Pendiente")
            document.add_paragraph("")

        def append_table_section(title, headers, rows):
            document.add_heading(title, level=2)
            if not rows:
                document.add_paragraph("Sin registros.")
                document.add_paragraph("")
                return
            table = document.add_table(rows=len(rows) + 1, cols=len(headers))
            header_cells = table.rows[0].cells
            for idx, header in enumerate(headers):
                header_cells[idx].text = header
            for row_values in rows:
                new_row = table.add_row().cells
                for idx, value in enumerate(row_values):
                    new_row[idx].text = str(value or "")
            document.add_paragraph("")

        add_heading_and_text("1. Antecedentes", analysis.get('antecedentes'))
        append_table_section(
            "2. Tabla de clientes",
            ["Cliente", "Tipo ID", "ID", "Flag", "Teléfonos", "Correos", "Direcciones", "Accionado"],
            context['client_rows'],
        )
        append_table_section(
            "3. Tabla de team members involucrados",
            ["Colaborador", "ID", "Flag", "División", "Área", "Servicio", "Puesto", "Agencia", "Código", "Falta", "Sanción"],
            context['team_rows'],
        )
        append_table_section(
            "4. Tabla de productos combinado",
            ["Registro", "ID", "Cliente", "Tipo", "Canal", "Proceso", "Cat.1", "Cat.2", "Modalidad", "Montos", "Reclamo/Analítica"],
            context['product_rows'],
        )
        document.add_heading("5. Resumen automatizado", level=2)
        document.add_paragraph(context['summary_sentence'])
        document.add_paragraph("")
        add_heading_and_text("6. Modus Operandi", analysis.get('modus_operandi'))
        add_heading_and_text("7. Hallazgos Principales", analysis.get('hallazgos'))
        add_heading_and_text("8. Descargo de colaboradores", analysis.get('descargos'))
        append_table_section(
            "9. Tabla de riesgos identificados",
            ["ID Riesgo", "Líder", "Criticidad", "Exposición US$", "Planes"],
            context['risk_rows'],
        )
        append_table_section(
            "10. Tabla de normas transgredidas",
            ["N° de norma", "Descripción", "Fecha de vigencia"],
            context['norm_rows'],
        )
        add_heading_and_text("11. Conclusiones", analysis.get('conclusiones'))
        add_heading_and_text("12. Recomendaciones y mejoras de procesos", analysis.get('recomendaciones'))
        return document

    def save_and_send(self):
        """Valida los datos y guarda CSVs normalizados y JSON en la carpeta de exportación."""
        self.flush_autosave()
        self.flush_logs_now()
        errors, warnings = self.validate_data()
        if errors:
            messagebox.showerror("Errores de validación", "\n".join(errors))
            log_event("validacion", f"Errores al guardar: {errors}", self.logs)
        if warnings:
            messagebox.showwarning("Advertencias de validación", "\n".join(warnings))
            log_event("validacion", f"Advertencias al guardar: {warnings}", self.logs)
        if errors:
            return
        folder = self._get_exports_folder()
        if not folder:
            return
        folder = Path(folder)
        # Reunir datos
        data = self.gather_data()
        self._normalize_export_amount_strings(data.get('productos'))
        # Completar claves foráneas con id_caso
        for c in data['clientes']:
            c['id_caso'] = data['caso']['id_caso']
        for t in data['colaboradores']:
            t['id_caso'] = data['caso']['id_caso']
        for p in data['productos']:
            p['id_caso'] = data['caso']['id_caso']
        for r in data['reclamos']:
            r['id_caso'] = data['caso']['id_caso']
        for i in data['involucramientos']:
            i['id_caso'] = data['caso']['id_caso']
        for risk in data['riesgos']:
            risk['id_caso'] = data['caso']['id_caso']
        for norm in data['normas']:
            norm['id_caso'] = data['caso']['id_caso']
        # Guardar CSVs
        case_id = data['caso']['id_caso'] or 'caso'
        created_files = []

        def write_csv(file_name, rows, header):
            path = folder / f"{case_id}_{file_name}"
            with path.open('w', newline='', encoding="utf-8") as f:
                writer = csv.DictWriter(f, fieldnames=header)
                writer.writeheader()
                for row in rows:
                    writer.writerow(row)
            created_files.append(path)
        # CASOS
        write_csv('casos.csv', [data['caso']], ['id_caso', 'tipo_informe', 'categoria1', 'categoria2', 'modalidad', 'canal', 'proceso'])
        # CLIENTES
        write_csv('clientes.csv', data['clientes'], ['id_cliente', 'id_caso', 'tipo_id', 'flag', 'telefonos', 'correos', 'direcciones', 'accionado'])
        # COLABORADORES
        write_csv('colaboradores.csv', data['colaboradores'], ['id_colaborador', 'id_caso', 'flag', 'division', 'area', 'servicio', 'puesto', 'nombre_agencia', 'codigo_agencia', 'tipo_falta', 'tipo_sancion'])
        # PRODUCTOS
        write_csv('productos.csv', data['productos'], ['id_producto', 'id_caso', 'id_cliente', 'categoria1', 'categoria2', 'modalidad', 'canal', 'proceso', 'fecha_ocurrencia', 'fecha_descubrimiento', 'monto_investigado', 'tipo_moneda', 'monto_perdida_fraude', 'monto_falla_procesos', 'monto_contingencia', 'monto_recuperado', 'monto_pago_deuda', 'tipo_producto'])
        # PRODUCTO_RECLAMO
        write_csv('producto_reclamo.csv', data['reclamos'], ['id_reclamo', 'id_caso', 'id_producto', 'nombre_analitica', 'codigo_analitica'])
        # INVOLUCRAMIENTO
        write_csv('involucramiento.csv', data['involucramientos'], ['id_producto', 'id_caso', 'id_colaborador', 'monto_asignado'])
        # DETALLES_RIESGO
        write_csv('detalles_riesgo.csv', data['riesgos'], ['id_riesgo', 'id_caso', 'lider', 'descripcion', 'criticidad', 'exposicion_residual', 'planes_accion'])
        # DETALLES_NORMA
        write_csv('detalles_norma.csv', data['normas'], ['id_norma', 'id_caso', 'descripcion', 'fecha_vigencia'])
        # ANALISIS
        write_csv('analisis.csv', [dict({'id_caso': data['caso']['id_caso']}, **data['analisis'])], ['id_caso', 'antecedentes', 'modus_operandi', 'hallazgos', 'descargos', 'conclusiones', 'recomendaciones'])
        # LOGS
        if self.logs:
            write_csv('logs.csv', self.logs, ['timestamp', 'tipo', 'mensaje'])
        # Guardar JSON
        json_path = folder / f"{case_id}_version.json"
        with json_path.open('w', encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        created_files.append(json_path)
        # Guardar informe Markdown
        report_path = folder / f"{case_id}_informe.md"
        with report_path.open('w', encoding='utf-8') as f:
            f.write(self.build_markdown_report(data))
        created_files.append(report_path)
        # Guardar informe Word
        docx_path = folder / f"{case_id}_informe.docx"
        word_document = self.build_word_report(data)
        word_document.save(docx_path)
        created_files.append(docx_path)
        self._mirror_exports_to_external_drive(created_files, case_id)
        messagebox.showinfo(
            "Datos guardados",
            (
                f"Los archivos se han guardado como {case_id}_*.csv, {case_id}_version.json, "
                f"{case_id}_informe.md y {case_id}_informe.docx."
            ),
        )
        log_event("navegacion", "Datos guardados y enviados", self.logs)
        self.flush_logs_now()

    def _mirror_exports_to_external_drive(self, file_paths, case_id: str) -> None:
        normalized_sources = [Path(path) for path in file_paths or [] if path]
        if not normalized_sources:
            return
        external_base = self._get_external_drive_path()
        if not external_base:
            return
        case_label = case_id or 'caso'
        case_folder = external_base / case_label
        try:
            case_folder.mkdir(parents=True, exist_ok=True)
        except OSError as exc:
            message = f"No se pudo crear la carpeta de respaldo {case_folder}: {exc}"
            log_event("validacion", message, self.logs)
            if not getattr(self, '_suppress_messagebox', False):
                messagebox.showwarning("Copia pendiente", message)
            return
        autosave_abs = Path(AUTOSAVE_FILE).resolve()
        failures = []
        for source in normalized_sources:
            if not source.exists():
                continue
            if source.resolve() == autosave_abs:
                continue
            destination = case_folder / source.name
            try:
                shutil.copy2(source, destination)
            except OSError as exc:
                failures.append((source, exc))
                log_event(
                    "validacion",
                    f"No se pudo copiar {source} a {destination}: {exc}",
                    self.logs,
                )
        if failures and not getattr(self, '_suppress_messagebox', False):
            lines = [
                "Se exportaron los archivos, pero algunos no se copiaron al respaldo externo:",
            ]
            for source, exc in failures:
                lines.append(f"- {source.name}: {exc}")
            messagebox.showwarning("Copia incompleta", "\n".join(lines))

    def save_temp_version(self, data=None):
        """Guarda una versión temporal del estado actual del formulario.

        Este método recoge los datos actuales mediante ``gather_data`` y los
        escribe en un archivo JSON con un sufijo de marca de tiempo. El
        fichero se guarda en el mismo directorio que el script y se nombra
        ``<id_caso>_temp_<YYYYMMDD_HHMMSS>.json``. Si no se ha especificado
        un ID de caso todavía, se utiliza ``caso`` como prefijo. El planificador
        de autosave lo ejecuta de forma diferida para consolidar múltiples
        ediciones cercanas.

        Examples:
            >>> app.save_temp_version()
            # Crea un archivo como ``2025-0001_temp_20251114_154501.json`` con
            # el contenido completo del formulario.
        """
        data = data or self.gather_data()
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        case_id = data.get('caso', {}).get('id_caso', '') or 'caso'
        filename = f"{case_id}_temp_{timestamp}.json"
        json_payload = json.dumps(data, ensure_ascii=False, indent=2)
        target_path = Path(BASE_DIR) / filename
        primary_written = False
        try:
            target_path.parent.mkdir(parents=True, exist_ok=True)
        except OSError as exc:
            log_event(
                "validacion",
                f"No se pudo preparar la carpeta local para la versión temporal: {exc}",
                self.logs,
            )
        else:
            try:
                target_path.write_text(json_payload, encoding='utf-8')
                primary_written = True
            except OSError as ex:
                # Registrar en el log pero no interrumpir
                log_event(
                    "validacion",
                    f"Error guardando versión temporal en la carpeta principal: {ex}",
                    self.logs,
                )
        external_base = self._get_external_drive_path()
        if not external_base:
            return
        case_folder = Path(external_base) / case_id
        try:
            case_folder.mkdir(parents=True, exist_ok=True)
        except OSError as exc:
            log_event(
                "validacion",
                f"No se pudo preparar la carpeta externa para {case_id}: {exc}",
                self.logs,
            )
            return
        mirror_path = case_folder / filename
        if primary_written:
            try:
                shutil.copy2(target_path, mirror_path)
                return
            except OSError as exc:
                log_event(
                    "validacion",
                    f"No se pudo copiar la versión temporal a la carpeta externa: {exc}",
                    self.logs,
                )
        try:
            mirror_path.write_text(json_payload, encoding='utf-8')
        except OSError as exc:
            log_event(
                "validacion",
                f"No se pudo escribir la versión temporal en la carpeta externa: {exc}",
                self.logs,
            )


# ---------------------------------------------------------------------------
# API pública

__all__ = ["FraudCaseApp", "should_autofill_field"]


# ---------------------------------------------------------------------------
# Ejecución de la aplicación

def run_app():
    root = tk.Tk()
    app = FraudCaseApp(root)
    root.mainloop()