from __future__ import annotations

import logging
from collections import defaultdict
from collections.abc import Mapping
from dataclasses import dataclass, field
from decimal import Decimal
from pathlib import Path
from typing import Any, Dict, Iterable, List, Mapping, Optional, Set

try:  # python-docx es opcional en tiempo de ejecución
    from docx import Document as DocxDocument
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import RGBColor, Pt
except ImportError:  # pragma: no cover - se usa el respaldo integrado
    DocxDocument = None
    RGBColor = None
    Pt = None
    WD_ALIGN_PARAGRAPH = None

DOCX_AVAILABLE = DocxDocument is not None
DOCX_MISSING_MESSAGE = (
    "La dependencia opcional 'python-docx' no está instalada. "
    "Instálala con 'pip install python-docx' para habilitar el informe Word."
)

import settings
from validators import parse_decimal_amount, sanitize_rich_text
from report.styling_enhancer import apply_cell_shading, apply_header_band, style_section_heading, style_table, style_title


PLACEHOLDER = "-"
LOGGER = logging.getLogger(__name__)


def _group_relations_by_product(items: Iterable[Any]) -> dict[str, list[Mapping[str, Any]]]:
    grouped: dict[str, list[Mapping[str, Any]]] = defaultdict(list)
    for item in items or []:
        if not isinstance(item, Mapping):
            continue
        producto_id = item.get("id_producto", "")
        grouped[producto_id].append(item)
    return grouped


def _iter_product_combinations(
    productos: Iterable[Any],
    reclamos_por_producto: Mapping[str, list[Mapping[str, Any]]],
    involucramientos_por_producto: Mapping[str, list[Mapping[str, Any]]],
):
    for product in productos or []:
        if not isinstance(product, Mapping):
            continue
        producto_id = product.get("id_producto", "")
        invs = involucramientos_por_producto.get(producto_id) or [None]
        claims = reclamos_por_producto.get(producto_id) or [None]
        for inv in invs:
            for claim in claims:
                yield product, inv or {}, claim or {}


def build_llave_tecnica_rows(case_data: Mapping[str, Any] | CaseData) -> tuple[list[dict[str, str]], list[str]]:
    """Construye las combinaciones de llave técnica a partir de los datos capturados.

    Cada fila incluye los campos del tab "caso" combinados con los
    identificadores que conforman la llave técnica: producto, cliente,
    colaborador (Team Member) y reclamo, además de la fecha de
    ocurrencia asociada al producto.
    """

    caso = case_data.get("caso", {}) if isinstance(case_data, Mapping) else {}
    investigator = caso.get("investigador") if isinstance(caso, Mapping) else {}
    investigator = investigator if isinstance(investigator, Mapping) else {}
    case_occurrence = caso.get("fecha_de_ocurrencia")
    case_discovery = caso.get("fecha_de_descubrimiento")
    base_row = {
        "id_caso": caso.get("id_caso", ""),
        "tipo_informe": caso.get("tipo_informe", ""),
        "categoria1": caso.get("categoria1", ""),
        "categoria2": caso.get("categoria2", ""),
        "modalidad": caso.get("modalidad", ""),
        "canal": caso.get("canal", ""),
        "proceso": caso.get("proceso", ""),
        "fecha_de_ocurrencia": case_occurrence or "",
        "fecha_de_descubrimiento": case_discovery or "",
        "centro_costos": caso.get("centro_costos", "") or caso.get("centro_costo", ""),
        "matricula_investigador": caso.get("matricula_investigador", ""),
        "investigador_nombre": caso.get("investigador_nombre", "")
        or investigator.get("nombre", ""),
        "investigador_cargo": caso.get("investigador_cargo", "")
        or investigator.get("cargo", ""),
    }

    header = list(base_row.keys()) + [
        "id_producto",
        "id_cliente",
        "id_colaborador",
        "id_cliente_involucrado",
        "tipo_involucrado",
        "id_reclamo",
        "fecha_ocurrencia",
    ]

    productos = case_data.get("productos") if isinstance(case_data, Mapping) else []
    reclamos = case_data.get("reclamos") if isinstance(case_data, Mapping) else []
    involucramientos = case_data.get("involucramientos") if isinstance(case_data, Mapping) else []

    reclamos_por_producto = _group_relations_by_product(reclamos)
    involucramientos_por_producto = _group_relations_by_product(involucramientos)

    rows: list[dict[str, str]] = []
    for product, inv, claim in _iter_product_combinations(
        productos, reclamos_por_producto, involucramientos_por_producto
    ):
        involucrado_tipo = inv.get("tipo_involucrado", "") if isinstance(inv, Mapping) else ""
        if not involucrado_tipo and isinstance(inv, Mapping):
            if inv.get("id_cliente_involucrado"):
                involucrado_tipo = "cliente"
            elif inv.get("id_colaborador"):
                involucrado_tipo = "colaborador"
        product_occurrence, _product_discovery = _resolve_product_dates(
            product,
            case_occurrence,
            case_discovery,
        )
        rows.append(
            {
                **base_row,
                "id_producto": product.get("id_producto", ""),
                "id_cliente": product.get("id_cliente", ""),
                "id_colaborador": inv.get("id_colaborador", "")
                if involucrado_tipo != "cliente"
                else "",
                "id_cliente_involucrado": inv.get("id_cliente_involucrado", "")
                if involucrado_tipo == "cliente"
                else "",
                "tipo_involucrado": involucrado_tipo,
                "id_reclamo": claim.get("id_reclamo", ""),
                "fecha_ocurrencia": product_occurrence or "",
            }
        )
    return rows, header


def _event_placeholder(value: object, placeholder: str) -> str:
    if value is None:
        return placeholder
    text = str(value)
    if not text.strip():
        return placeholder
    if text == placeholder:
        return f"\\{placeholder}"
    return text


def _has_meaningful_value(value: object) -> bool:
    if value is None:
        return False
    return bool(str(value).strip())


def _resolve_product_dates(
    product: Mapping[str, Any],
    case_occurrence: object,
    case_discovery: object,
) -> tuple[object, object]:
    product_occurrence = product.get("fecha_ocurrencia")
    product_discovery = product.get("fecha_descubrimiento")
    resolved_occurrence = (
        product_occurrence
        if _has_meaningful_value(product_occurrence)
        else case_occurrence
    )
    resolved_discovery = (
        product_discovery
        if _has_meaningful_value(product_discovery)
        else case_discovery
    )
    return resolved_occurrence, resolved_discovery


def build_event_rows(case_data: Mapping[str, Any] | CaseData) -> tuple[list[dict[str, str]], list[str]]:
    """Combina productos con clientes, reclamos e involucramientos vinculados.

    Usa la misma lógica de combinaciones que ``build_llave_tecnica_rows``
    para garantizar que cada fila conserve la llave técnica
    ``[id_caso, id_producto, id_cliente, id_colaborador, id_reclamo, fecha_ocurrencia]``
    y añade atributos relevantes de producto, cliente y colaborador. Cuando
    no exista la relación correspondiente, se rellenan los campos con
    cadenas vacías.
    """

    caso = case_data.get("caso", {}) if isinstance(case_data, Mapping) else {}
    investigator = caso.get("investigador") if isinstance(caso, Mapping) else {}
    investigator = investigator if isinstance(investigator, Mapping) else {}
    raw_analysis = case_data.get("analisis", {}) if isinstance(case_data, Mapping) else {}
    analysis_texts = normalize_analysis_texts(raw_analysis)
    placeholder = settings.EVENTOS_PLACEHOLDER
    case_occurrence = caso.get("fecha_de_ocurrencia")
    case_discovery = caso.get("fecha_de_descubrimiento")
    base_row = {
        "id_caso": _event_placeholder(caso.get("id_caso"), placeholder),
        "tipo_informe": _event_placeholder(caso.get("tipo_informe"), placeholder),
        "categoria1": _event_placeholder(caso.get("categoria1"), placeholder),
        "categoria2": _event_placeholder(caso.get("categoria2"), placeholder),
        "modalidad": _event_placeholder(caso.get("modalidad"), placeholder),
        "canal": _event_placeholder(caso.get("canal"), placeholder),
        "proceso": _event_placeholder(caso.get("proceso"), placeholder),
        "fecha_de_ocurrencia": _event_placeholder(case_occurrence, placeholder),
        "fecha_de_descubrimiento": _event_placeholder(case_discovery, placeholder),
        "centro_costos": _event_placeholder(
            caso.get("centro_costos") or caso.get("centro_costo"), placeholder
        ),
        "matricula_investigador": _event_placeholder(
            caso.get("matricula_investigador"), placeholder
        ),
        "investigador_nombre": _event_placeholder(
            caso.get("investigador_nombre") or investigator.get("nombre"), placeholder
        ),
        "investigador_cargo": _event_placeholder(
            caso.get("investigador_cargo") or investigator.get("cargo"), placeholder
        ),
        "comentario_breve": _event_placeholder(
            analysis_texts.get("comentario_breve"), placeholder
        ),
        "comentario_amplio": _event_placeholder(
            analysis_texts.get("comentario_amplio"), placeholder
        ),
    }
    header = list(settings.EVENTOS_HEADER_CANONICO)

    productos = case_data.get("productos") if isinstance(case_data, Mapping) else []
    reclamos = case_data.get("reclamos") if isinstance(case_data, Mapping) else []
    involucramientos = case_data.get("involucramientos") if isinstance(case_data, Mapping) else []
    clientes = case_data.get("clientes") if isinstance(case_data, Mapping) else []
    colaboradores = case_data.get("colaboradores") if isinstance(case_data, Mapping) else []

    reclamos_por_producto = _group_relations_by_product(reclamos)
    involucramientos_por_producto = _group_relations_by_product(involucramientos)
    clientes_por_id = {
        client.get("id_cliente", ""): client for client in clientes if isinstance(client, Mapping)
    }
    colaboradores_por_id = {
        collaborator.get("id_colaborador", ""): collaborator
        for collaborator in colaboradores
        if isinstance(collaborator, Mapping)
    }

    rows: list[dict[str, str]] = []
    for product, inv, claim in _iter_product_combinations(
        productos, reclamos_por_producto, involucramientos_por_producto
    ):
        client = clientes_por_id.get(product.get("id_cliente", ""), {})
        inv_tipo = inv.get("tipo_involucrado", "") if isinstance(inv, Mapping) else ""
        if not inv_tipo and isinstance(inv, Mapping):
            if inv.get("id_cliente_involucrado"):
                inv_tipo = "cliente"
            elif inv.get("id_colaborador"):
                inv_tipo = "colaborador"
        is_client_involvement = inv_tipo == "cliente"
        collaborator = (
            colaboradores_por_id.get(inv.get("id_colaborador", ""), {})
            if not is_client_involvement
            else {}
        )
        involved_client_id = inv.get("id_cliente_involucrado", "") if is_client_involvement else ""
        involved_client = clientes_por_id.get(involved_client_id, {}) if is_client_involvement else {}
        categoria1 = product.get("categoria1") or base_row.get("categoria1")
        categoria2 = product.get("categoria2") or base_row.get("categoria2")
        modalidad = product.get("modalidad") or base_row.get("modalidad")
        canal = product.get("canal") or base_row.get("canal")
        proceso = product.get("proceso") or base_row.get("proceso")
        product_occurrence, product_discovery = _resolve_product_dates(
            product,
            case_occurrence,
            case_discovery,
        )
        fecha_cese = collaborator.get("fecha_carta_renuncia") if not is_client_involvement else None
        telefono_relacionado = client.get("telefonos")
        correos_relacionado = client.get("correos")
        direcciones_relacionado = client.get("direcciones")
        accionado_relacionado = client.get("accionado")
        monto_perdida_fraude = product.get("monto_perdida_fraude")
        monto_falla_procesos = product.get("monto_falla_procesos")
        monto_contingencia = product.get("monto_contingencia")
        monto_recuperado = product.get("monto_recuperado")
        monto_pago_deuda = product.get("monto_pago_deuda")
        canonical_row = {
            "case_id": _event_placeholder(caso.get("id_caso"), placeholder),
            "tipo_informe": _event_placeholder(caso.get("tipo_informe"), placeholder),
            "categoria_1": _event_placeholder(categoria1, placeholder),
            "categoria_2": _event_placeholder(categoria2, placeholder),
            "modalidad": _event_placeholder(modalidad, placeholder),
            "tipo_de_producto": _event_placeholder(product.get("tipo_producto"), placeholder),
            "canal": _event_placeholder(canal, placeholder),
            "proceso_impactado": _event_placeholder(proceso, placeholder),
            "product_id": _event_placeholder(product.get("id_producto"), placeholder),
            "cod_operation": placeholder,
            "monto_investigado": _event_placeholder(product.get("monto_investigado"), placeholder),
            "tipo_moneda": _event_placeholder(product.get("tipo_moneda"), placeholder),
            "tipo_id_cliente_involucrado": _event_placeholder(
                involved_client.get("tipo_id"), placeholder
            ),
            "client_id_involucrado": _event_placeholder(involved_client_id, placeholder),
            "flag_cliente_involucrado": _event_placeholder(
                involved_client.get("flag"), placeholder
            ),
            "nombres_cliente_involucrado": _event_placeholder(
                involved_client.get("nombres"), placeholder
            ),
            "apellidos_cliente_involucrado": _event_placeholder(
                involved_client.get("apellidos"), placeholder
            ),
            "matricula_colaborador_involucrado": _event_placeholder(
                inv.get("id_colaborador") if not is_client_involvement else None, placeholder
            ),
            "apellido_paterno_involucrado": _event_placeholder(
                collaborator.get("apellidos") if not is_client_involvement else None, placeholder
            ),
            "apellido_materno_involucrado": placeholder,
            "nombres_involucrado": _event_placeholder(
                collaborator.get("nombres") if not is_client_involvement else None, placeholder
            ),
            "division": _event_placeholder(
                collaborator.get("division") if not is_client_involvement else None, placeholder
            ),
            "area": _event_placeholder(
                collaborator.get("area") if not is_client_involvement else None, placeholder
            ),
            "servicio": _event_placeholder(
                collaborator.get("servicio") if not is_client_involvement else None, placeholder
            ),
            "nombre_agencia": _event_placeholder(
                collaborator.get("nombre_agencia") if not is_client_involvement else None, placeholder
            ),
            "codigo_agencia": _event_placeholder(
                collaborator.get("codigo_agencia") if not is_client_involvement else None,
                placeholder,
            ),
            "puesto": _event_placeholder(
                collaborator.get("puesto") if not is_client_involvement else None, placeholder
            ),
            "fecha_cese": _event_placeholder(fecha_cese, placeholder),
            "tipo_de_falta": _event_placeholder(
                collaborator.get("tipo_falta") if not is_client_involvement else None, placeholder
            ),
            "tipo_sancion": _event_placeholder(
                collaborator.get("tipo_sancion") if not is_client_involvement else None,
                placeholder,
            ),
            "fecha_ocurrencia_caso": _event_placeholder(case_occurrence, placeholder),
            "fecha_descubrimiento_caso": _event_placeholder(case_discovery, placeholder),
            "fecha_ocurrencia": _event_placeholder(product_occurrence, placeholder),
            "fecha_descubrimiento": _event_placeholder(product_discovery, placeholder),
            "monto_fraude_interno_soles": _event_placeholder(monto_perdida_fraude, placeholder),
            "monto_fraude_interno_dolares": placeholder,
            "monto_fraude_externo_soles": placeholder,
            "monto_fraude_externo_dolares": placeholder,
            "monto_falla_en_proceso_soles": _event_placeholder(
                monto_falla_procesos, placeholder
            ),
            "monto_falla_en_proceso_dolares": placeholder,
            "monto_contingencia_soles": _event_placeholder(monto_contingencia, placeholder),
            "monto_contingencia_dolares": placeholder,
            "monto_recuperado_soles": _event_placeholder(monto_recuperado, placeholder),
            "monto_recuperado_dolares": placeholder,
            "monto_pagado_soles": _event_placeholder(monto_pago_deuda, placeholder),
            "monto_pagado_dolares": placeholder,
            "comentario_breve": _event_placeholder(
                analysis_texts.get("comentario_breve"), placeholder
            ),
            "comentario_amplio": _event_placeholder(
                analysis_texts.get("comentario_amplio"), placeholder
            ),
            "id_reclamo": _event_placeholder(claim.get("id_reclamo"), placeholder),
            "nombre_analitica": _event_placeholder(
                claim.get("nombre_analitica"), placeholder
            ),
            "codigo_analitica": _event_placeholder(
                claim.get("codigo_analitica"), placeholder
            ),
            "telefonos_cliente_relacionado": _event_placeholder(
                telefono_relacionado, placeholder
            ),
            "correos_cliente_relacionado": _event_placeholder(
                correos_relacionado, placeholder
            ),
            "direcciones_cliente_relacionado": _event_placeholder(
                direcciones_relacionado, placeholder
            ),
            "accionado_cliente_relacionado": _event_placeholder(
                accionado_relacionado, placeholder
            ),
        }
        rows.append(
            {
                **{field: placeholder for field in header},
                **canonical_row,
                **base_row,
                "categoria1": _event_placeholder(categoria1, placeholder),
                "categoria2": _event_placeholder(categoria2, placeholder),
                "modalidad": _event_placeholder(modalidad, placeholder),
                "canal": _event_placeholder(canal, placeholder),
                "proceso": _event_placeholder(proceso, placeholder),
                "id_producto": _event_placeholder(product.get("id_producto"), placeholder),
                "id_cliente": _event_placeholder(product.get("id_cliente"), placeholder),
                "id_colaborador": _event_placeholder(
                    inv.get("id_colaborador") if not is_client_involvement else None, placeholder
                ),
                "id_cliente_involucrado": _event_placeholder(
                    involved_client_id if is_client_involvement else None, placeholder
                ),
                "tipo_involucrado": _event_placeholder(inv_tipo, placeholder),
                "id_reclamo": _event_placeholder(claim.get("id_reclamo"), placeholder),
                "fecha_ocurrencia": _event_placeholder(product_occurrence, placeholder),
                "fecha_descubrimiento": _event_placeholder(product_discovery, placeholder),
                "tipo_producto": _event_placeholder(product.get("tipo_producto"), placeholder),
                "tipo_moneda": _event_placeholder(product.get("tipo_moneda"), placeholder),
                "monto_investigado": _event_placeholder(
                    product.get("monto_investigado"), placeholder
                ),
                "monto_perdida_fraude": _event_placeholder(monto_perdida_fraude, placeholder),
                "monto_falla_procesos": _event_placeholder(monto_falla_procesos, placeholder),
                "monto_contingencia": _event_placeholder(monto_contingencia, placeholder),
                "monto_recuperado": _event_placeholder(monto_recuperado, placeholder),
                "monto_pago_deuda": _event_placeholder(monto_pago_deuda, placeholder),
                "nombre_analitica": _event_placeholder(
                    claim.get("nombre_analitica"), placeholder
                ),
                "codigo_analitica": _event_placeholder(
                    claim.get("codigo_analitica"), placeholder
                ),
                "cliente_nombres": _event_placeholder(client.get("nombres"), placeholder),
                "cliente_apellidos": _event_placeholder(client.get("apellidos"), placeholder),
                "cliente_tipo_id": _event_placeholder(client.get("tipo_id"), placeholder),
                "cliente_flag": _event_placeholder(client.get("flag"), placeholder),
                "cliente_telefonos": _event_placeholder(client.get("telefonos"), placeholder),
                "cliente_correos": _event_placeholder(client.get("correos"), placeholder),
                "cliente_direcciones": _event_placeholder(
                    client.get("direcciones"), placeholder
                ),
                "cliente_accionado": _event_placeholder(client.get("accionado"), placeholder),
                "colaborador_flag": _event_placeholder(
                    collaborator.get("flag") if not is_client_involvement else None, placeholder
                ),
                "colaborador_nombres": _event_placeholder(
                    collaborator.get("nombres") if not is_client_involvement else None,
                    placeholder,
                ),
                "colaborador_apellidos": _event_placeholder(
                    collaborator.get("apellidos") if not is_client_involvement else None,
                    placeholder,
                ),
                "colaborador_division": _event_placeholder(
                    collaborator.get("division") if not is_client_involvement else None,
                    placeholder,
                ),
                "colaborador_area": _event_placeholder(
                    collaborator.get("area") if not is_client_involvement else None, placeholder
                ),
                "colaborador_servicio": _event_placeholder(
                    collaborator.get("servicio") if not is_client_involvement else None,
                    placeholder,
                ),
                "colaborador_puesto": _event_placeholder(
                    collaborator.get("puesto") if not is_client_involvement else None,
                    placeholder,
                ),
                "colaborador_fecha_carta_inmediatez": _event_placeholder(
                    collaborator.get("fecha_carta_inmediatez")
                    if not is_client_involvement
                    else None,
                    placeholder,
                ),
                "colaborador_fecha_carta_renuncia": _event_placeholder(
                    collaborator.get("fecha_carta_renuncia") if not is_client_involvement else None,
                    placeholder,
                ),
                "colaborador_nombre_agencia": _event_placeholder(
                    collaborator.get("nombre_agencia") if not is_client_involvement else None,
                    placeholder,
                ),
                "colaborador_codigo_agencia": _event_placeholder(
                    collaborator.get("codigo_agencia") if not is_client_involvement else None,
                    placeholder,
                ),
                "colaborador_tipo_falta": _event_placeholder(
                    collaborator.get("tipo_falta") if not is_client_involvement else None,
                    placeholder,
                ),
                "colaborador_tipo_sancion": _event_placeholder(
                    collaborator.get("tipo_sancion") if not is_client_involvement else None,
                    placeholder,
                ),
                "monto_asignado": _event_placeholder(inv.get("monto_asignado"), placeholder),
            }
        )

    return rows, header


@dataclass
class CaseData(Mapping):
    """Estructura normalizada del caso y sus entidades relacionadas."""

    caso: Dict[str, Any]
    clientes: List[Dict[str, Any]]
    colaboradores: List[Dict[str, Any]]
    productos: List[Dict[str, Any]]
    reclamos: List[Dict[str, Any]]
    involucramientos: List[Dict[str, Any]]
    riesgos: List[Dict[str, Any]]
    normas: List[Dict[str, Any]]
    analisis: Dict[str, Any]
    encabezado: Dict[str, Any]
    operaciones: List[Dict[str, Any]]
    anexos: List[Dict[str, Any]]
    firmas: List[Dict[str, Any]]
    recomendaciones_categorias: Dict[str, Any]
    _dict_cache: Dict[str, Any] = field(default=None, init=False, repr=False)

    def as_dict(self) -> Dict[str, Any]:
        if self._dict_cache is None:
            self._dict_cache = {
                "caso": self.caso,
                "clientes": self.clientes,
                "colaboradores": self.colaboradores,
                "productos": self.productos,
                "reclamos": self.reclamos,
                "involucramientos": self.involucramientos,
                "riesgos": self.riesgos,
                "normas": self.normas,
                "analisis": self.analisis,
                "encabezado": self.encabezado,
                "operaciones": self.operaciones,
                "anexos": self.anexos,
                "firmas": self.firmas,
                "recomendaciones_categorias": self.recomendaciones_categorias,
            }
        return self._dict_cache

    def __getitem__(self, key: str) -> Any:  # type: ignore[override]
        return self.as_dict()[key]

    def __iter__(self):  # type: ignore[override]
        return iter(self.as_dict())

    def __len__(self):  # type: ignore[override]
        return len(self.as_dict())

    def get(self, key: str, default: Any = None) -> Any:  # type: ignore[override]
        return self.as_dict().get(key, default)

    @classmethod
    def from_mapping(cls, payload: Mapping[str, Any]) -> "CaseData":
        return cls(
            caso=dict(payload.get("caso") or {}),
            clientes=list(payload.get("clientes") or []),
            colaboradores=list(payload.get("colaboradores") or []),
            productos=list(payload.get("productos") or []),
            reclamos=list(payload.get("reclamos") or []),
            involucramientos=list(payload.get("involucramientos") or []),
            riesgos=list(payload.get("riesgos") or []),
            normas=list(payload.get("normas") or []),
            analisis=dict(payload.get("analisis") or {}),
            encabezado=dict(payload.get("encabezado") or {}),
            operaciones=list(payload.get("operaciones") or []),
            anexos=list(payload.get("anexos") or []),
            firmas=list(payload.get("firmas") or []),
            recomendaciones_categorias=dict(payload.get("recomendaciones_categorias") or {}),
        )


def _normalize_report_segment(value: str | None, placeholder: str) -> str:
    text = (value or "").strip() or placeholder
    for ch in '\\/:*?"<>|':
        text = text.replace(ch, "_")
    return text.replace(" ", "_")


def _extract_analysis_text(value: Any) -> str:
    if isinstance(value, Mapping):
        return str(value.get("text") or "")
    return str(value or "")


@dataclass
class TagRange:
    tag: str
    start: int
    end: int


@dataclass
class RichTextLine:
    text: str
    block_tags: Set[str]
    inline_tags: List[TagRange]


def _tk_index_to_offset(index: Any, text: str) -> Optional[int]:
    try:
        line_str, col_str = str(index).split(".")
        line = int(line_str)
        column = int(col_str)
    except (ValueError, AttributeError):
        return None

    if line < 1 or column < 0:
        return None

    segments = text.splitlines(True)
    if line > len(segments):
        return None

    offset = sum(len(segment) for segment in segments[: line - 1]) + column
    if offset > len(text):
        return None
    return offset


def _normalize_tag_ranges(tags: Iterable[Mapping[str, Any]], text: str) -> List[TagRange]:
    normalized: List[TagRange] = []
    for tag_data in tags:
        tag_name = tag_data.get("tag") if isinstance(tag_data, Mapping) else None
        start = _tk_index_to_offset(tag_data.get("start"), text) if isinstance(tag_data, Mapping) else None
        end = _tk_index_to_offset(tag_data.get("end"), text) if isinstance(tag_data, Mapping) else None
        if not tag_name or start is None or end is None or end <= start:
            continue
        normalized.append(TagRange(tag_name, start, end))
    return normalized


def _parse_rich_text_entry(entry: Any) -> tuple[str, List[TagRange]]:
    if isinstance(entry, Mapping):
        text = sanitize_rich_text(entry.get("text"), settings.RICH_TEXT_MAX_CHARS)
        tags = _normalize_tag_ranges(entry.get("tags") or [], text)
    else:
        text = sanitize_rich_text(_extract_analysis_text(entry), settings.RICH_TEXT_MAX_CHARS)
        tags = []
    return text, tags


def _split_rich_text_lines(entry: Any) -> List[RichTextLine]:
    text, tag_ranges = _parse_rich_text_entry(entry)
    if not text and not tag_ranges:
        return []

    lines_with_breaks = text.splitlines(True) or [""]
    line_spans: List[tuple[str, int, int]] = []
    offset = 0
    for raw_line in lines_with_breaks:
        line_text = raw_line.rstrip("\n")
        line_length = len(line_text)
        line_spans.append((line_text, offset, offset + line_length))
        offset += len(raw_line)

    rich_lines: List[RichTextLine] = [RichTextLine(text=span[0], block_tags=set(), inline_tags=[]) for span in line_spans]

    for tag in tag_ranges:
        for idx, (_, start, end) in enumerate(line_spans):
            if tag.end <= start or tag.start >= end:
                continue
            if tag.tag in {"header", "list", "table"}:
                rich_lines[idx].block_tags.add(tag.tag)
            elif tag.tag == "bold":
                relative_start = max(tag.start - start, 0)
                relative_end = min(tag.end, end) - start
                if relative_end > relative_start:
                    rich_lines[idx].inline_tags.append(
                        TagRange(tag.tag, relative_start, relative_end)
                    )

    return rich_lines


def normalize_analysis_texts(analysis: Mapping[str, Any] | None) -> Dict[str, str]:
    keys = [
        "antecedentes",
        "modus_operandi",
        "hallazgos",
        "descargos",
        "conclusiones",
        "recomendaciones",
        "comentario_breve",
        "comentario_amplio",
    ]
    payload = analysis or {}
    normalized = {
        name: _parse_rich_text_entry(payload.get(name))[0]
        for name in keys
    }
    for name, value in payload.items():
        if name in normalized:
            continue
        normalized[name] = _parse_rich_text_entry(value)[0]
    return normalized


def _safe_text(value: Any, *, placeholder: str = PLACEHOLDER) -> str:
    text = str(value or "").strip()
    return text or placeholder


def _inline_markdown(text: str, inline_tags: List[TagRange]) -> str:
    if not inline_tags:
        return text
    events = []
    for tag in inline_tags:
        if tag.tag != "bold":
            continue
        events.append((tag.start, "start"))
        events.append((tag.end, "end"))
    events.sort(key=lambda item: (item[0], 0 if item[1] == "end" else 1))

    chunks: List[str] = []
    cursor = 0
    for position, kind in events:
        if position > cursor:
            chunks.append(text[cursor:position])
        marker = "**"
        chunks.append(marker)
        cursor = position

    chunks.append(text[cursor:])
    return "".join(chunks)


def _markdown_from_rich_text(entry: Any) -> str:
    lines = _split_rich_text_lines(entry)
    if not lines:
        return ""

    output: List[str] = []
    table_buffer: List[str] = []

    def flush_table_buffer():
        if not table_buffer:
            return
        output.append("```")
        output.extend(table_buffer)
        output.append("```")
        table_buffer.clear()

    for line in lines:
        rendered_line = _inline_markdown(line.text, line.inline_tags)
        if "table" in line.block_tags:
            table_buffer.append(rendered_line)
            continue

        flush_table_buffer()

        if "header" in line.block_tags:
            rendered_line = f"### {rendered_line}".rstrip()
        if "list" in line.block_tags:
            rendered_line = f"- {rendered_line}".rstrip()

        output.append(rendered_line)

    flush_table_buffer()
    return "\n".join(output)


def _inline_segments(text: str, inline_tags: List[TagRange]) -> List[tuple[str, bool]]:
    if not inline_tags:
        return [(text, False)]

    points: List[tuple[int, str]] = []
    for tag in inline_tags:
        if tag.tag != "bold":
            continue
        points.append((tag.start, "start"))
        points.append((tag.end, "end"))

    points.sort()
    segments: List[tuple[str, bool]] = []
    last_index = 0
    depth = 0
    for index, action in points:
        if index > last_index:
            segments.append((text[last_index:index], depth > 0))
        depth += 1 if action == "start" else -1
        last_index = index
    if last_index < len(text):
        segments.append((text[last_index:], depth > 0))
    return segments or [(text, False)]


def _add_rich_text_paragraphs(document, entry: Any) -> None:
    lines = _split_rich_text_lines(entry)
    if not lines:
        document.add_paragraph(PLACEHOLDER)
        return

    table_buffer: List[str] = []

    def flush_table_buffer():
        if not table_buffer:
            return
        paragraph = document.add_paragraph()
        run = paragraph.add_run("\n".join(table_buffer))
        run.font.name = "Courier New"
        table_buffer.clear()

    for line in lines:
        segments = _inline_segments(line.text, line.inline_tags)
        rendered_line = "".join(segment for segment, _ in segments)

        if "table" in line.block_tags:
            table_buffer.append(rendered_line)
            continue

        flush_table_buffer()

        if "header" in line.block_tags:
            paragraph = document.add_heading(level=3)
        elif "list" in line.block_tags:
            paragraph = document.add_paragraph(style="List Bullet")
        else:
            paragraph = document.add_paragraph()

        for text_segment, is_bold in segments:
            run = paragraph.add_run(text_segment)
            if is_bold:
                run.bold = True

    flush_table_buffer()

def _format_decimal_value(value: Optional[Decimal]) -> str:
    if value is None:
        return PLACEHOLDER
    return f"{value.quantize(Decimal('0.01'))}"


def _sum_amounts(items: Iterable[Mapping[str, Any]], key: str) -> Decimal:
    total = Decimal("0")
    for item in items:
        amount = parse_decimal_amount(item.get(key)) if isinstance(item, Mapping) else None
        if amount is not None:
            total += amount
    return total


def _aggregate_amounts(
    products: List[Dict[str, Any]],
    encabezado: Mapping[str, Any],
) -> Dict[str, Optional[Decimal]]:
    def get_amount(key: str, fallback_key: Optional[str] = None) -> Optional[Decimal]:
        raw_value = encabezado.get(key) if isinstance(encabezado, Mapping) else None
        if raw_value not in (None, ""):
            parsed = parse_decimal_amount(raw_value)
            if parsed is not None:
                return parsed
        if fallback_key:
            return _sum_amounts(products, fallback_key)
        return None

    perdida_total = get_amount("perdida_total")
    if perdida_total is None:
        perdida_total = _sum_amounts(products, "monto_perdida_fraude")

    return {
        "investigado": get_amount("importe_investigado", "monto_investigado"),
        "contingencia": get_amount("contingencia", "monto_contingencia"),
        "perdida_total": perdida_total if perdida_total != Decimal("0") else perdida_total,
        "normal": get_amount("normal"),
        "vencido": get_amount("vencido"),
        "judicial": get_amount("judicial"),
        "castigo": get_amount("castigo"),
    }


def build_report_filename(tipo_informe: str | None, case_id: str | None, extension: str) -> str:
    safe_tipo_informe = _normalize_report_segment(tipo_informe, "Generico")
    safe_case_id = _normalize_report_segment(case_id, "caso")
    return f"Informe_{safe_tipo_informe}_{safe_case_id}.{extension}"


def _create_word_document():
    if not DOCX_AVAILABLE or DocxDocument is None:
        raise RuntimeError(DOCX_MISSING_MESSAGE)
    template_path = getattr(settings, "REPORT_TEMPLATE_PATH", None)
    if template_path:
        candidate = Path(template_path)
        try:
            if candidate.is_file():
                return DocxDocument(candidate)
            LOGGER.warning("Plantilla de reporte no encontrada en %s. Se usará un documento en blanco.", candidate)
        except Exception as exc:  # pragma: no cover - defensivo ante archivos corruptos
            LOGGER.warning(
                "No se pudo cargar la plantilla de reporte %s. Se usará un documento en blanco. Detalle: %s",
                candidate,
                exc,
            )
    return DocxDocument()


def _build_report_context(case_data: CaseData):
    case = case_data.caso
    analysis = normalize_analysis_texts(case_data.analisis)
    clients = case_data.clientes
    team = case_data.colaboradores
    products = case_data.productos
    operaciones = case_data.operaciones
    riesgos = case_data.riesgos
    normas = case_data.normas
    encabezado = case_data.encabezado or {}
    reclamos = case_data.reclamos or []
    recomendaciones = case_data.recomendaciones_categorias or {}

    destinatarios = encabezado.get("dirigido_a")
    if not destinatarios:
        destinatarios_set = sorted(
            {
                " - ".join(
                    filter(
                        None,
                        [
                            str(col.get("division", "")).strip(),
                            str(col.get("area", "")).strip(),
                            str(col.get("servicio", "")).strip(),
                        ],
                    )
                )
                for col in team
                if any([col.get("division"), col.get("area"), col.get("servicio")])
            }
        )
        destinatarios = ", ".join([d for d in destinatarios_set if d])
    destinatarios_text = destinatarios or PLACEHOLDER

    amounts = _aggregate_amounts(products, encabezado)
    categoria = " / ".join(
        filter(None, [str(case.get("categoria1", "")).strip(), str(case.get("categoria2", "")).strip()])
    )
    tipologia = _safe_text(encabezado.get("tipologia_evento") or case.get("tipologia_evento") or case.get("modalidad"))
    procesos = encabezado.get("procesos_impactados") or ", ".join(
        sorted({str(prod.get("proceso", "")).strip() for prod in products if prod.get("proceso")})
    )
    analiticas = encabezado.get("analitica_contable")
    if not analiticas:
        codigos_analitica = sorted(
            {str(rec.get("codigo_analitica", "")).strip() for rec in reclamos if rec.get("codigo_analitica")}
        )
        analiticas = ", ".join(filter(None, codigos_analitica))

    productos_texto = encabezado.get("producto") or ", ".join(
        sorted({str(prod.get("tipo_producto", "")).strip() or str(prod.get("producto", "")).strip() for prod in products})
    )
    reclamos_ids = sorted(
        {str(rec.get("id_reclamo", "")).strip() for rec in reclamos if rec.get("id_reclamo")}
    )
    reclamos_count = encabezado.get("numero_reclamos")
    if not reclamos_count:
        reclamos_count = len([rec for rec in reclamos if rec]) or len(reclamos_ids)
    reclamos_ids_text = ", ".join(reclamos_ids) if reclamos_ids else PLACEHOLDER

    header_headers = [
        "Dirigido a",
        "Referencia",
        "Área de Reporte",
        "Fecha de reporte",
        "Categoría del evento",
        "Tipología de evento",
        "Importe investigado",
        "Contingencia",
        "Pérdida total",
        "Normal",
        "Vencido",
        "Judicial",
        "Castigo",
        "Analítica Contable",
        "Centro de Costos",
        "Producto",
        "Procesos impactados",
        "N° de Reclamos",
        "ID de Reclamos",
    ]

    referencia = _safe_text(
        encabezado.get("referencia")
        or case.get("referencia")
        or (
            f"{len(team)} colaboradores investigados, {len(products)} productos afectados, "
            f"monto investigado total {_format_decimal_value(amounts['investigado'])} y modalidad {case.get('modalidad', '')}."
        )
    )

    header_row = [
        destinatarios_text,
        referencia,
        _safe_text(encabezado.get("area_reporte") or case.get("area_reporte")),
        _safe_text(encabezado.get("fecha_reporte") or case.get("fecha_reporte")),
        _safe_text(categoria),
        tipologia,
        _format_decimal_value(amounts.get("investigado")),
        _format_decimal_value(amounts.get("contingencia")),
        _format_decimal_value(amounts.get("perdida_total")),
        _format_decimal_value(amounts.get("normal")),
        _format_decimal_value(amounts.get("vencido")),
        _format_decimal_value(amounts.get("judicial")),
        _format_decimal_value(amounts.get("castigo")),
        _safe_text(analiticas),
        _safe_text(encabezado.get("centro_costos") or case.get("centro_costos")),
        _safe_text(productos_texto),
        _safe_text(procesos),
        _safe_text(reclamos_count),
        _safe_text(reclamos_ids_text),
    ]

    def _collaborator_name(record: Mapping[str, Any]) -> str:
        parts = [record.get("nombres"), record.get("apellidos")]
        joined = " ".join(filter(None, parts)).strip()
        if joined:
            return joined
        return _safe_text(record.get("nombre_completo") or record.get("nombres_apellidos") or record.get("nombre"))

    collaborator_rows = [
        [
            _collaborator_name(col),
            _safe_text(col.get("id_colaborador") or col.get("matricula"), placeholder="-"),
            _safe_text(col.get("puesto") or col.get("cargo"), placeholder="-"),
            _safe_text(col.get("tipo_falta") or col.get("falta"), placeholder="-"),
            _safe_text(col.get("fecha_carta_inmediatez") or col.get("fecha_carta_inmediate"), placeholder="-"),
            _safe_text(col.get("fecha_carta_renuncia"), placeholder="-"),
        ]
        for col in team
    ]

    client_rows = [
        [
            _safe_text(client.get("id_cliente"), placeholder="-"),
            _safe_text(client.get("nombres"), placeholder="-"),
            _safe_text(client.get("apellidos"), placeholder="-"),
            _safe_text(client.get("tipo_id"), placeholder="-"),
            _safe_text(client.get("flag"), placeholder="-"),
            _safe_text(client.get("telefonos"), placeholder="-"),
            _safe_text(client.get("correos"), placeholder="-"),
            _safe_text(client.get("direcciones"), placeholder="-"),
            _safe_text(client.get("accionado"), placeholder="-"),
        ]
        for client in clients
    ]

    event_rows, _ = build_event_rows(case_data)
    combined_product_rows = [
        [
            _safe_text(row.get("id_producto"), placeholder="-"),
            _safe_text(row.get("id_cliente"), placeholder="-"),
            _safe_text(row.get("tipo_producto"), placeholder="-"),
            _safe_text(row.get("canal"), placeholder="-"),
            _safe_text(row.get("proceso"), placeholder="-"),
            _format_decimal_value(parse_decimal_amount(row.get("monto_investigado"))),
            _safe_text(row.get("nombre_analitica"), placeholder="-"),
            _safe_text(
                " ".join(
                    filter(
                        None,
                        [row.get("cliente_nombres"), row.get("cliente_apellidos")],
                    )
                ),
                placeholder="-",
            ),
            _safe_text(
                " ".join(
                    filter(
                        None,
                        [row.get("colaborador_nombres"), row.get("colaborador_apellidos")],
                    )
                ),
                placeholder="-",
            ),
        ]
        for row in event_rows
    ]

    def _build_placeholder_operation_rows(count: int = 3) -> List[List[str]]:
        placeholder_cells = [PLACEHOLDER] * 10
        return [[str(idx)] + placeholder_cells for idx in range(1, count + 1)]

    operation_table_rows = _build_placeholder_operation_rows()
    if operaciones and not products:
        operation_table_rows = []
        total_desembolsado = Decimal("0")
        total_saldo = Decimal("0")

        for idx, operation in enumerate(operaciones, start=1):
            desembolsado = parse_decimal_amount(operation.get("importe_desembolsado"))
            saldo = parse_decimal_amount(operation.get("saldo_deudor"))

            if desembolsado is not None:
                total_desembolsado += desembolsado
            if saldo is not None:
                total_saldo += saldo

            operation_table_rows.append(
                [
                    _safe_text(operation.get("numero") or idx, placeholder=str(idx)),
                    _safe_text(operation.get("fecha_aprobacion")),
                    _safe_text(operation.get("cliente")),
                    _safe_text(operation.get("ingreso_bruto_mensual")),
                    _safe_text(operation.get("empresa_empleadora")),
                    _safe_text(operation.get("vendedor_inmueble")),
                    _safe_text(operation.get("vendedor_credito")),
                    _safe_text(operation.get("producto")),
                    _format_decimal_value(desembolsado),
                    _format_decimal_value(saldo),
                    _safe_text(operation.get("status")),
                ]
            )

        if operation_table_rows:
            operation_table_rows.append(
                [
                    "Totales",
                    PLACEHOLDER,
                    PLACEHOLDER,
                    PLACEHOLDER,
                    PLACEHOLDER,
                    PLACEHOLDER,
                    PLACEHOLDER,
                    PLACEHOLDER,
                    _format_decimal_value(total_desembolsado),
                    _format_decimal_value(total_saldo),
                    PLACEHOLDER,
                ]
            )

    risk_rows = [
        [
            _safe_text(risk.get("lider") or risk.get("lider_riesgo"), placeholder="-"),
            _safe_text(risk.get("id_riesgo") or risk.get("id_riesgo_grc"), placeholder="-"),
            _safe_text(risk.get("descripcion") or risk.get("descripcion_riesgo"), placeholder="-"),
            _safe_text(risk.get("criticidad") or risk.get("criticidad_riesgo"), placeholder="-"),
            _format_decimal_value(parse_decimal_amount(risk.get("exposicion_residual"))),
            _safe_text(risk.get("planes_accion") or risk.get("id_plan_accion"), placeholder="-"),
        ]
        for risk in riesgos
    ]

    norm_rows = []
    for norm in normas:
        if not isinstance(norm, Mapping):
            norm = {}
        norm_rows.append(
            [
                _safe_text(norm.get("id_norma") or norm.get("norma"), placeholder="-"),
                _safe_text(
                    norm.get("acapite_inciso") or norm.get("acapite") or norm.get("inciso"),
                    placeholder="-",
                ),
                _safe_text(norm.get("fecha_vigencia") or norm.get("vigencia"), placeholder="-"),
                _safe_text(
                    norm.get("descripcion") or norm.get("detalle_norma") or norm.get("detalle"),
                    placeholder="-",
                ),
                _safe_text(
                    norm.get("detalle_norma")
                    or norm.get("detalle")
                    or norm.get("descripcion"),
                    placeholder="-",
                ),
            ]
        )

    def _normalize_recommendation_list(value: Any) -> List[str]:
        if isinstance(value, str):
            value = [value] if value.strip() else []
        if not isinstance(value, list):
            return []
        return [str(item).strip() for item in value if str(item).strip()]

    rec_operativas = _normalize_recommendation_list(recomendaciones.get("operativo") or recomendaciones.get("operativas"))
    rec_laborales = _normalize_recommendation_list(recomendaciones.get("laboral") or recomendaciones.get("laborales"))
    rec_legales = _normalize_recommendation_list(recomendaciones.get("legal") or recomendaciones.get("legales"))

    if not (rec_operativas or rec_laborales or rec_legales):
        text = analysis.get("recomendaciones", "")
        if text and text != PLACEHOLDER:
            rec_operativas = [text]

    investigator = case.get("investigador") if isinstance(case, Mapping) else {}
    firmas: List[Dict[str, Any]] = []
    if isinstance(investigator, Mapping):
        nombre_investigador = investigator.get("nombre") or ""
        cargo_investigador = investigator.get("cargo") or "Investigador Principal"
        if nombre_investigador or investigator.get("matricula"):
            firmas.append({"nombre": nombre_investigador, "cargo": cargo_investigador})
    matricula_investigador = case.get("matricula_investigador") if isinstance(case, Mapping) else None
    if matricula_investigador and not firmas:
        firmas.append({"nombre": "", "cargo": "Investigador Principal"})

    return {
        "case": case,
        "analysis": analysis,
        "header_headers": header_headers,
        "header_row": header_row,
        "collaborator_rows": collaborator_rows,
        "client_rows": client_rows,
        "combined_product_rows": combined_product_rows,
        "operation_rows": operation_table_rows,
        "risk_rows": risk_rows,
        "norm_rows": norm_rows,
        "recomendaciones": {
            "laboral": rec_laborales,
            "operativo": rec_operativas,
            "legal": rec_legales,
        },
        "anexos": case_data.anexos or [],
        "firmas": firmas,
    }


def _md_table(headers: Iterable[str], rows: List[List[Any]], *, placeholder: str = PLACEHOLDER) -> List[str]:
    if not rows:
        return [placeholder]
    safe = lambda cell: str(cell or '').replace('|', '\\|')
    lines = ["| " + " | ".join(headers) + " |", "| " + " | ".join(['---'] * len(headers)) + " |"]
    for row in rows:
        lines.append("| " + " | ".join(safe(col) for col in row) + " |")
    return lines


def _md_list(items: List[str]) -> List[str]:
    if not items:
        return [PLACEHOLDER]
    return [f"- {item}" for item in items]


def _format_anexos(anexos: List[Dict[str, Any]]) -> List[str]:
    lines: List[str] = []
    for idx, anexo in enumerate(anexos, start=1):
        if isinstance(anexo, Mapping):
            titulo = anexo.get("titulo") or anexo.get("title") or anexo.get("nombre")
            descripcion = anexo.get("descripcion") or anexo.get("detalle") or ""
            label = titulo or f"Anexo {idx}"
            text = f"{label}" + (f" - {descripcion}" if descripcion else "")
            if text.strip():
                lines.append(text.strip())
        elif str(anexo).strip():
            lines.append(str(anexo).strip())
    return lines


def _format_firmas(firmas: List[Dict[str, Any]]) -> List[str]:
    lines: List[str] = []
    for firma in firmas:
        if isinstance(firma, Mapping):
            nombre = firma.get("nombre") or firma.get("responsable") or ""
            cargo = firma.get("cargo") or firma.get("puesto") or ""
            if nombre or cargo:
                parts = [part for part in [nombre, cargo] if part]
                lines.append(" – ".join(parts))
        elif str(firma).strip():
            lines.append(str(firma).strip())
    return lines


def _section_state(has_data: bool) -> str:
    return "Con información" if has_data else PLACEHOLDER


def _build_sections_summary(context: Mapping[str, Any], analysis: Mapping[str, Any]) -> List[List[str]]:
    recommendations = context["recomendaciones"]
    return [
        ["Encabezado Institucional", "Tabla", _section_state(any(val != PLACEHOLDER for val in context["header_row"]))],
        ["Antecedentes", "Narrativa", _section_state(bool(analysis.get("antecedentes")))],
        ["Colaboradores", "Tabla", _section_state(bool(context["collaborator_rows"]))],
        ["Tabla de clientes", "Tabla", _section_state(bool(context["client_rows"]))],
        [
            "Tabla de productos combinado",
            "Tabla",
            _section_state(bool(context["combined_product_rows"])),
        ],
        ["Modus operandi", "Narrativa", _section_state(bool(analysis.get("modus_operandi")))],
        ["Principales Hallazgos", "Tabla + texto", _section_state(bool(context["operation_rows"]))],
        ["Descargos", "Narrativa", _section_state(bool(analysis.get("descargos")))],
        [
            "Riesgos identificados y debilidades de los controles",
            "Tabla",
            _section_state(bool(context["risk_rows"])),
        ],
        ["Normas transgredidas", "Tabla", _section_state(bool(context["norm_rows"]))],
        ["Conclusiones", "Narrativa", _section_state(bool(analysis.get("conclusiones")))],
        [
            "Recomendaciones y Mejoras de Procesos",
            "Listas",
            _section_state(
                bool(recommendations["laboral"] or recommendations["operativo"] or recommendations["legal"])
            ),
        ],
        ["Anexos", "Lista", _section_state(bool(context["anexos"]))],
        ["Firma", "Lista", _section_state(bool(context["firmas"]))],
    ]


def _build_header_markdown_table(header_values: Mapping[str, Any]) -> List[str]:
    """Genera una tabla HTML que replica la estructura del encabezado DOCX."""

    def _value(key: str, *, emphasize: bool = False) -> str:
        text = _safe_text(header_values.get(key, PLACEHOLDER))
        return f"<strong>{text}</strong>" if emphasize else text

    def _cell(tag: str, content: str, *, colspan: int = 1) -> str:
        span_attr = f' colspan="{colspan}"' if colspan > 1 else ""
        return f"<{tag}{span_attr}>{content}</{tag}>"

    def _row(*cells: str) -> str:
        return "<tr>" + "".join(cells) + "</tr>"

    rows = [
        _row(_cell("th", "Dirigido a"), _cell("td", _value("Dirigido a"), colspan=3)),
        _row(_cell("th", "Referencia"), _cell("td", _value("Referencia"), colspan=3)),
        _row(
            _cell("th", "Área de Reporte"),
            _cell("td", _value("Área de Reporte")),
            _cell("th", "Fecha de reporte"),
            _cell("td", _value("Fecha de reporte")),
        ),
        _row(
            _cell("th", "Categoría del evento"),
            _cell("td", _value("Categoría del evento")),
            _cell("th", "Tipología de evento"),
            _cell("td", _value("Tipología de evento")),
        ),
        _row(
            _cell("th", "Importe investigado"),
            _cell("td", _value("Importe investigado", emphasize=True)),
            _cell("th", "Contingencia"),
            _cell("td", _value("Contingencia", emphasize=True)),
        ),
        _row(
            _cell("th", "Pérdida total"),
            _cell("td", _value("Pérdida total", emphasize=True)),
            _cell("th", "Normal"),
            _cell("td", _value("Normal", emphasize=True)),
        ),
        _row(
            _cell("th", "Vencido"),
            _cell("td", _value("Vencido", emphasize=True)),
            _cell("th", "Judicial"),
            _cell("td", _value("Judicial", emphasize=True)),
        ),
        _row(_cell("th", "Castigo"), _cell("td", _value("Castigo", emphasize=True), colspan=3)),
        _row(
            _cell("th", "Analítica Contable"),
            _cell("td", _value("Analítica Contable")),
            _cell("th", "Centro de Costos"),
            _cell("td", _value("Centro de Costos")),
        ),
        _row(
            _cell("th", "Producto"),
            _cell("td", _value("Producto")),
            _cell("th", "Procesos impactados"),
            _cell("td", _value("Procesos impactados")),
        ),
        _row(_cell("th", "N° de Reclamos"), _cell("td", _value("N° de Reclamos"), colspan=3)),
        _row(_cell("th", "ID de Reclamos"), _cell("td", _value("ID de Reclamos"), colspan=3)),
    ]

    return ["<table>", "  <tbody>", *[f"    {row}" for row in rows], "  </tbody>", "</table>", ""]


def build_md(case_data: CaseData) -> str:
    context = _build_report_context(case_data)
    case = context["case"]
    analysis = context["analysis"]
    raw_analysis = case_data.analisis or {}
    analysis_markdown = {
        name: _markdown_from_rich_text(raw_analysis.get(name))
        for name in [
            "antecedentes",
            "modus_operandi",
            "hallazgos",
            "descargos",
            "conclusiones",
        ]
    }

    header_lines = [
        "**BANCO DE CRÉDITO – BCP**",
        "**SEGURIDAD CORPORATIVA, INTELIGENCIA & CRIMEN CIBERNÉTICO**",
        "**INVESTIGACIONES & CIBERCRIMINOLOGÍA**",
        f"**Informe de Gerencia** {case.get('tipo_informe', '')} N° {case.get('id_caso', '')}",
        f"{_safe_text(case.get('lugar'))}, {_safe_text(case.get('fecha_informe'))}",
        "",
        "## Encabezado Institucional",
        "<em>Resumen estructurado del informe con énfasis en los montos críticos.</em>",
    ]

    header_values = dict(zip(context["header_headers"], context["header_row"]))

    lines = list(header_lines)
    lines.extend(_build_header_markdown_table(header_values))
    lines.extend(
        [
            "",
            "## Antecedentes",
            analysis_markdown.get("antecedentes") or PLACEHOLDER,
            "",
            "## Detalle de los Colaboradores Involucrados",
        ]
    )
    lines.extend(
        _md_table(
            [
                "Nombres y Apellidos",
                "Matrícula",
                "Cargo",
                "Falta cometida",
                "Fecha Carta de Inmediatez",
                "Fecha Carta de Renuncia",
            ],
            context["collaborator_rows"],
        )
    )
    lines.extend(
        [
            "",
            "## Modus operandi",
            analysis_markdown.get("modus_operandi") or PLACEHOLDER,
            "",
            "## Principales Hallazgos",
        ]
    )
    lines.extend(
        _md_table(
            [
                "N°",
                "Fecha de aprobación",
                "Cliente / DNI",
                "Ingreso Bruto Mensual",
                "Empresa Empleadora",
                "Vendedor del Inmueble",
                "Vendedor del Crédito",
                "Producto",
                "Importe Desembolsado",
                "Saldo Deudor",
                "Status (BCP/SBS)",
            ],
            context["operation_rows"],
        )
    )
    lines.extend(
        [
            "",
            analysis_markdown.get("hallazgos") or PLACEHOLDER,
            "",
            "## Descargos",
            analysis_markdown.get("descargos") or PLACEHOLDER,
            "",
            "## Riesgos identificados y debilidades de los controles",
        ]
    )
    lines.extend(
        _md_table(
            [
                "Líder del riesgo",
                "ID Riesgo (GRC)",
                "Descripción del riesgo de fraude",
                "Criticidad del riesgo",
                "Exposición residual (USD)",
                "ID Plan de Acción",
            ],
            context["risk_rows"],
        )
    )
    lines.extend(
        [
            "",
            "## Normas transgredidas",
        ]
    )
    lines.extend(
        _md_table(
            [
                "Norma/Política",
                "Acápite/Inciso",
                "Fecha de vigencia",
                "Descripción",
                "Detalle de Norma",
            ],
            context["norm_rows"],
        )
    )
    lines.extend(
        [
            "",
            "## Conclusiones",
            analysis_markdown.get("conclusiones") or PLACEHOLDER,
            "",
            "## Recomendaciones y Mejoras de Procesos",
            "### De carácter laboral",
        ]
    )
    lines.extend(_md_list(context["recomendaciones"]["laboral"]))
    lines.extend(["", "### De carácter operativo"])
    lines.extend(_md_list(context["recomendaciones"]["operativo"]))
    lines.extend(["", "### De carácter legal"])
    lines.extend(_md_list(context["recomendaciones"]["legal"]))
    lines.extend(["", "## Anexos"])
    lines.extend(_md_list(_format_anexos(context["anexos"])))
    lines.extend(["", "## Firma"])
    lines.extend(_md_list(_format_firmas(context["firmas"])))
    lines.extend(["", "## Resumen de Secciones y Tablas del Informe"])
    lines.extend(
        _md_table(
            ["Sección", "Tipo", "Estado"],
            _build_sections_summary(context, analysis),
        )
    )
    return "\n".join(lines)


def build_docx(case_data: CaseData, path: Path | str) -> Path:
    document = _create_word_document()
    context = _build_report_context(case_data)
    case = context["case"]
    analysis = context["analysis"]
    raw_analysis = case_data.analisis or {}

    def _is_nuevo_riesgo_row(row: List[Any]) -> bool:
        return any(str(value).strip().lower() == "nuevo riesgo" for value in row)

    def add_paragraphs(lines: List[str]) -> List[Any]:
        paragraphs: List[Any] = []
        for line in lines:
            paragraphs.append(document.add_paragraph(line))
        return paragraphs

    def append_table(
        headers: List[str],
        rows: List[List[Any]],
        highlight_predicate: Optional[Any] = None,
    ) -> None:
        if not rows:
            document.add_paragraph(PLACEHOLDER)
            return
        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for idx, header in enumerate(headers):
            table.rows[0].cells[idx].text = header
        highlighted_rows: Set[int] = set()
        for row_index, row in enumerate(rows, start=1):
            docx_row = table.add_row()
            if highlight_predicate and highlight_predicate(row):
                highlighted_rows.add(row_index)
                for cell in docx_row.cells:
                    apply_cell_shading(cell, "FFEBEE")
            for idx, value in enumerate(row):
                docx_row.cells[idx].text = str(value or "")
        style_table(table, zebra_skip_rows=highlighted_rows)

    def add_list(items: List[str]) -> None:
        if not items:
            document.add_paragraph(PLACEHOLDER)
            return
        for item in items:
            document.add_paragraph(item, style="List Bullet")

    header_lines = [
        "BANCO DE CRÉDITO – BCP",
        "SEGURIDAD CORPORATIVA, INTELIGENCIA & CRIMEN CIBERNÉTICO",
        "INVESTIGACIONES & CIBERCRIMINOLOGÍA",
        f"Informe de Gerencia {case.get('tipo_informe', '')} N° {case.get('id_caso', '')}",
        f"{_safe_text(case.get('lugar'))}, {_safe_text(case.get('fecha_informe'))}",
    ]

    header_paragraphs = [document.add_paragraph(header_lines[0])] + add_paragraphs(
        header_lines[1:]
    )
    for paragraph in header_paragraphs:
        style_title(paragraph)
    style_section_heading(document.add_heading("Encabezado Institucional", level=2))
    header_values = dict(zip(context["header_headers"], context["header_row"]))
    header_table = document.add_table(rows=12, cols=4)

    def _set_cells(row_idx: int, col_idx: int, label: str, value: Any) -> None:
        header_table.rows[row_idx].cells[col_idx].text = label
        header_table.rows[row_idx].cells[col_idx + 1].text = str(value or "")

    def _merge_value(row_idx: int, start_col: int, end_col: int) -> None:
        if start_col < end_col:
            header_table.cell(row_idx, start_col).merge(header_table.cell(row_idx, end_col))

    _set_cells(0, 0, "Dirigido a", header_values.get("Dirigido a", PLACEHOLDER))
    _merge_value(0, 1, 3)
    _set_cells(1, 0, "Referencia", header_values.get("Referencia", PLACEHOLDER))
    _merge_value(1, 1, 3)

    _set_cells(2, 0, "Área de Reporte", header_values.get("Área de Reporte", PLACEHOLDER))
    header_table.rows[2].cells[2].text = "Fecha de reporte"
    header_table.rows[2].cells[3].text = str(header_values.get("Fecha de reporte", PLACEHOLDER) or "")

    _set_cells(3, 0, "Categoría del evento", header_values.get("Categoría del evento", PLACEHOLDER))
    header_table.rows[3].cells[2].text = "Tipología de evento"
    header_table.rows[3].cells[3].text = str(header_values.get("Tipología de evento", PLACEHOLDER) or "")

    _set_cells(4, 0, "Importe investigado", header_values.get("Importe investigado", PLACEHOLDER))
    header_table.rows[4].cells[2].text = "Contingencia"
    header_table.rows[4].cells[3].text = str(header_values.get("Contingencia", PLACEHOLDER) or "")

    _set_cells(5, 0, "Pérdida total", header_values.get("Pérdida total", PLACEHOLDER))
    header_table.rows[5].cells[2].text = "Normal"
    header_table.rows[5].cells[3].text = str(header_values.get("Normal", PLACEHOLDER) or "")

    _set_cells(6, 0, "Vencido", header_values.get("Vencido", PLACEHOLDER))
    header_table.rows[6].cells[2].text = "Judicial"
    header_table.rows[6].cells[3].text = str(header_values.get("Judicial", PLACEHOLDER) or "")

    _set_cells(7, 0, "Castigo", header_values.get("Castigo", PLACEHOLDER))
    _merge_value(7, 1, 3)

    _set_cells(8, 0, "Analítica Contable", header_values.get("Analítica Contable", PLACEHOLDER))
    header_table.rows[8].cells[2].text = "Centro de Costos"
    header_table.rows[8].cells[3].text = str(header_values.get("Centro de Costos", PLACEHOLDER) or "")

    _set_cells(9, 0, "Producto", header_values.get("Producto", PLACEHOLDER))
    header_table.rows[9].cells[2].text = "Procesos impactados"
    header_table.rows[9].cells[3].text = str(header_values.get("Procesos impactados", PLACEHOLDER) or "")

    _set_cells(10, 0, "N° de Reclamos", header_values.get("N° de Reclamos", PLACEHOLDER))
    _merge_value(10, 1, 3)

    _set_cells(11, 0, "ID de Reclamos", header_values.get("ID de Reclamos", PLACEHOLDER))
    _merge_value(11, 1, 3)
    apply_header_band(header_table.rows[:3], alignment=WD_ALIGN_PARAGRAPH.LEFT)
    for paragraph in header_table.rows[2].cells[3].paragraphs:
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    monetary_rows = [
        (4, 1),
        (4, 3),
        (5, 1),
        (5, 3),
        (6, 1),
        (6, 3),
        (7, 1),
    ]
    for row_idx, col_idx in monetary_rows:
        for run in header_table.rows[row_idx].cells[col_idx].paragraphs[0].runs:
            run.font.bold = True
    style_section_heading(document.add_heading("1. Antecedentes", level=2))
    _add_rich_text_paragraphs(document, raw_analysis.get("antecedentes"))
    style_section_heading(document.add_heading("Detalle de los Colaboradores Involucrados", level=2))
    append_table(
        [
            "Nombres y Apellidos",
            "Matrícula",
            "Cargo",
            "Falta cometida",
            "Fecha Carta de Inmediatez",
            "Fecha Carta de Renuncia",
        ],
        context["collaborator_rows"],
    )
    style_section_heading(document.add_heading("Modus operandi", level=2))
    _add_rich_text_paragraphs(document, raw_analysis.get("modus_operandi"))
    style_section_heading(document.add_heading("Principales Hallazgos", level=2))
    append_table(
        [
            "N°",
            "Fecha de aprobación",
            "Cliente / DNI",
            "Ingreso Bruto Mensual",
            "Empresa Empleadora",
            "Vendedor del Inmueble",
            "Vendedor del Crédito",
            "Producto",
            "Importe Desembolsado",
            "Saldo Deudor",
            "Status (BCP/SBS)",
        ],
        context["operation_rows"],
    )
    _add_rich_text_paragraphs(document, raw_analysis.get("hallazgos"))
    style_section_heading(document.add_heading("Descargos", level=2))
    _add_rich_text_paragraphs(document, raw_analysis.get("descargos"))
    style_section_heading(document.add_heading("Riesgos identificados y debilidades de los controles", level=2))
    append_table(
        [
            "Líder del riesgo",
            "ID Riesgo (GRC)",
            "Descripción del riesgo de fraude",
            "Criticidad del riesgo",
            "Exposición residual (USD)",
            "ID Plan de Acción",
        ],
        context["risk_rows"],
        highlight_predicate=_is_nuevo_riesgo_row,
    )
    style_section_heading(document.add_heading("Normas transgredidas", level=2))
    append_table(
        [
            "Norma/Política",
            "Acápite/Inciso",
            "Fecha de vigencia",
            "Descripción",
            "Detalle de Norma",
        ],
        context["norm_rows"],
    )
    style_section_heading(document.add_heading("Tabla de clientes", level=2))
    append_table(
        [
            "ID Cliente",
            "Nombres",
            "Apellidos",
            "Tipo ID",
            "Flag",
            "Teléfonos",
            "Correos",
            "Direcciones",
            "Accionado",
        ],
        context["client_rows"],
    )
    style_section_heading(document.add_heading("Tabla de productos combinado", level=2))
    append_table(
        [
            "ID Producto",
            "ID Cliente",
            "Tipo de producto",
            "Canal",
            "Proceso",
            "Monto investigado",
            "Analítica",
            "Cliente",
            "Colaborador",
        ],
        context["combined_product_rows"],
    )
    style_section_heading(document.add_heading("Conclusiones", level=2))
    _add_rich_text_paragraphs(document, raw_analysis.get("conclusiones"))
    style_section_heading(document.add_heading("Recomendaciones y Mejoras de Procesos", level=2))
    style_section_heading(document.add_heading("De carácter laboral", level=3))
    add_list(context["recomendaciones"]["laboral"])
    style_section_heading(document.add_heading("De carácter operativo", level=3))
    add_list(context["recomendaciones"]["operativo"])
    style_section_heading(document.add_heading("De carácter legal", level=3))
    add_list(context["recomendaciones"]["legal"])
    style_section_heading(document.add_heading("Anexos", level=2))
    add_list(_format_anexos(context["anexos"]))
    style_section_heading(document.add_heading("Firma", level=2))
    add_list(_format_firmas(context["firmas"]))
    style_section_heading(document.add_heading("Resumen de Secciones y Tablas del Informe", level=2))
    append_table(["Sección", "Tipo", "Estado"], _build_sections_summary(context, analysis))
    document.save(path)
    return Path(path)


def save_md(case_data: CaseData, path: Path | str) -> Path:
    output_path = Path(path)
    output_path.write_text(build_md(case_data), encoding='utf-8')
    return output_path
