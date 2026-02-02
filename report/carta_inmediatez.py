"""Generador de cartas de inmediatez.

Este módulo encapsula la lógica para construir cartas de inmediatez en
formato DOCX, asignar identificadores correlativos por año y registrar
las generaciones en archivos CSV locales e históricos.
"""

from __future__ import annotations

import csv
import socket
from dataclasses import dataclass
from datetime import datetime
from importlib import util as importlib_util
from pathlib import Path
from typing import Iterable, Mapping, Sequence

from validators import normalize_without_accents, sanitize_rich_text

DOCX_AVAILABLE = importlib_util.find_spec("docx") is not None
if DOCX_AVAILABLE:
    from docx import Document
else:  # pragma: no cover - dependencia opcional
    Document = None


class CartaInmediatezError(Exception):
    """Error de negocio al generar cartas de inmediatez."""


@dataclass(slots=True)
class CartaContext:
    case_id: str
    investigator_name: str
    investigator_id: str
    generation_date: datetime


class CartaInmediatezGenerator:
    """Gestiona la generación y persistencia de cartas de inmediatez."""

    CSV_FIELDS = (
        "numero_caso",
        "fecha_generacion",
        "mes",
        "investigador_principal",
        "matricula_investigador",
        "matricula_team_member",
        "Tipo",
        "codigo_agencia",
        "agencia",
        "Numero_de_Carta",
        "Tipo_entrevista",
    )

    HISTORY_FIELDS = (
        "id_carta",
        "matricula_team_member",
        "nombres_team_member",
        "apellidos_team_member",
        "fecha_creacion",
        "numero_caso",
        "matricula_investigador",
        "hostname",
    )

    _HISTORY_FIELD_MAP = {
        "id_carta": "Numero_de_Carta",
        "matricula_team_member": "matricula_team_member",
        "numero_caso": "numero_caso",
    }

    _HISTORY_KEYS = ("numero_caso", "matricula_team_member", "Numero_de_Carta")

    def __init__(
        self,
        exports_dir: Path,
        external_dir: Path | None,
        *,
        renderer=None,
        docx_available: bool | None = None,
    ):
        self.exports_dir = Path(exports_dir)
        self.external_dir = Path(external_dir) if external_dir else None
        self.renderer = renderer or self._render_with_docx
        self.docx_available = DOCX_AVAILABLE if docx_available is None else bool(docx_available)
        self.template_path = self.exports_dir / "cartas" / "plantilla_carta_inmediatez.docx"

    @staticmethod
    def _normalize_identifier(identifier: str | None) -> str:
        return (identifier or "").strip().upper()

    @staticmethod
    def _sanitize_csv_value(value: str | None) -> str:
        sanitized = sanitize_rich_text("" if value is None else value, max_chars=None)
        if sanitized.startswith(("=", "+", "-", "@")):
            return f"'{sanitized}"
        return sanitized

    def _ensure_directories(self) -> Path:
        cartas_dir = self.exports_dir / "cartas"
        cartas_dir.mkdir(parents=True, exist_ok=True)
        self.exports_dir.mkdir(parents=True, exist_ok=True)
        if self.external_dir:
            self.external_dir.mkdir(parents=True, exist_ok=True)
        return cartas_dir

    def _load_records(self, paths: Iterable[Path]) -> list[dict[str, str]]:
        records: list[dict[str, str]] = []
        seen_rows: set[tuple[str, str, str]] = set()
        for path in paths:
            if not path or not path.exists():
                continue
            try:
                with path.open(newline="", encoding="utf-8") as handle:
                    reader = csv.DictReader(line for line in handle if line.strip())
                    fieldnames = [field.strip() for field in (reader.fieldnames or [])]
                    use_history_map = "id_carta" in fieldnames
                    for row in reader:
                        if use_history_map:
                            normalized_row = {
                                target: row.get(source, "") for source, target in self._HISTORY_FIELD_MAP.items()
                            }
                        else:
                            normalized_row = {field: row.get(field, "") for field in self.CSV_FIELDS}
                        normalized = tuple(
                            (normalized_row.get(key) or "").strip() for key in self._HISTORY_KEYS
                        )
                        if normalized in seen_rows:
                            continue
                        seen_rows.add(normalized)
                        records.append(normalized_row)
            except OSError as exc:
                raise CartaInmediatezError(f"No se pudo leer el historial de cartas: {exc}") from exc
        return records

    def _write_records(self, path: Path, rows: Iterable[dict[str, str]], fields: Sequence[str]) -> None:
        existing = path.exists()
        try:
            with path.open("a", newline="", encoding="utf-8") as handle:
                writer = csv.DictWriter(handle, fieldnames=fields)
                if not existing:
                    writer.writeheader()
                for row in rows:
                    writer.writerow({field: self._sanitize_csv_value(row.get(field)) for field in fields})
        except OSError as exc:
            raise CartaInmediatezError(f"No se pudo actualizar {path.name}: {exc}") from exc

    def _ensure_history_schema(self, path: Path) -> None:
        if not path.exists():
            return
        try:
            with path.open(newline="", encoding="utf-8") as handle:
                reader = csv.DictReader(line for line in handle if line.strip())
                fieldnames = [field.strip() for field in (reader.fieldnames or [])]
                if "id_carta" in fieldnames:
                    return
                rows = list(reader)
        except OSError as exc:
            raise CartaInmediatezError(f"No se pudo leer el historial de cartas: {exc}") from exc
        upgraded_rows = [
            {
                "id_carta": row.get("Numero_de_Carta", ""),
                "matricula_team_member": row.get("matricula_team_member", ""),
                "nombres_team_member": "",
                "apellidos_team_member": "",
                "fecha_creacion": row.get("fecha_generacion", ""),
                "numero_caso": row.get("numero_caso", ""),
                "matricula_investigador": row.get("matricula_investigador", ""),
                "hostname": "",
            }
            for row in rows
        ]
        try:
            with path.open("w", newline="", encoding="utf-8") as handle:
                writer = csv.DictWriter(handle, fieldnames=self.HISTORY_FIELDS)
                writer.writeheader()
                for row in upgraded_rows:
                    writer.writerow({field: self._sanitize_csv_value(row.get(field)) for field in self.HISTORY_FIELDS})
        except OSError as exc:
            raise CartaInmediatezError(f"No se pudo actualizar {path.name}: {exc}") from exc

    def _parse_last_sequence(self, records: Sequence[Mapping[str, str]], year: int) -> int:
        max_value = 0
        for row in records:
            card_id = (row.get("Numero_de_Carta") or "").strip()
            if not card_id.endswith(str(year)):
                continue
            try:
                prefix, _year = card_id.split("-", 1)
                if int(_year) != year:
                    continue
                max_value = max(max_value, int(prefix))
            except (ValueError, TypeError):
                continue
        return max_value

    def _allocate_numbers(self, count: int, records: Sequence[Mapping[str, str]], year: int) -> list[str]:
        start = self._parse_last_sequence(records, year) + 1
        return [f"{index:03d}-{year}" for index in range(start, start + count)]

    def _render_with_docx(self, template_path: Path, output_path: Path, placeholders: Mapping[str, str]) -> None:
        if not self.docx_available or Document is None:
            raise CartaInmediatezError(
                "No se puede generar la carta porque falta la dependencia python-docx."
            )
        if not template_path.exists():
            template_path.parent.mkdir(parents=True, exist_ok=True)
            document = Document()
            document.add_paragraph("Lima, {{FECHA_LARGA}}")
            document.add_paragraph()
            document.add_paragraph("Señora")
            document.add_paragraph("{{NOMBRE_COMPLETO}}")
            document.add_paragraph("Matrícula {{MATRICULA}}")
            document.add_paragraph()
            document.add_paragraph("Presente.-")
            document.add_paragraph()
            document.add_paragraph("Sr(a). {{APELLIDOS}}")
            document.add_paragraph()
            document.add_paragraph(
                "Nos dirigimos a usted para hacerle saber que recientemente el Banco ha tomado "
                "conocimiento de determinadas irregularidades ocurridas en el Área {{AREA}} a la cual "
                "usted pertenece."
            )
            document.add_paragraph()
            document.add_paragraph(
                "Con el objeto de determinar la existencia de posibles responsabilidades, así como "
                "recopilar los elementos de prueba necesarios para obtener una determinación del caso, "
                "el Banco ha dispuesto realizar una investigación de los hechos."
            )
            document.add_paragraph()
            document.add_paragraph(
                "En tal sentido, le solicitamos, conforme a lo establecido en el Reglamento Interno de Trabajo "
                "del Banco, nos brinde su colaboración para el esclarecimiento de los hechos materia del proceso "
                "investigatorio referido, agradeciéndole esté usted a disposición del funcionario que se designe, "
                "en las oportunidades que sea requerida y mientras dure el mismo."
            )
            document.add_paragraph()
            document.add_paragraph(
                "Finalmente cumplimos con informarle que esta comunicación no significa de modo alguno una "
                "sanción disciplinaria ni la imputación de responsabilidad alguna."
            )
            document.add_paragraph()
            document.add_paragraph("Atentamente,")
            document.add_paragraph()
            document.add_paragraph("BANCO DE CREDITO DEL PERU BCP")
            document.add_paragraph()
            table = document.add_table(rows=2, cols=2)
            table.style = "Table Grid"
            table.autofit = True
            table.rows[0].cells[0].text = "------------------------------"
            table.rows[0].cells[1].text = "------------------------------"
            table.rows[1].cells[0].text = "Funcionario"
            table.rows[1].cells[1].text = "Funcionario"
            document.add_paragraph()
            document.add_paragraph("c.c.: Dr. Juan Kam – Gerencia de Relaciones Laborales")
            document.add_paragraph("Carta N° {{NUMERO_CARTA}}")
            document.save(template_path)
        document = Document(template_path)

        def _replace_in_text(text: str) -> str:
            updated = text
            for key, value in placeholders.items():
                updated = updated.replace(f"{{{{{key}}}}}", value)
            return updated

        for paragraph in document.paragraphs:
            paragraph.text = _replace_in_text(paragraph.text)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell.text = _replace_in_text(cell.text)
        document.save(output_path)

    def _build_placeholder_map(self, context: CartaContext, row: dict[str, str], member: Mapping[str, str]) -> dict[str, str]:
        full_name = row.get("matricula_team_member", "")
        if member.get("nombres") or member.get("apellidos"):
            full_name = " ".join(
                part for part in (member.get("nombres", "").strip(), member.get("apellidos", "").strip()) if part
            )
        full_name = sanitize_rich_text(full_name, max_chars=None)
        display_name = full_name or row.get("matricula_team_member", "")
        last_name = sanitize_rich_text(member.get("apellidos", ""), max_chars=None)
        area = sanitize_rich_text(member.get("area", ""), max_chars=None)
        return {
            "NUMERO_CARTA": row.get("Numero_de_Carta", ""),
            "NUMERO_CASO": row.get("numero_caso", ""),
            "COLABORADOR": display_name,
            "PUESTO": sanitize_rich_text(member.get("puesto", ""), max_chars=None),
            "AGENCIA": sanitize_rich_text(member.get("nombre_agencia", ""), max_chars=None),
            "INVESTIGADOR": row.get("investigador_principal", ""),
            "FECHA": row.get("fecha_generacion", ""),
            "FECHA_LARGA": self._format_long_date(context.generation_date),
            "NOMBRE_COMPLETO": display_name,
            "MATRICULA": row.get("matricula_team_member", ""),
            "APELLIDOS": last_name,
            "AREA": area,
        }

    def _determine_tipo(self, division: str) -> str:
        normalized = normalize_without_accents(division or "").lower()
        if "comercial" in normalized or normalized == "dcc":
            return "Agencia"
        return "Sede"

    @staticmethod
    def _format_long_date(date_value: datetime) -> str:
        months = [
            "enero",
            "febrero",
            "marzo",
            "abril",
            "mayo",
            "junio",
            "julio",
            "agosto",
            "septiembre",
            "octubre",
            "noviembre",
            "diciembre",
        ]
        month_name = months[date_value.month - 1]
        return f"{date_value.day:02d} {month_name} {date_value.year}"

    def _build_row(
        self,
        context: CartaContext,
        member: Mapping[str, str],
        numero_carta: str,
    ) -> dict[str, str]:
        flag = (member.get("flag") or "").strip().lower()
        return {
            "numero_caso": context.case_id,
            "fecha_generacion": context.generation_date.strftime("%Y-%m-%d"),
            "mes": context.generation_date.strftime("%m"),
            "investigador_principal": context.investigator_name,
            "matricula_investigador": context.investigator_id,
            "matricula_team_member": self._normalize_identifier(member.get("id_colaborador")),
            "Tipo": self._determine_tipo(member.get("division", "")),
            "codigo_agencia": member.get("codigo_agencia", ""),
            "agencia": member.get("nombre_agencia", ""),
            "Numero_de_Carta": numero_carta,
            "Tipo_entrevista": "Involucrado" if flag == "involucrado" else "Informativo",
        }

    def _build_history_row(
        self,
        context: CartaContext,
        member: Mapping[str, str],
        numero_carta: str,
        hostname: str,
    ) -> dict[str, str]:
        return {
            "id_carta": numero_carta,
            "matricula_team_member": self._normalize_identifier(member.get("id_colaborador")),
            "nombres_team_member": sanitize_rich_text(member.get("nombres", ""), max_chars=None),
            "apellidos_team_member": sanitize_rich_text(member.get("apellidos", ""), max_chars=None),
            "fecha_creacion": context.generation_date.strftime("%Y-%m-%d"),
            "numero_caso": context.case_id,
            "matricula_investigador": context.investigator_id,
            "hostname": hostname,
        }

    def _ensure_no_duplicates(
        self,
        records: Sequence[Mapping[str, str]],
        case_id: str,
        member_ids: Iterable[str],
    ) -> None:
        normalized_case = self._normalize_identifier(case_id)
        normalized_members = {self._normalize_identifier(mid) for mid in member_ids}
        for record in records:
            if self._normalize_identifier(record.get("numero_caso")) != normalized_case:
                continue
            member_id = self._normalize_identifier(record.get("matricula_team_member"))
            if member_id and member_id in normalized_members:
                raise CartaInmediatezError(
                    f"Ya existe una carta para el caso {normalized_case} y la matrícula {member_id}."
                )

    def generate_cartas(
        self, case_payload: Mapping[str, object], members: Sequence[Mapping[str, str]]
    ) -> dict[str, list[Path] | list[dict[str, str]]]:
        if not members:
            raise CartaInmediatezError("Debes seleccionar al menos un colaborador para generar la carta.")

        case = case_payload.get("caso") if isinstance(case_payload, Mapping) else {}
        investigator = case.get("investigador") if isinstance(case, Mapping) else {}
        context = CartaContext(
            case_id=self._normalize_identifier((case or {}).get("id_caso")),
            investigator_name=sanitize_rich_text((investigator or {}).get("nombre", ""), max_chars=None),
            investigator_id=self._normalize_identifier((investigator or {}).get("matricula")),
            generation_date=datetime.now(),
        )
        if not context.case_id:
            raise CartaInmediatezError("El número de caso es obligatorio para generar cartas.")
        if not context.investigator_id:
            raise CartaInmediatezError("La matrícula del investigador es obligatoria para generar cartas.")

        cartas_dir = self._ensure_directories()
        history_paths = [
            self.exports_dir / "h_cartas_inmediatez.csv",
            self.exports_dir / "cartas_inmediatez.csv",
        ]
        if self.external_dir:
            history_paths.append(self.external_dir / "h_cartas_inmediatez.csv")
        existing_records = self._load_records(history_paths)

        member_ids = [member.get("id_colaborador", "") for member in members]
        self._ensure_no_duplicates(existing_records, context.case_id, member_ids)

        numbers = self._allocate_numbers(len(members), existing_records, context.generation_date.year)
        created_rows: list[dict[str, str]] = []
        created_history_rows: list[dict[str, str]] = []
        created_files: list[Path] = []
        hostname = socket.gethostname()

        for member, numero_carta in zip(members, numbers):
            row = self._build_row(context, member, numero_carta)
            created_rows.append(row)
            created_history_rows.append(self._build_history_row(context, member, numero_carta, hostname))
            placeholders = self._build_placeholder_map(context, row, member)
            output_name = f"carta_{row['matricula_team_member'] or 'colaborador'}_{numero_carta}.docx"
            output_path = cartas_dir / output_name
            self.renderer(self.template_path, output_path, placeholders)
            created_files.append(output_path)

        self._write_records(self.exports_dir / "cartas_inmediatez.csv", created_rows, self.CSV_FIELDS)
        history_path = self.exports_dir / "h_cartas_inmediatez.csv"
        self._ensure_history_schema(history_path)
        self._write_records(history_path, created_history_rows, self.HISTORY_FIELDS)
        if self.external_dir:
            external_history_path = self.external_dir / "h_cartas_inmediatez.csv"
            self._ensure_history_schema(external_history_path)
            self._write_records(
                external_history_path,
                created_history_rows,
                self.HISTORY_FIELDS,
            )

        return {"files": created_files, "rows": created_rows}
