from __future__ import annotations

import csv
from datetime import datetime
from pathlib import Path

import pytest

from report.carta_inmediatez import CartaInmediatezError, CartaInmediatezGenerator


def _stub_renderer(_template_path: Path, output_path: Path, placeholders: dict[str, str]) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(placeholders.get("NUMERO_CARTA", ""))


def _write_history(path: Path, rows: list[list[str]]) -> None:
    headers = [
        "numero_caso",
        "fecha_generacion",
        "mes",
        "investigador_principal",
        "matricula_investigador",
        "matricula_team_member",
        "Tipo",
        "codigo_agencia",
        "agencia",
        "Numero_de_Carta",
        "Tipo_entrevista",
    ]
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerow(headers)
        for row in rows:
            writer.writerow(row)


def _write_extended_history(path: Path, rows: list[list[str]]) -> None:
    headers = [
        "id_carta",
        "matricula_team_member",
        "nombres_team_member",
        "apellidos_team_member",
        "fecha_creacion",
        "numero_caso",
        "matricula_investigador",
        "hostname",
    ]
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerow(headers)
        for row in rows:
            writer.writerow(row)


def test_carta_generator_appends_history_and_files(tmp_path: Path) -> None:
    year = datetime.now().year
    exports_dir = tmp_path / "exports"
    external_dir = tmp_path / "external"
    previous_row = [
        "2024-0001",
        f"{year}-01-01",
        "01",
        "Investigador",
        "INV001",
        "TM001",
        "Sede",
        "000001",
        "Agencia Central",
        f"002-{year}",
        "Involucrado",
    ]
    _write_history(exports_dir / "h_cartas_inmediatez.csv", [previous_row])
    generator = CartaInmediatezGenerator(exports_dir, external_dir, renderer=_stub_renderer, docx_available=True)

    case_payload = {
        "caso": {
            "id_caso": f"{year}-0101",
            "investigador": {"matricula": "INV999", "nombre": "Investigador X"},
        }
    }
    members = [
        {
            "id_colaborador": "TM010",
            "nombres": "Ana",
            "apellidos": "Pérez",
            "puesto": "Analista",
            "nombre_agencia": "Agencia Central",
            "codigo_agencia": "000123",
            "flag": "Involucrado",
            "division": "Division Comercial",
        },
        {
            "id_colaborador": "TM011",
            "nombres": "Luis",
            "apellidos": "Diaz",
            "puesto": "Gestor",
            "nombre_agencia": "Agencia Central",
            "codigo_agencia": "000123",
            "flag": "Relacionado",
            "division": "Sede Norte",
        },
    ]

    result = generator.generate_cartas(case_payload, members)

    assert len(result["files"]) == 2
    main_history = exports_dir / "cartas_inmediatez.csv"
    history_rows = list(csv.DictReader(main_history.open(encoding="utf-8")))
    assert {row["Numero_de_Carta"] for row in history_rows} == {f"003-{year}", f"004-{year}"}
    external_history = external_dir / "h_cartas_inmediatez.csv"
    assert external_history.exists()
    extended_history = exports_dir / "h_cartas_inmediatez.csv"
    extended_rows = list(csv.DictReader(extended_history.open(encoding="utf-8")))
    assert {row["id_carta"] for row in extended_rows} == {f"002-{year}", f"003-{year}", f"004-{year}"}
    assert {row["numero_caso"] for row in extended_rows} == {f"{year}-0101", "2024-0001"}
    assert {row["matricula_investigador"] for row in extended_rows} == {"INV001", "INV999"}
    new_rows = [row for row in extended_rows if row["numero_caso"] == f"{year}-0101"]
    assert all(row["hostname"] for row in new_rows)
    assert all(path.exists() for path in result["files"])


def test_carta_generator_detects_duplicate(tmp_path: Path) -> None:
    year = datetime.now().year
    exports_dir = tmp_path / "exports"
    case_id = f"{year}-0101"
    duplicate_row = [
        case_id,
        f"{year}-02-02",
        "02",
        "Investigador",
        "INV001",
        "TM010",
        "Sede",
        "000001",
        "Agencia Central",
        f"001-{year}",
        "Involucrado",
    ]
    _write_history(exports_dir / "cartas_inmediatez.csv", [duplicate_row])
    generator = CartaInmediatezGenerator(exports_dir, None, renderer=_stub_renderer, docx_available=True)

    case_payload = {
        "caso": {"id_caso": case_id, "investigador": {"matricula": "INV999", "nombre": "Investigador X"}}
    }
    members = [
        {
            "id_colaborador": "TM010",
            "nombres": "Ana",
            "apellidos": "Pérez",
            "puesto": "Analista",
            "nombre_agencia": "Agencia Central",
            "codigo_agencia": "000123",
            "flag": "Involucrado",
            "division": "Division Comercial",
        }
    ]

    with pytest.raises(CartaInmediatezError):
        generator.generate_cartas(case_payload, members)


def test_carta_generator_reads_extended_history_for_duplicates(tmp_path: Path) -> None:
    year = datetime.now().year
    exports_dir = tmp_path / "exports"
    case_id = f"{year}-0101"
    duplicate_row = [
        f"001-{year}",
        "TM010",
        "Ana",
        "Pérez",
        f"{year}-02-02",
        case_id,
        "INV001",
        "host-01",
    ]
    _write_extended_history(exports_dir / "h_cartas_inmediatez.csv", [duplicate_row])
    generator = CartaInmediatezGenerator(exports_dir, None, renderer=_stub_renderer, docx_available=True)

    case_payload = {
        "caso": {"id_caso": case_id, "investigador": {"matricula": "INV999", "nombre": "Investigador X"}}
    }
    members = [
        {
            "id_colaborador": "TM010",
            "nombres": "Ana",
            "apellidos": "Pérez",
            "puesto": "Analista",
            "nombre_agencia": "Agencia Central",
            "codigo_agencia": "000123",
            "flag": "Involucrado",
            "division": "Division Comercial",
        }
    ]

    with pytest.raises(CartaInmediatezError):
        generator.generate_cartas(case_payload, members)


def test_carta_generator_rejects_duplicate_card_ids_in_history(tmp_path: Path) -> None:
    year = datetime.now().year
    exports_dir = tmp_path / "exports"
    rows = [
        [
            f"{year}-0001",
            f"{year}-03-01",
            "03",
            "Investigador",
            "INV001",
            "TM010",
            "Sede",
            "000001",
            "Agencia Central",
            f"005-{year}",
            "Involucrado",
        ],
        [
            f"{year}-0002",
            f"{year}-03-02",
            "03",
            "Investigador",
            "INV002",
            "TM011",
            "Sede",
            "000002",
            "Agencia Norte",
            f"005-{year}",
            "Informativo",
        ],
    ]
    _write_history(exports_dir / "h_cartas_inmediatez.csv", rows)
    generator = CartaInmediatezGenerator(exports_dir, None, renderer=_stub_renderer, docx_available=True)

    case_payload = {
        "caso": {
            "id_caso": f"{year}-0101",
            "investigador": {"matricula": "INV999", "nombre": "Investigador X"},
        }
    }
    members = [
        {
            "id_colaborador": "TM012",
            "nombres": "Ana",
            "apellidos": "Pérez",
            "puesto": "Analista",
            "nombre_agencia": "Agencia Central",
            "codigo_agencia": "000123",
            "flag": "Involucrado",
            "division": "Division Comercial",
        }
    ]

    with pytest.raises(CartaInmediatezError):
        generator.generate_cartas(case_payload, members)


def test_carta_generator_skips_existing_card_ids(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    year = datetime.now().year
    exports_dir = tmp_path / "exports"
    existing_row = [
        f"{year}-0001",
        f"{year}-01-01",
        "01",
        "Investigador",
        "INV001",
        "TM010",
        "Sede",
        "000001",
        "Agencia Central",
        f"001-{year}",
        "Involucrado",
    ]
    _write_history(exports_dir / "h_cartas_inmediatez.csv", [existing_row])
    generator = CartaInmediatezGenerator(exports_dir, None, renderer=_stub_renderer, docx_available=True)
    monkeypatch.setattr(generator, "_parse_last_sequence", lambda *_: 0)

    case_payload = {
        "caso": {"id_caso": f"{year}-0102", "investigador": {"matricula": "INV999", "nombre": "Investigador X"}}
    }
    members = [
        {
            "id_colaborador": "TM020",
            "nombres": "Luis",
            "apellidos": "Diaz",
            "puesto": "Gestor",
            "nombre_agencia": "Agencia Central",
            "codigo_agencia": "000123",
            "flag": "Relacionado",
            "division": "Sede Norte",
        }
    ]

    result = generator.generate_cartas(case_payload, members)

    assert result["rows"][0]["Numero_de_Carta"] == f"002-{year}"


def test_default_template_matches_required_layout(tmp_path: Path) -> None:
    pytest.importorskip("docx")
    exports_dir = tmp_path / "exports"
    generator = CartaInmediatezGenerator(exports_dir, None, docx_available=True)

    placeholders = {
        "FECHA_LARGA": "01 enero 2024",
        "NOMBRE_COMPLETO": "Ana Pérez",
        "MATRICULA": "T12345",
        "APELLIDOS": "Pérez",
        "AREA": "Área Comercial",
        "NUMERO_CARTA": "001-2024",
        "NUMERO_CASO": "2024-0001",
        "COLABORADOR": "Ana Pérez",
        "PUESTO": "Analista",
        "AGENCIA": "Agencia Central",
        "INVESTIGADOR": "Investigador X",
        "FECHA": "2024-01-01",
    }
    output_path = exports_dir / "cartas" / "carta_prueba.docx"
    output_path.parent.mkdir(parents=True, exist_ok=True)

    generator._render_with_docx(generator.template_path, output_path, placeholders)

    from docx import Document  # imported after importorskip to avoid optional dependency issues

    document = Document(output_path)
    text_blocks = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    assert f"Lima, {placeholders['FECHA_LARGA']}" in text_blocks
    assert "Señora" in text_blocks
    assert f"Matrícula {placeholders['MATRICULA']}" in text_blocks
    assert any("irregularidades" in paragraph and placeholders["AREA"] in paragraph for paragraph in text_blocks)
    assert any(paragraph.startswith("Carta N°") and placeholders["NUMERO_CARTA"] in paragraph for paragraph in text_blocks)
    assert document.tables, "La plantilla debe incluir el bloque de firmas en tabla para evitar superposiciones."
    signature_table = document.tables[0]
    assert "Funcionario" in signature_table.rows[1].cells[0].text
    assert "Funcionario" in signature_table.rows[1].cells[1].text
    assert generator.template_path.exists()


def test_placeholder_map_uses_member_details(tmp_path: Path) -> None:
    captured: list[dict[str, str]] = []

    def renderer(_template_path: Path, output_path: Path, placeholders: dict[str, str]) -> None:
        captured.append(placeholders)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text("placeholder-carta")

    year = datetime.now().year
    exports_dir = tmp_path / "exports"
    generator = CartaInmediatezGenerator(exports_dir, None, renderer=renderer, docx_available=True)

    case_payload = {
        "caso": {
            "id_caso": f"{year}-0101",
            "investigador": {"matricula": "inv123", "nombre": "Investigador X"},
        }
    }
    members = [
        {
            "id_colaborador": "tm010",
            "nombres": "Ana",
            "apellidos": "Pérez Gómez",
            "puesto": "Analista",
            "area": "Área Comercial",
            "nombre_agencia": "Agencia Central",
            "codigo_agencia": "000123",
            "flag": "Involucrado",
            "division": "Division Comercial",
        }
    ]

    generator.generate_cartas(case_payload, members)

    assert captured, "El renderer debe recibir los placeholders generados."
    placeholders = captured[0]
    assert placeholders["NOMBRE_COMPLETO"] == "Ana Pérez Gómez"
    assert placeholders["APELLIDOS"] == "Pérez Gómez"
    assert placeholders["MATRICULA"] == "TM010"
    assert placeholders["AREA"] == "Área Comercial"
