import pytest

import report_builder
import settings

docx = pytest.importorskip("docx")
from docx import Document  # noqa: E402


@pytest.fixture
def minimal_case_data():
    return report_builder.CaseData.from_mapping(
        {
            "caso": {"id_caso": "2024-1111", "tipo_informe": "Inicial", "fecha_informe": "2024-01-01"},
            "encabezado": {},
            "clientes": [],
            "colaboradores": [],
            "productos": [],
            "reclamos": [],
            "involucramientos": [],
            "riesgos": [],
            "normas": [],
            "analisis": {},
            "operaciones": [],
            "anexos": [],
            "firmas": [],
            "recomendaciones_categorias": {},
        }
    )


@pytest.mark.skipif(not report_builder.DOCX_AVAILABLE, reason="python-docx es requerido")
def test_build_docx_uses_template_when_available(tmp_path, monkeypatch, minimal_case_data):
    template_dir = tmp_path / "templates"
    template_dir.mkdir()
    template_path = template_dir / "report_template.dotx"
    template_doc = Document()
    template_doc.add_paragraph("Contenido de plantilla")
    template_doc.save(template_path)

    monkeypatch.setattr(settings, "REPORT_TEMPLATE_PATH", template_path)

    output_path = tmp_path / "exports" / "plantilla.docx"
    output_path.parent.mkdir()
    result_path = report_builder.build_docx(minimal_case_data, output_path)

    assert result_path.exists()
    built_doc = Document(result_path)
    assert any(paragraph.text == "Contenido de plantilla" for paragraph in built_doc.paragraphs)
    assert any("BANCO DE CRÉDITO" in paragraph.text for paragraph in built_doc.paragraphs)


@pytest.mark.skipif(not report_builder.DOCX_AVAILABLE, reason="python-docx es requerido")
def test_build_docx_falls_back_when_template_missing(tmp_path, monkeypatch, caplog, minimal_case_data):
    missing_template = tmp_path / "templates" / "missing_template.dotx"
    monkeypatch.setattr(settings, "REPORT_TEMPLATE_PATH", missing_template)

    output_path = tmp_path / "exports" / "fallback.docx"
    output_path.parent.mkdir()
    caplog.set_level("WARNING")
    result_path = report_builder.build_docx(minimal_case_data, output_path)

    assert result_path.exists()
    assert any("documento en blanco" in message for message in caplog.messages)

    built_doc = Document(result_path)
    assert any("Informe de Gerencia" in paragraph.text for paragraph in built_doc.paragraphs)
