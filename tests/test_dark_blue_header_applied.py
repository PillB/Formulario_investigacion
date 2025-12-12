from pathlib import Path

import pytest
pytest.importorskip("docx")
from docx import Document
from docx.oxml.ns import qn

import report_builder
from report.styling_enhancer import BCP_DARK_BLUE, WHITE
from tests.test_report_builder import sample_case_data


@pytest.mark.skipif(
    not report_builder.DOCX_AVAILABLE, reason=report_builder.DOCX_MISSING_MESSAGE
)
def test_dark_blue_header_applied(tmp_path: Path, sample_case_data):
    output = tmp_path / "header.docx"

    report_builder.build_docx(sample_case_data, output)

    document = Document(output)
    header_cells = document.tables[0].rows[0].cells

    fills = set()
    for cell in header_cells:
        shading = cell._tc.xpath(".//w:shd")
        assert shading, "El encabezado debe incluir sombreado"
        fills.add(shading[0].get(qn("w:fill")))

        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                assert run.font.bold is True
                assert run.font.color.rgb == WHITE
                assert run.font.name == "Segoe UI Semibold"
                assert run.font.size is not None
                assert run.font.size.pt == pytest.approx(11)

    assert fills == {str(BCP_DARK_BLUE)}
